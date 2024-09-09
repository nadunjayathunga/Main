import os
from datetime import datetime, timedelta

import pandas as pd
import matplotlib.pyplot as plt
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
from io import BytesIO
from data import company_data, company_info, doc_styles, table_style, cogs_ledger_map
import statistics
import numpy as np
from itertools import islice
import sys


def data_sources(company_id: int) -> dict:
    path = f'C:\Masters\{company_info[company_id]["data"]["file_name"]}.xlsx'
    fGL: pd.DataFrame = pd.read_excel(io=path, sheet_name='fGL',
                                      usecols=['Bussiness Unit Name', 'Cost Center', 'Voucher Date', 'Credit Amount',
                                               'Debit Amount', 'Narration', 'Ledger_Code', 'Voucher Number',
                                               'Transaction Type'], engine='calamine')
    dEmployee: pd.DataFrame = pd.read_excel(io=path, sheet_name='dEmployee',
                                            usecols=['Employee_Code',
                                                     'Employee_Name', 'Dept', 'doj', 'nationality', 'Gender',
                                                     'termination_date', 'ba', 'hra', 'tra', 'ma', 'oa', 'pda',
                                                     'travel_cost', 'leave_policy', 'emp_type', 'dob',
                                                     'termination_date', 'Designation'],
                                            index_col='Employee_Code',
                                            dtype={'Employee_Name': str, 'Dept': str, 'doj': 'datetime64[ns]',
                                                   'nationality': str, 'Gender': str, 'ba': float, 'hra': float,
                                                   'tra': float, 'ma': float, 'oa': float,
                                                   'pda': float, 'travel_cost': int, 'leave_policy': str,
                                                   'emp_type': str, 'dob': 'datetime64[ns]', 'Designation': str})
    dCoAAdler: pd.DataFrame = pd.read_excel(io=path, sheet_name='dCoAAdler', index_col='Ledger_Code',
                                            usecols=['Third_Level_Group_Name', 'First_Level_Group_Name', 'Ledger_Code',
                                                     'Ledger_Name', 'Second_Level_Group_Name',
                                                     'Fourth_Level_Group_Name'])
    dCustomers: pd.DataFrame = pd.read_excel(io=path, sheet_name='dCustomers',
                                             usecols=['Customer_Code', 'Ledger_Code', 'Cus_Name', 'Type', 'Credit_Days',
                                                      'Date_Established'])
    fOutSourceInv: pd.DataFrame = pd.read_excel(io=path,
                                                usecols=['Invoice_Number', 'Invoice_Date', 'Customer_Code', 'Order_ID',
                                                         'Net_Amount'], sheet_name='fOutSourceInv')
    fAMCInv: pd.DataFrame = pd.read_excel(io=path, sheet_name='fAMCInv',
                                          usecols=['Invoice_Number', 'Invoice_Date', 'Customer_Code', 'Net_Amount',
                                                   'Order_ID'])
    fProInv: pd.DataFrame = pd.read_excel(io=path, sheet_name='fProInv',
                                          usecols=['Invoice_Number', 'Invoice_Date', 'Order_ID', 'Customer_Code',
                                                   'Net_Amount'])
    fCreditNote: pd.DataFrame = pd.read_excel(io=path, sheet_name='fCreditNote',
                                              usecols=['Invoice_Number', 'Invoice_Date', 'Ledger_Code', 'Net_Amount',
                                                       'Order_ID'])
    fBudget: pd.DataFrame = pd.read_excel(io=path,
                                          usecols=['FY', 'Ledger_Code', 'Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul',
                                                   'Aug', 'Sep', 'Oct', 'Nov', 'Dec'], sheet_name='fBudget')
    fCollection: pd.DataFrame = pd.read_excel(io=path, usecols=['Ledger_Code', 'Invoice Number', 'Invoice Amount',
                                                                'Payment Voucher Number', 'Payment Date',
                                                                'Invoice Date'], sheet_name='fCollection',
                                              date_format={
                                                  'Invoice Date': '%d-%b-%y'},
                                              dtype={'Payment Voucher Number': 'str'})
    dCusOrder: pd.DataFrame = pd.read_excel(io=path, usecols=['Order_ID', 'Customer_Code', 'Employee_Code'],
                                            sheet_name='dCusOrder')
    dContracts: pd.DataFrame = pd.read_excel(io=path, usecols=['Order_ID', 'Customer_Code', 'Employee_Code'],
                                             sheet_name='dContracts')
    dOrderAMC: pd.DataFrame = pd.read_excel(io=path, usecols=['Order_ID', 'Customer_Code', 'Employee_Code'],
                                            sheet_name='dOrderAMC')
    fTimesheet: pd.DataFrame = pd.read_excel(io=path, sheet_name='fTimesheet',
                                             usecols=['cost_center', 'job_id', 'v_date'],
                                             dtype={'cost_center': str, 'job_id': str}, parse_dates=['v_date'],
                                             date_format='%Y-%m-%d %H:%M:%S')
    fOT: pd.DataFrame = pd.read_excel(io=path, sheet_name='fOT', usecols=['date', 'Employee_Code', 'job_id', 'net'],
                                      dtype={'date': str, 'Employee_Code': str, 'job_id': str, 'net': float},
                                      engine='calamine')
    dExclude: pd.DataFrame = pd.read_excel(sheet_name='dExclude', io=path)
    return {'fGL': fGL, 'dEmployee': dEmployee, 'dCoAAdler': dCoAAdler, 'fOutSourceInv': fOutSourceInv,
            'fAMCInv': fAMCInv, 'fProInv': fProInv, 'fCreditNote': fCreditNote, 'dCustomers': dCustomers,
            'fBudget': fBudget, 'fCollection': fCollection, 'dContracts': dContracts, 'dCusOrder': dCusOrder,
            'dOrderAMC': dOrderAMC, 'fTimesheet': fTimesheet, 'fOT': fOT, 'dExclude': dExclude}


def first_page(document, report_date: datetime):
    new_section = document.sections[-1]
    new_section.left_margin = Inches(0.4)
    new_section.right_margin = Inches(0.4)
    new_section.top_margin = Inches(0.3)
    new_section.bottom_margin = Inches(0.1)
    new_section.header_distance = Inches(0.1)
    new_section.footer_distance = Inches(0.1)
    logo = document.add_picture(
        f'C:\Masters\images\logo\{company_info[company_id]["data"]["abbr"]}-logo.png')
    logo = document.paragraphs[-1]
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first = document.add_paragraph()
    first.add_run('\n\n\n')
    first_run = first.add_run(
        company_info[company_id]["data"]["long_name"].upper())
    first_run.bold = True
    first_run.font.size = Pt(24)
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER

    second = document.add_paragraph()
    second_run = second.add_run(
        f'For the period ended {report_date.strftime("%Y-%b-%d")}')
    second_run.font.size = Pt(24)
    second.alignment = WD_ALIGN_PARAGRAPH.CENTER

    third = document.add_paragraph()
    third_run = third.add_run('COMPREHENSIVE FINANCIAL STATEMENT ANALYSIS')
    third_run.font.size = Pt(24)
    third.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return document


def business_unit(row, dEmployee: pd.DataFrame, dCoAAdler: pd.DataFrame) -> str:
    elv_groups: list = dCoAAdler.loc[
        dCoAAdler['First_Level_Group_Name'].isin(['Material Parts & Consumables - Projects',
                                                  'Maintenance - Projects', 'Depreciation - Projects',
                                                  'Others - Projects', 'Projects Revenue'])].index.tolist()
    ledger_code: str = row['Ledger_Code']
    cc: str = row['Cost Center']

    if ledger_code in elv_groups:
        return 'ELV-ESS'
    if pd.isna(cc) or cc == '':
        return 'GUARDING-ESS'
    dept = dEmployee.loc[dEmployee.index == cc, 'Dept'].squeeze() if cc in dEmployee.index else 'GUARDING-ESS'
    return 'ELV-ESS' if dept == 'ELV' else 'GUARDING-ESS'


def receipts_recorded(data: pd.DataFrame) -> pd.DataFrame:
    data.rename(columns={'Ledger_Code': 'ledger_code', 'Invoice Number': 'invoice_number',
                         'Invoice Date': 'invoice_date', 'Invoice Amount': 'invoice_amount',
                         'Payment Voucher Number': 'voucher_number', 'Payment Date': 'voucher_date'}, inplace=True)
    data['invoice_date'] = pd.to_datetime(data['invoice_date'], errors='coerce')
    nulldf: pd.DataFrame = data.loc[data['voucher_date'].isna()]
    multidates: pd.DataFrame = data.loc[~data['voucher_date'].isna() & data['voucher_date'].str.contains(pat=',')]
    data['voucher_date'] = pd.to_datetime(data['voucher_date'], errors='coerce')
    singledate: pd.DataFrame = data.loc[~data['voucher_date'].isna()]
    singledate[['voucher_number', 'voucher_amount']] = singledate['voucher_number'].apply(
        lambda name: pd.Series(name.split("-", 1)))

    final_collection_df: pd.DataFrame = pd.DataFrame(columns=[
        'invoice_number', 'ledger_code', 'invoice_date', 'invoice_amount',
        'voucher_number', 'voucher_amount', 'voucher_date'])

    for _, row in multidates.iterrows():
        pv_number = row['voucher_number']
        voucher_number = [None] if isinstance(pv_number, float) else [voucher.split(sep='-')[0] for voucher in
                                                                      pv_number.split(sep=';')]
        voucher_amount = [None] if isinstance(pv_number, float) else [float(voucher.split(sep='-')[1]) for voucher in
                                                                      pv_number.split(sep=';')]
        voucher_date = [None] if isinstance(row['voucher_date'], float) else row['voucher_date'].split(sep=',')
        invoice_number = [row['invoice_number']] if isinstance(pv_number, float) else [row['invoice_number'] for _ in
                                                                                       range(len(voucher_number))]
        ledger_code = [row['ledger_code']] if isinstance(pv_number, float) else [row['ledger_code'] for _ in
                                                                                 range(len(voucher_number))]
        invoice_date = [row['invoice_date']] if isinstance(pv_number, float) else [row['invoice_date'] for _ in
                                                                                   range(len(voucher_number))]
        invoice_amount = [row['invoice_amount']] if isinstance(pv_number, float) else [row['invoice_amount'] for _ in
                                                                                       range(len(voucher_number))]
        collection_df: pd.DataFrame = pd.DataFrame(
            data={'invoice_number': invoice_number, 'ledger_code': ledger_code, 'invoice_date': invoice_date,
                  'invoice_amount': invoice_amount,
                  'voucher_number': voucher_number, 'voucher_amount': voucher_amount, 'voucher_date': voucher_date})
        final_collection_df = pd.concat(i for i in [final_collection_df, collection_df] if not i.empty)
    final_collection_df = pd.concat(i for i in [final_collection_df, nulldf, singledate] if not i.empty)
    final_collection_df['voucher_date'] = pd.to_datetime(final_collection_df['voucher_date'], errors='coerce')
    final_collection_df['voucher_amount'].fillna(value=0, inplace=True)
    final_collection_df['voucher_amount'] = pd.to_numeric(final_collection_df['voucher_amount'])
    return final_collection_df


def preprocessing(data: dict) -> dict:
    fGL: pd.DataFrame = data['fGL']
    dEmployee: pd.DataFrame = data['dEmployee']
    dCoAAdler: pd.DataFrame = data['dCoAAdler']
    fOutSourceInv: pd.DataFrame = data['fOutSourceInv']
    fAMCInv: pd.DataFrame = data['fAMCInv']
    fProInv: pd.DataFrame = data['fProInv']
    fCreditNote: pd.DataFrame = data['fCreditNote']
    dCustomers: pd.DataFrame = data['dCustomers']
    fBudget: pd.DataFrame = data['fBudget']
    fCollection: pd.DataFrame = data['fCollection']
    dContracts: pd.DataFrame = data['dContracts']
    dCusOrder: pd.DataFrame = data['dCusOrder']
    dOrderAMC: pd.DataFrame = data['dOrderAMC']
    fTimesheet: pd.DataFrame = data['fTimesheet']
    fOT: pd.DataFrame = data['fOT']
    dExclude: pd.DataFrame = data['dExclude']

    fGL['Cost Center'] = fGL['Cost Center'].str.split(
        '|', expand=True)[0].str.strip()  # ESS0012 | GAURAV VASHISTH
    fGL['Bussiness Unit Name'] = fGL.apply(
        business_unit, axis=1, args=[dEmployee, dCoAAdler])
    fGL['Narration'] = fGL['Narration'].fillna('')
    fGL['Amount'] = fGL['Credit Amount'] - fGL['Debit Amount']
    fGL.drop(columns=['Credit Amount', 'Debit Amount'], inplace=True)
    fGL.loc[:, 'Voucher Date'] = fGL['Voucher Date'] + pd.offsets.MonthEnd(0)
    dContracts['Order_ID'] = dContracts['Order_ID'].str.split('-', expand=True)[0]
    fOutSourceInv = pd.merge(
        left=fOutSourceInv, right=dCustomers, on='Customer_Code', how='left')
    fAMCInv = pd.merge(left=fAMCInv, right=dCustomers,
                       on='Customer_Code', how='left')
    fProInv = pd.merge(left=fProInv, right=dCustomers,
                       on='Customer_Code', how='left')
    fCreditNote['Net_Amount'] = fCreditNote['Net_Amount'] * -1
    fCreditNote = pd.merge(
        left=fCreditNote, right=dCustomers, on='Ledger_Code', how='left')
    fInvoices: pd.DataFrame = pd.concat(
        [fOutSourceInv, fAMCInv, fProInv, fCreditNote])
    dJobs: pd.DataFrame = pd.concat([dContracts, dCusOrder, dOrderAMC], ignore_index=True)
    fInvoices = pd.merge(left=fInvoices, right=dJobs[['Order_ID', 'Employee_Code']], on='Order_ID',
                         how='left').sort_values(by='Invoice_Date', ascending=True)
    fInvoices['Invoice_Date'] = fInvoices['Invoice_Date'].apply(lambda row: row + relativedelta(day=31))
    fBudget = pd.melt(fBudget, id_vars=[
        'FY', 'Ledger_Code'], var_name='Month', value_name='Amount')
    fBudget['Voucher Date'] = fBudget.apply(
        lambda x: pd.to_datetime(f'{x["FY"]}-{x["Month"]}-01') + relativedelta(day=31), axis=1)
    fBudget.drop(columns=['FY', 'Month'], inplace=True)
    fBudget = fBudget.loc[fBudget['Amount'] != 0]
    fBudget = pd.merge(left=fBudget, right=dCoAAdler,
                       on='Ledger_Code', how='left')
    fBudget['Bussiness Unit Name'] = 'GUARDING-ESS'
    fBudget.loc[fBudget['Fourth_Level_Group_Name'] == 'Expenses', 'Amount'] *= -1
    fCollection = receipts_recorded(data=fCollection)
    fOT['date'] = fOT['date'].str.split(' ', expand=True)[4].str.strip()
    fOT['date'] = pd.to_datetime(fOT['date'], format='%b-%Y') + pd.offsets.MonthEnd(0)
    fOT.fillna(0, inplace=True)
    fOT = fOT.loc[fOT['net'] != 0]
    fOT.loc[:, 'net'] = fOT['net'] * -1
    fTimesheet = fTimesheet.loc[~fTimesheet['job_id'].isin(['discharged', 'not_joined'])]
    fTimesheet.loc[:, 'v_date'] = fTimesheet['v_date'] + pd.offsets.MonthEnd(0)
    return {'fGL': fGL, 'dEmployee': dEmployee, 'dCoAAdler': dCoAAdler, 'fInvoices': fInvoices, 'fBudget': fBudget,
            'dCustomers': dCustomers, 'fCollection': fCollection, 'dJobs': dJobs, 'fTimesheet': fTimesheet, 'fOT': fOT,
            'dExclude': dExclude}


def coa_ordering(dCoAAdler: pd.DataFrame) -> list:
    other_income_df: pd.DataFrame = dCoAAdler.loc[dCoAAdler['Third_Level_Group_Name'] == 'Indirect Income'].copy()
    coa_df: pd.DataFrame = dCoAAdler.loc[dCoAAdler['Third_Level_Group_Name'] != 'Indirect Income'].copy()

    coa_df.sort_index(inplace=True)
    coa_df.reset_index(inplace=True)
    coa_list: list = coa_df['Ledger_Code'].tolist()

    other_income_df.sort_index(inplace=True)
    other_income_df.reset_index(inplace=True)
    other_inc: list = other_income_df['Ledger_Code'].tolist()

    coa_sort_order: dict = coa_df['Ledger_Name'].reset_index().reset_index().set_index(keys='Ledger_Name')[
        'index'].to_dict()

    first_level: np.ndarray = coa_df['First_Level_Group_Name'].unique()
    for i in first_level:
        coa_sort_order[i] = coa_list.index(
            coa_df.loc[(coa_df['First_Level_Group_Name'] == i), 'Ledger_Code'].max()) + 0.1
    second_level: np.ndarray = coa_df['Second_Level_Group_Name'].unique()
    for i in second_level:
        coa_sort_order[i] = coa_list.index(
            coa_df.loc[(coa_df['Second_Level_Group_Name'] == i), 'Ledger_Code'].max()) + 0.2
    third_level: np.ndarray = coa_df['Third_Level_Group_Name'].unique()
    for i in third_level:
        coa_sort_order[i] = coa_list.index(
            coa_df.loc[(coa_df['Third_Level_Group_Name'] == i), 'Ledger_Code'].max()) + 0.3
    forth_level: np.ndarray = coa_df['Fourth_Level_Group_Name'].unique()
    for i in forth_level:
        coa_sort_order[i] = coa_list.index(
            coa_df.loc[(coa_df['Fourth_Level_Group_Name'] == i), 'Ledger_Code'].max()) + 0.4
    coa_sort_order['Gross Profit'] = coa_sort_order['Cost of Sales'] + 0.1

    for i, j in enumerate(other_inc):
        coa_sort_order[other_income_df.loc[other_income_df['Ledger_Code'] == j, 'Ledger_Name'].iloc[0]] = \
            coa_sort_order['Gross Profit'] + i / 10
    first_level_other_inc: np.ndarray = other_income_df['First_Level_Group_Name'].unique()
    for i in first_level_other_inc:
        coa_sort_order[i] = coa_sort_order[other_income_df.loc[other_income_df['Ledger_Code'] == other_income_df.loc[
            (other_income_df['First_Level_Group_Name'] == i), 'Ledger_Code'].max(), 'Ledger_Name'].iloc[0]] + 0.1
    second_level_other_inc: np.ndarray = other_income_df['Second_Level_Group_Name'].unique()
    for i in second_level_other_inc:
        coa_sort_order[i] = coa_sort_order[other_income_df.loc[other_income_df['Ledger_Code'] == other_income_df.loc[
            (other_income_df['Second_Level_Group_Name'] == i), 'Ledger_Code'].max(), 'Ledger_Name'].iloc[0]] + 0.2
    third_level_other_inc: np.ndarray = other_income_df['Third_Level_Group_Name'].unique()
    for i in third_level_other_inc:
        coa_sort_order[i] = coa_sort_order[other_income_df.loc[other_income_df['Ledger_Code'] == other_income_df.loc[
            (other_income_df['Third_Level_Group_Name'] == i), 'Ledger_Code'].max(), 'Ledger_Name'].iloc[0]] + 0.3
    forth_level_other_inc: np.ndarray = other_income_df['Fourth_Level_Group_Name'].unique()
    for i in forth_level_other_inc:
        coa_sort_order[i] = coa_sort_order[other_income_df.loc[other_income_df['Ledger_Code'] == other_income_df.loc[
            (other_income_df['Fourth_Level_Group_Name'] == i), 'Ledger_Code'].max(), 'Ledger_Name'].iloc[0]] + 0.4
    coa_sort_order['Total Overhead'] = coa_sort_order['Expenses'] + 0.1
    coa_sort_order['Net Profit'] = coa_sort_order['Expenses'] + 0.1

    value = coa_sort_order.pop('Direct Income')
    coa_sort_order['Total Revenue'] = value

    # sorted_data = dict(sorted(coa_sort_order.items(), key=lambda item: item[1]))

    for i in ['Due From Related Parties', 'Due To Related Parties', 'Total Equity & Liabilities']:
        if i == 'Due From Related Parties':
            coa_sort_order['Due From Related Parties'] = coa_sort_order.get('Current Assets') - 0.01
        elif i == 'Due To Related Parties':
            coa_sort_order['Due To Related Parties'] = coa_sort_order.get('Current Liabilities') - 0.01
        else:
            coa_sort_order['Total Equity & Liabilities'] = coa_sort_order.get('Equity') + 0.01

    sorted_data = dict(sorted(coa_sort_order.items(), key=lambda item: item[1]))
    return sorted_data


def apply_style_properties(run, properties):
    if 'bold' in properties:
        run.bold = properties['bold']
    if 'size' in properties:
        run.font.size = Pt(properties['size'])
    if 'name' in properties:
        run.font.name = properties['name']
    if 'color' in properties:
        run.font.color.rgb = RGBColor(*properties['color'])


def style_picker(name: str) -> dict:
    return [i[name] for i in doc_styles if name in i][0]


def table_formatter(table_name, style_name: dict, special: list):
    # Set the table style
    table_name.style = 'Table Grid'

    # Get the style configuration
    style = table_style[style_name]

    for element in ['th', 'td']:
        if element == 'th':
            # Format header cells
            for th_row in table_name.rows[0].cells:
                for paragraph in th_row.paragraphs:
                    run = paragraph.runs[0]
                    run.font.size = Pt(style[f'{element}_style']['font_size'])
                    run.font.name = style[f'{element}_style']['font_name']
                    run.font.color.rgb = RGBColor(*style[f'{element}_style']['font_color'])
                    run.bold = style[f'{element}_style']['bold']

                # Set header cell background color
                cell_xml_element = th_row._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement('w:shd')
                shade_obj.set(qn('w:fill'), style[f'{element}_style']['cell_color'])
                table_cell_properties.append(shade_obj)

        else:
            # Format table data cells
            for row_index in range(1, len(table_name.rows)):
                row_has_special = False

                # Check if any cell in the row meets the special criterion
                for cell in table_name.rows[row_index].cells:
                    if cell.text.strip() in special:
                        row_has_special = True
                        break

                # Apply formatting to the entire row if a special cell is found
                for cell in table_name.rows[row_index].cells:
                    cell_style = style[f'{element}_sp_style'] if row_has_special else style[f'{element}_style']

                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(cell_style['font_size'])
                            run.font.name = cell_style['font_name']
                            run.font.color.rgb = RGBColor(*cell_style['font_color'])
                            run.bold = cell_style['bold']

                    # Set cell background color
                    cell_xml_element = cell._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement('w:shd')
                    shade_obj.set(qn('w:fill'), cell_style['cell_color'])
                    table_cell_properties.append(shade_obj)


def empctc(row, dEmployee: pd.DataFrame) -> float:
    policy: str = dEmployee.loc[(dEmployee['Employee_Code'] == row['Employee_Code']), 'leave_policy']
    basic: float = dEmployee.loc[(dEmployee['Employee_Code'] == row['Employee_Code']), 'ba']
    gross: float = \
        dEmployee.loc[
            (dEmployee['Employee_Code'] == row['Employee_Code']), ['ba', 'hra', 'tra', 'ma', 'oa', 'pda']].sum(
            axis=1).values[0]
    ticket_amt: float = dEmployee.loc[(dEmployee['Employee_Code'] == row['Employee_Code']), 'travel_cost']
    try:
        leave_policy, ticket_policy = int(policy.split(sep='-')[1].strip().split(sep=' ')[0]), \
            policy.split(sep='-')[-1].strip().split(sep=' ')[0]
    except IndexError:
        leave_policy, ticket_policy = 1, 1
    eos: float = basic * 12 / 365 * 30 / 12
    leave: float = gross * 12 / 365 * leave_policy / 12
    ticket: float = ticket_amt / (12 if ticket_policy == 'Yearly' else 24)
    ctc: float = gross + eos + leave + ticket
    return ctc


def profitandlossheads(data: pd.DataFrame, start_date: datetime, end_date: datetime, bu: list) -> pd.DataFrame:
    gp_filt = (data['Third_Level_Group_Name'].isin(['Cost of Sales', 'Direct Income'])) & (
            data['Voucher Date'] >= start_date) & (data['Voucher Date'] <= end_date) & (
                  data['Bussiness Unit Name'].isin(bu))
    gp: float = data.loc[gp_filt, 'Amount'].sum()
    oh_filt = (data['Third_Level_Group_Name'].isin(['Overhead', 'Finance Cost'])) & (
            data['Voucher Date'] >= start_date) & (data['Voucher Date'] <= end_date) & (
                  data['Bussiness Unit Name'].isin(bu))
    overhead: float = data.loc[oh_filt, 'Amount'].sum()
    np_filt = (data['Fourth_Level_Group_Name'].isin(['Expenses', 'Income'])) & (data['Voucher Date'] >= start_date) & (
            data['Voucher Date'] <= end_date) & (data['Bussiness Unit Name'].isin(bu))
    np: float = data.loc[np_filt, 'Amount'].sum()
    rev_filt = (data['Third_Level_Group_Name'].isin(['Direct Income'])) & (
            data['Voucher Date'] >= start_date) & (data['Voucher Date'] <= end_date) & (
                   data['Bussiness Unit Name'].isin(bu))
    rev: float = data.loc[rev_filt, 'Amount'].sum()
    gp_row: pd.DataFrame = pd.DataFrame(data={'Amount': [gp], 'Voucher Date': [end_date]}, index=['Gross Profit'])
    oh_row: pd.DataFrame = pd.DataFrame(data={'Amount': [overhead], 'Voucher Date': [end_date]},
                                        index=['Total Overhead'])
    np_row: pd.DataFrame = pd.DataFrame(data={'Amount': [np], 'Voucher Date': [end_date]}, index=['Net Profit'])
    rev_row: pd.DataFrame = pd.DataFrame(data={'Amount': [rev], 'Voucher Date': [end_date]}, index=['Total Revenue'])
    pl_summary: pd.DataFrame = pd.concat([gp_row, oh_row, np_row, rev_row], ignore_index=False)
    return pl_summary


def profitandloss(data: pd.DataFrame, start_date: datetime, bu: list, end_date: datetime, basic_pl: bool = False,
                  mid_pl: bool = False, full_pl: bool = False) -> dict:
    df_basic: pd.DataFrame = pd.DataFrame(data={'Voucher Date': [], 'Amount': []})
    df_basic_bud: pd.DataFrame = pd.DataFrame(data={'Voucher Date': [], 'Amount': []})
    df_mid: pd.DataFrame = pd.DataFrame(data={'Voucher Date': [], 'Amount': []})
    df_mid_bud: pd.DataFrame = pd.DataFrame(data={'Voucher Date': [], 'Amount': []})
    df_full: pd.DataFrame = pd.DataFrame(data={'Voucher Date': [], 'Amount': []})
    df_full_bud: pd.DataFrame = pd.DataFrame(data={'Voucher Date': [], 'Amount': []})
    basic: pd.DataFrame = pd.DataFrame()
    basic_bud: pd.DataFrame = pd.DataFrame()
    mid: pd.DataFrame = pd.DataFrame()
    mid_bud: pd.DataFrame = pd.DataFrame()
    full: pd.DataFrame = pd.DataFrame()
    full_bud: pd.DataFrame = pd.DataFrame()
    month_end_dates = pd.date_range(start=start_date, end=end_date, freq='ME')
    for end in month_end_dates:
        start: datetime = end + relativedelta(day=1)
        indirect_inc_filt = data['Third_Level_Group_Name'].isin(['Indirect Income']) & (
                data['Voucher Date'] >= start) & (data['Voucher Date'] <= end) & (
                                data['Bussiness Unit Name'].isin(bu))
        indirect_inc_brief: pd.DataFrame = data.loc[
            indirect_inc_filt, ['First_Level_Group_Name', 'Voucher Date', 'Amount']].groupby(
            by=['Voucher Date', 'First_Level_Group_Name'], as_index=False).sum().rename(
            columns={'First_Level_Group_Name': 'Description'})
        indirect_inc_brief = indirect_inc_brief.loc[indirect_inc_brief['Amount'] != 0]
        indirect_inc_filt_bud = fBudget['Third_Level_Group_Name'].isin(['Indirect Income']) & (
                fBudget['Voucher Date'] >= start) & (fBudget['Voucher Date'] <= end) & (
                                    fBudget['Bussiness Unit Name'].isin(bu))
        indirect_inc_brief_bud: pd.DataFrame = fBudget.loc[
            indirect_inc_filt_bud, ['First_Level_Group_Name', 'Voucher Date', 'Amount']].groupby(
            by=['Voucher Date', 'First_Level_Group_Name'], as_index=False).sum().rename(
            columns={'First_Level_Group_Name': 'Description'})
        indirect_inc_brief_bud = indirect_inc_brief_bud.loc[indirect_inc_brief_bud['Amount'] != 0]

        overhead_brief_filt = (data['Third_Level_Group_Name'].isin(['Overhead', 'Finance Cost'])) & (
                data['Voucher Date'] >= start) & (data['Voucher Date'] <= end) & (
                                  data['Bussiness Unit Name'].isin(bu))
        overhead_brief_filt_bud = (fBudget['Third_Level_Group_Name'].isin(['Overhead', 'Finance Cost'])) & (
                fBudget['Voucher Date'] >= start) & (fBudget['Voucher Date'] <= end) & (
                                      fBudget['Bussiness Unit Name'].isin(bu))
        summary_actual: pd.DataFrame = profitandlossheads(data=data, start_date=start, end_date=end, bu=bu)
        summary_budget: pd.DataFrame = profitandlossheads(data=fBudget, start_date=start, end_date=end, bu=bu)
        trade_account_filt = data['Third_Level_Group_Name'].isin(['Cost of Sales', 'Direct Income']) & (
                data['Voucher Date'] >= start) & (data['Voucher Date'] <= end) & (
                                 data['Bussiness Unit Name'].isin(bu))
        # basic version
        if basic_pl:
            trade_account_filt_bud = fBudget['Third_Level_Group_Name'].isin(['Cost of Sales', 'Direct Income']) & (
                    fBudget['Voucher Date'] >= start) & (fBudget['Voucher Date'] <= end) & (
                                         fBudget['Bussiness Unit Name'].isin(bu))
            overhead_brief_basic: pd.DataFrame = data.loc[
                overhead_brief_filt, ['First_Level_Group_Name', 'Voucher Date', 'Amount']].groupby(
                by=['Voucher Date', 'First_Level_Group_Name'], as_index=False).sum().rename(
                columns={'First_Level_Group_Name': 'Description'})
            overhead_brief_basic_bud: pd.DataFrame = fBudget.loc[
                overhead_brief_filt_bud, ['First_Level_Group_Name', 'Voucher Date', 'Amount']].groupby(
                by=['Voucher Date', 'First_Level_Group_Name'], as_index=False).sum().rename(
                columns={'First_Level_Group_Name': 'Description'})
            trade_account_brief: pd.DataFrame = data.loc[
                trade_account_filt, ['First_Level_Group_Name', 'Voucher Date', 'Amount']].groupby(
                by=['Voucher Date', 'First_Level_Group_Name'], as_index=False).sum().rename(
                columns={'First_Level_Group_Name': 'Description'})
            trade_account_brief_bud: pd.DataFrame = fBudget.loc[
                trade_account_filt_bud, ['First_Level_Group_Name', 'Voucher Date', 'Amount']].groupby(
                by=['Voucher Date', 'First_Level_Group_Name'], as_index=False).sum().rename(
                columns={'First_Level_Group_Name': 'Description'})
            basic: pd.DataFrame = pd.concat(
                i for i in [trade_account_brief, indirect_inc_brief, overhead_brief_basic] if not i.empty).rename(
                columns={'First_Level_Group_Name': 'Description'})
            basic_bud: pd.DataFrame = pd.concat(
                [trade_account_brief_bud, indirect_inc_brief_bud, overhead_brief_basic_bud]).rename(
                columns={'First_Level_Group_Name': 'Description'})
            basic = basic.loc[basic['Amount'] != 0].set_index(keys='Description')
            basic_bud = basic_bud.loc[basic_bud['Amount'] != 0].set_index(keys='Description')
            df_basic = pd.concat(i for i in [basic, summary_actual, df_basic] if not i.empty)
            df_basic_bud = pd.concat(i for i in [basic_bud, summary_budget, df_basic_bud] if not i.empty)

        # mid version
        if mid_pl:
            trade_account_mid: pd.DataFrame = data.loc[
                trade_account_filt, ['Second_Level_Group_Name', 'Voucher Date', 'Amount']].groupby(
                by=['Voucher Date', 'Second_Level_Group_Name'], as_index=False).sum().rename(
                columns={'Second_Level_Group_Name': 'Description'})
            trade_account_mid_bud: pd.DataFrame = fBudget.loc[
                trade_account_filt_bud, ['Second_Level_Group_Name', 'Voucher Date', 'Amount']].groupby(
                by=['Voucher Date', 'Second_Level_Group_Name'], as_index=False).sum().rename(
                columns={'Second_Level_Group_Name': 'Description'})
            overhead_brief_mid: pd.DataFrame = data.loc[
                overhead_brief_filt, ['First_Level_Group_Name', 'Voucher Date', 'Amount']].groupby(
                by=['Voucher Date', 'First_Level_Group_Name'], as_index=False).sum().rename(
                columns={'First_Level_Group_Name': 'Description'})
            overhead_brief_mid_bud: pd.DataFrame = fBudget.loc[
                overhead_brief_filt_bud, ['First_Level_Group_Name', 'Voucher Date', 'Amount']].groupby(
                by=['Voucher Date', 'First_Level_Group_Name'], as_index=False).sum().rename(
                columns={'First_Level_Group_Name': 'Description'})
            mid = pd.concat(
                [trade_account_mid, indirect_inc_brief, overhead_brief_mid])
            mid_bud = pd.concat(
                [trade_account_mid_bud, indirect_inc_brief_bud, overhead_brief_mid_bud])
            mid = mid.loc[mid['Amount'] != 0].set_index(keys='Description')
            mid_bud = mid_bud.loc[mid_bud['Amount'] != 0].set_index(keys='Description')
            df_mid = pd.concat([mid, summary_actual, df_mid])
            df_mid_bud = pd.concat([mid_bud, summary_budget, df_mid_bud])

        # full version
        if full_pl:
            detailed_filt = data['Fourth_Level_Group_Name'].isin(['Income', 'Expenses']) & (
                    data['Voucher Date'] >= start) & (data['Voucher Date'] <= end) & (
                                data['Bussiness Unit Name'].isin(bu))
            detailed_filt_bud = fBudget['Third_Level_Group_Name'].isin(
                ['Indirect Income', 'Overhead', 'Finance Cost', 'Direct Income', 'Cost of Sales']) & (
                                        fBudget['Voucher Date'] >= start) & (fBudget['Voucher Date'] <= end) & (
                                    fBudget['Bussiness Unit Name'].isin(bu))
            full = data.loc[detailed_filt, ['Ledger_Name', 'Voucher Date', 'Amount']].groupby(
                by=['Voucher Date', 'Ledger_Name'], as_index=False).sum().rename(columns={'Ledger_Name': 'Description'})
            full_bud = fBudget.loc[detailed_filt_bud, ['Ledger_Name', 'Voucher Date', 'Amount']].groupby(
                by=['Voucher Date', 'Ledger_Name'], as_index=False).sum().rename(columns={'Ledger_Name': 'Description'})
            full = full.loc[full['Amount'] != 0].set_index(keys='Description')
            full_bud = full_bud.loc[full_bud['Amount'] != 0].set_index(keys='Description')
            df_full = pd.concat(i for i in [df_full, summary_actual, full] if not i.empty)
            df_full_bud = pd.concat(i for i in [df_full_bud, summary_budget, full_bud] if not i.empty)
    cy_cp_basic: pd.DataFrame = df_basic.loc[df_basic['Voucher Date'] == end_date].drop(
        columns=['Voucher Date']).reset_index().rename(columns={'index': 'Description'})
    cy_cp_basic_bud: pd.DataFrame = df_basic_bud.loc[df_basic_bud['Voucher Date'] == end_date].drop(
        columns=['Voucher Date']).reset_index().rename(columns={'index': 'Description'})

    cy_pp_basic: pd.DataFrame = df_basic.loc[
        df_basic['Voucher Date'] == end_date - relativedelta(months=1) + relativedelta(day=31)].drop(
        columns=['Voucher Date']).reset_index().rename(columns={'index': 'Description'})

    cy_ytd_basic: pd.DataFrame = df_basic.loc[(df_basic['Voucher Date'] <= end_date) & (
            df_basic['Voucher Date'] >= datetime(year=end_date.year, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['Voucher Date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()
    cy_ytd_basic_bud: pd.DataFrame = df_basic_bud.loc[(df_basic_bud['Voucher Date'] <= end_date) & (
            df_basic_bud['Voucher Date'] >= datetime(year=end_date.year, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['Voucher Date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()

    py_cp_basic: pd.DataFrame = df_basic.loc[
        df_basic['Voucher Date'] == datetime(year=end_date.year - 1, month=end_date.month, day=end_date.day)].drop(
        columns=['Voucher Date']).reset_index().rename(columns={'index': 'Description'})

    py_ytd_basic: pd.DataFrame = df_basic.loc[
        (df_basic['Voucher Date'] <= end_date - relativedelta(years=1) + relativedelta(day=31)) & (
                df_basic['Voucher Date'] >= datetime(year=end_date.year - 1, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['Voucher Date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()

    cy_ytd_basic_monthwise: pd.DataFrame = df_basic.loc[(df_basic['Voucher Date'] <= end_date) & (
            df_basic['Voucher Date'] >= datetime(year=end_date.year, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).pivot_table(index='Description', columns='Voucher Date', values='Amount',
                                                      aggfunc='sum', fill_value=0).reset_index()

    cy_cp_mid: pd.DataFrame = df_mid.loc[df_mid['Voucher Date'] == end_date].drop(
        columns=['Voucher Date']).reset_index().rename(columns={'index': 'Description'})
    cy_cp_mid_bud: pd.DataFrame = df_mid_bud.loc[df_mid['Voucher Date'] == end_date].drop(
        columns=['Voucher Date']).reset_index().rename(columns={'index': 'Description'})

    cy_pp_mid: pd.DataFrame = df_mid.loc[
        df_mid['Voucher Date'] == end_date - relativedelta(months=1) + relativedelta(day=31)].drop(
        columns=['Voucher Date']).reset_index().rename(columns={'index': 'Description'})

    cy_ytd_mid: pd.DataFrame = df_mid.loc[
        (df_mid['Voucher Date'] <= end_date) & (
                df_mid['Voucher Date'] >= datetime(year=end_date.year, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['Voucher Date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()
    cy_ytd_mid_bud: pd.DataFrame = df_mid_bud.loc[
        (df_mid_bud['Voucher Date'] <= end_date) & (
                df_mid_bud['Voucher Date'] >= datetime(year=end_date.year, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['Voucher Date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()

    py_cp_mid: pd.DataFrame = df_mid.loc[
        df_mid['Voucher Date'] == datetime(year=end_date.year - 1, month=end_date.month, day=end_date.day)].drop(
        columns=['Voucher Date']).reset_index().rename(columns={'index': 'Description'})

    py_ytd_mid: pd.DataFrame = df_mid.loc[
        (df_mid['Voucher Date'] <= end_date - relativedelta(years=1) + relativedelta(day=31)) & (
                df_mid['Voucher Date'] >= datetime(year=end_date.year - 1, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['Voucher Date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()

    cy_cp_full: pd.DataFrame = df_full.loc[df_full['Voucher Date'] == end_date].drop(
        columns=['Voucher Date']).reset_index().rename(columns={'index': 'Description'})
    cy_cp_full_bud: pd.DataFrame = df_full_bud.loc[df_full_bud['Voucher Date'] == end_date].drop(
        columns=['Voucher Date']).reset_index().rename(columns={'index': 'Description'})

    cy_pp_full: pd.DataFrame = df_full.loc[
        df_full['Voucher Date'] == end_date - relativedelta(months=1) + relativedelta(day=31)].drop(
        columns=['Voucher Date']).reset_index().rename(columns={'index': 'Description'})

    cy_ytd_full: pd.DataFrame = df_full.loc[(df_full['Voucher Date'] <= end_date) & (
            df_full['Voucher Date'] >= datetime(year=end_date.year, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['Voucher Date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()
    cy_ytd_full_bud: pd.DataFrame = df_full_bud.loc[(df_full_bud['Voucher Date'] <= end_date) & (
            df_full_bud['Voucher Date'] >= datetime(year=end_date.year, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['Voucher Date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()

    py_cp_full: pd.DataFrame = df_full.loc[
        df_full['Voucher Date'] == datetime(year=end_date.year - 1, month=end_date.month, day=end_date.day)].drop(
        columns=['Voucher Date']).reset_index().rename(columns={'index': 'Description'})

    py_ytd_full: pd.DataFrame = df_full.loc[
        (df_full['Voucher Date'] <= end_date - relativedelta(years=1) + relativedelta(day=31)) & (
                df_full['Voucher Date'] >= datetime(year=end_date.year - 1, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['Voucher Date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()

    return {'df_basic': {'cy_cp_basic': cy_cp_basic, 'cy_pp_basic': cy_pp_basic, 'cy_ytd_basic': cy_ytd_basic,
                         'py_cp_basic': py_cp_basic, 'py_ytd_basic': py_ytd_basic, 'cy_cp_basic_bud': cy_cp_basic_bud,
                         'cy_ytd_basic_bud': cy_ytd_basic_bud, 'cy_ytd_basic_monthwise': cy_ytd_basic_monthwise},

            'df_mid': {'cy_cp_mid': cy_cp_mid, 'cy_pp_mid': cy_pp_mid, 'cy_ytd_mid': cy_ytd_mid, 'py_cp_mid': py_cp_mid,
                       'py_ytd_mid': py_ytd_mid, 'cy_cp_mid_bud': cy_cp_mid_bud, 'cy_ytd_mid_bud': cy_ytd_mid_bud},

            'df_full': {'cy_cp_full': cy_cp_full, 'cy_pp_full': cy_pp_full, 'cy_ytd_full': cy_ytd_full,
                        'py_cp_full': py_cp_full, 'py_ytd_full': py_ytd_full, 'cy_cp_full_bud': cy_cp_full_bud,
                        'cy_ytd_full_bud': cy_ytd_full_bud}}


def offset_bal(end_date: datetime, data: pd.DataFrame) -> float:
    offset_accounts: list = company_info[0]['data']['offset_accounts']
    offset_filt = (data['Voucher Date'] <= end_date) & (
        data['Ledger_Code'].isin(offset_accounts))
    offset_df: float = data.loc[offset_filt, 'Amount'].sum()
    return offset_df


def interco_bal(data: pd.DataFrame, end_date: datetime) -> dict:
    interco_final: pd.DataFrame = pd.DataFrame()
    for entity in company_data:
        interco_ids: list = dCoAAdler.loc[dCoAAdler['Ledger_Name'].isin(
            company_data[entity]['names'])].index.tolist()
        interco_filt = (data['Voucher Date'] <= end_date) & (
            data['Ledger_Code'].isin(interco_ids))
        interco_df: pd.DataFrame = data.loc[interco_filt, ['Amount']]
        interco_df['Description'] = entity
        interco_df = interco_df.groupby(
            by=['Description'], as_index=False).sum()
        interco_final = pd.concat([interco_final, interco_df])
    interco_final = interco_final.loc[interco_final['Amount'] != 0]
    rpr: float = interco_final.loc[interco_final['Amount'] < 0, 'Amount'].sum()
    rpp: float = interco_final.loc[interco_final['Amount'] > 0, 'Amount'].sum()
    rpr_df: pd.DataFrame = interco_final.loc[interco_final['Amount'] < 0, [
        'Description', 'Amount']].sort_values(by='Amount', ascending=True)
    rpp_df: pd.DataFrame = interco_final.loc[interco_final['Amount'] > 0, [
        'Description', 'Amount']].sort_values(by='Amount', ascending=False)
    return {'rpr': rpr, 'rpp': rpp, 'rpr_df': rpr_df, 'rpp_df': rpp_df}


def balancesheet(data: pd.DataFrame, end_date: datetime) -> pd.DataFrame:
    offset_accounts: list = company_info[0]['data']['offset_accounts']
    # Sum total of offset_accounts is zero. i.e. PDC
    interco_acc_names: list = [i for j in [
        company_data[i]['names'] for i in company_data] for i in j]
    interco_acc_codes: list = dCoAAdler.loc[dCoAAdler['Ledger_Name'].isin(
        interco_acc_names)].index.tolist()

    exclude_bs_codes: list = offset_accounts + interco_acc_codes

    bs_filt = (data['Voucher Date'] <= end_date) & (~data['Ledger_Code'].isin(exclude_bs_codes)) & (
        data['Fourth_Level_Group_Name'].isin(['Assets', 'Liabilities', 'Equity']))
    is_filt = (data['Voucher Date'] <= end_date) & (
        data['Fourth_Level_Group_Name'].isin(['Income', 'Expenses']))

    dr_in_ap = data.loc[
        (data['Second_Level_Group_Name'] == 'Accounts Payables') & (data['Voucher Date'] <= end_date), ['Ledger_Code',
                                                                                                        'Amount']].groupby(
        by='Ledger_Code').sum()
    # returns negative figure
    dr_in_ap = dr_in_ap.loc[dr_in_ap['Amount'] < 0, 'Amount'].sum()
    cr_in_ar = data.loc[
        (data['Second_Level_Group_Name'] == 'Trade Receivables') & (data['Voucher Date'] <= end_date), ['Ledger_Code',
                                                                                                        'Amount']].groupby(
        by='Ledger_Code').sum()
    # returns positive figure
    cr_in_ar = cr_in_ar.loc[cr_in_ar['Amount'] > 0, 'Amount'].sum()

    bs_data: pd.DataFrame = data.loc[bs_filt, ['Second_Level_Group_Name', 'Amount']].groupby(
        by=['Second_Level_Group_Name'], as_index=False).sum().rename(
        columns={'Second_Level_Group_Name': 'Description'}).set_index(keys='Description')
    cum_profit: float = data.loc[is_filt, 'Amount'].sum()
    rounding_diff: float = data.loc[data['Voucher Date'] <= end_date, 'Amount'].sum()
    interco: dict = interco_bal(data=merged, end_date=end_date)
    rpr: float = interco.get('rpr')
    rpp: float = interco.get('rpp')
    rpr_row: pd.DataFrame = pd.DataFrame(data={'Amount': [rpr]}, index=[
        'Due From Related Parties'])
    rpp_row: pd.DataFrame = pd.DataFrame(
        data={'Amount': [rpp]}, index=['Due To Related Parties'])
    bs_data.loc['Accounts Payables',
    'Amount'] = bs_data.loc['Accounts Payables', 'Amount'] - dr_in_ap
    bs_data.loc['Other Receivable',
    'Amount'] = bs_data.loc['Other Receivable', 'Amount'] + dr_in_ap
    bs_data.loc['Trade Receivables',
    'Amount'] = bs_data.loc['Trade Receivables', 'Amount'] + cr_in_ar
    bs_data.loc['Accruals & Other Payables', 'Amount'] = bs_data.loc[
                                                             'Accruals & Other Payables', 'Amount'] - cr_in_ar - rounding_diff
    bs_data.loc['Retained Earnings',
    'Amount'] = bs_data.loc['Retained Earnings', 'Amount'] + cum_profit
    bs_data = pd.concat([bs_data, rpr_row, rpp_row], ignore_index=False)

    ca: float = bs_data.loc['Cash & Cash Equivalents', 'Amount'] + bs_data.loc['Inventory', 'Amount'] + bs_data.loc[
        'Other Receivable', 'Amount'] + bs_data.loc['Trade Receivables', 'Amount'] + bs_data.loc[
                    'Due From Related Parties', 'Amount']
    nca: float = (bs_data.loc['Intangible Assets', 'Amount'] if 'Intangible Assets' in bs_data.index else 0) + \
                 bs_data.loc[
                     'Property, Plant  & Equipment', 'Amount'] + \
                 bs_data.loc[
                     'Right of use Asset', 'Amount']
    equity: float = bs_data.loc['Retained Earnings', 'Amount'] + bs_data.loc['Share Capital', 'Amount'] + (bs_data.loc[
                                                                                                               'Statutory Reserves', 'Amount'] if 'Statutory Reserves' in bs_data.index else 0)
    cl: float = bs_data.loc['Accounts Payables', 'Amount'] + bs_data.loc['Accruals & Other Payables', 'Amount'] + \
                bs_data.loc[
                    'Due To Related Parties', 'Amount']
    ncl: float = bs_data.loc['Provisions', 'Amount'] + bs_data.loc[
        'Lease Liabilities', 'Amount'] if 'Lease Liabilities' in bs_data.index else 0

    ta: float = ca + nca
    tl: float = cl + ncl
    tle: float = tl + equity

    cl_row: pd.DataFrame = pd.DataFrame(
        data={'Amount': [cl]}, index=['Current Liabilities'])
    ncl_row: pd.DataFrame = pd.DataFrame(data={'Amount': [ncl]}, index=[
        'Non Current Liabilities'])
    tl_row: pd.DataFrame = pd.DataFrame(
        data={'Amount': [tl]}, index=['Liabilities'])
    ca_row: pd.DataFrame = pd.DataFrame(
        data={'Amount': [ca]}, index=['Current Assets'])
    nca_row: pd.DataFrame = pd.DataFrame(
        data={'Amount': [nca]}, index=['Non Current Assets'])
    ta_row: pd.DataFrame = pd.DataFrame(
        data={'Amount': [ta]}, index=['Assets'])
    equity_row: pd.DataFrame = pd.DataFrame(
        data={'Amount': [equity]}, index=['Equity'])
    tle_row: pd.DataFrame = pd.DataFrame(data={'Amount': [tle]}, index=[
        'Total Equity & Liabilities'])

    bs_data = pd.concat([bs_data, cl_row, ncl_row, tl_row, equity_row, ca_row, nca_row, ta_row, tle_row],
                        ignore_index=False)

    bs_data = bs_data.loc[bs_data['Amount'] != 0]
    bs_data['Amount'] = bs_data['Amount'] * -1
    return bs_data


def bsratios(bsdata: pd.DataFrame, pldata: pd.DataFrame, periods: list, end_date: datetime) -> dict:
    values: list = [np.nan] * len(bsdata.columns)
    df_ratios_bs = pd.DataFrame(data={'period': bsdata.columns.tolist(), 'cr': values, 'ato': values, 'roe': values})
    for period in periods:
        if period != datetime(year=end_date.year - 1, month=end_date.month, day=end_date.day):
            current_period: str = period.strftime('%Y-%m-%d')
            prior_year: str = int(period.strftime('%Y')) - 1 if int(period.strftime('%Y')) != 2020 else int(
                period.strftime('%Y'))
            previous_period: str = f"{prior_year}-{period.strftime('%m')}-{period.strftime('%d')}"
            # current_ratio https://corporatefinanceinstitute.com/resources/accounting/current-ratio-formula/  Liquidity ratio
            current_ratio: float = bsdata.loc[bsdata['Description'] == 'Current Assets', current_period].iloc[0] / - \
                bsdata.loc[bsdata['Description'] == 'Current Liabilities', current_period].iloc[0]
            # asset turnover ratio https://corporatefinanceinstitute.com/resources/accounting/asset-turnover-ratio/ efficiency
            asset_turnover: float = pldata.loc[pldata['Description'] == 'Total Revenue', current_period].iloc[0] / ((
                                                                                                                            bsdata.loc[
                                                                                                                                bsdata[
                                                                                                                                    'Description'] == 'Assets', current_period].iloc[
                                                                                                                                0] +
                                                                                                                            bsdata.loc[
                                                                                                                                bsdata[
                                                                                                                                    'Description'] == 'Assets', previous_period].iloc[
                                                                                                                                0]) / 2)
            # roe https://corporatefinanceinstitute.com/resources/accounting/what-is-return-on-equity-roe/ profitability
            roe: float = pldata.loc[pldata['Description'] == 'Net Profit', current_period].iloc[0] / ((-bsdata.loc[
                bsdata['Description'] == 'Equity', current_period].iloc[0] + -bsdata.loc[
                bsdata['Description'] == 'Equity', previous_period].iloc[0]) / 2) * 100
            df_ratios_bs.loc[df_ratios_bs['period'] == current_period, 'cr'] = current_ratio
            df_ratios_bs.loc[df_ratios_bs['period'] == current_period, 'ato'] = asset_turnover
            df_ratios_bs.loc[df_ratios_bs['period'] == current_period, 'roe'] = roe
    df_ratios_bs.dropna(inplace=True)
    return df_ratios_bs


def plratios(df_pl: pd.DataFrame, plcombined: pd.DataFrame) -> dict:
    plmeasures: dict = {
        'gp': {'cy_cp_basic': 0, 'cy_ytd_basic': 0, 'cy_pp_basic': 0, 'py_cp_basic': 0, 'py_ytd_basic': 0,
               'cy_cp_basic_bud': 0, 'cy_ytd_basic_bud': 0, 'cy_ytd_basic_monthwise': 0},
        'np': {'cy_cp_basic': 0, 'cy_ytd_basic': 0, 'cy_pp_basic': 0, 'py_cp_basic': 0, 'py_ytd_basic': 0,
               'cy_cp_basic_bud': 0, 'cy_ytd_basic_bud': 0, 'cy_ytd_basic_monthwise': 0},
        'ebitda': {'cy_cp_basic': 0, 'cy_ytd_basic': 0, 'cy_pp_basic': 0, 'py_cp_basic': 0, 'py_ytd_basic': 0,
                   'cy_cp_basic_bud': 0, 'cy_ytd_basic_bud': 0, 'cy_ytd_basic_monthwise': 0}}

    for measure in plmeasures.keys():
        for k, v in df_pl['df_basic'].items():
            if k == 'cy_ytd_basic_monthwise':
                financial: pd.DataFrame = v.loc[v['Description'].isin(['Gross Profit', 'Net Profit', 'Total Revenue',
                                                                       'Depreciation', 'Depreciation - Projects',
                                                                       'Interest Expenses'])]
                gp = financial.index[financial['Description'] == 'Gross Profit'][0]
                netp = financial.index[financial['Description'] == 'Net Profit'][0]
                rev = financial.index[financial['Description'] == 'Total Revenue'][0]
                dep = financial.index[financial['Description'] == 'Depreciation'][0]
                deppro = financial.index[financial['Description'] == 'Depreciation - Projects'][0]
                interest = financial.index[financial['Description'] == 'Interest Expenses'][0]
                financial = financial.transpose().reset_index().rename(columns={gp: 'Gross Profit', netp: 'Net Profit',
                                                                                rev: 'Total Revenue',
                                                                                'index': 'Description',
                                                                                dep: 'Depreciation',
                                                                                deppro: 'Depreciation - Projects',
                                                                                interest: 'Interest Expenses'}).drop(0)
                financial.loc[:, 'EBITDA'] = financial['Net Profit'] - financial['Depreciation'] - financial[
                    'Depreciation - Projects'] - financial['Interest Expenses']
                financial.drop(columns=['Depreciation', 'Depreciation - Projects', 'Interest Expenses'], inplace=True)
                plmeasures[measure][k] = financial
            else:
                df: pd.DataFrame = v.set_index('Description')

                if measure == 'gp':
                    ratio: float = df.loc['Gross Profit', 'Amount'] / df.loc['Total Revenue', 'Amount'] * 100
                if measure == 'np':
                    ratio: float = df.loc['Net Profit', 'Amount'] / df.loc['Total Revenue', 'Amount'] * 100
                if measure == 'ebitda':
                    ratio: float = (df.loc['Net Profit', 'Amount'] -
                                    df.loc['Depreciation', 'Amount'] if 'Depreciation' in df.index else 0 -
                                                                                                        df.loc[
                                                                                                            'Depreciation - Projects', 'Amount'] if 'Depreciation - Projects' in df.index else 0 -
                                                                                                                                                                                               df.loc[
                                                                                                                                                                                                   'Interest Expenses', 'Amount'] if 'Interest Expenses' in df.index else 0) / \
                                   df.loc['Total Revenue', 'Amount'] * 100
                plmeasures[measure][k] = ratio

    plcombined.fillna(0, inplace=True)
    plcombined.set_index('Description', inplace=True)
    values: list = [np.nan] * len(plcombined.columns)
    df_ratios = pd.DataFrame(
        data={'period': plcombined.columns.tolist(), 'gp': values, 'np': values, 'ebitda': values, 'revenue': values})

    for period in df_ratios['period']:
        revenue: float = plcombined.loc['Total Revenue', period]
        gp: float = plcombined.loc['Gross Profit', period]
        netp: float = plcombined.loc['Net Profit', period]
        interest: float = plcombined.loc['Interest Expenses', period]
        dep: float = plcombined.loc['Depreciation', period]
        depro: float = plcombined.loc['Depreciation - Projects', period]
        ebitda: float = netp + dep + depro + interest
        df_ratios.loc[df_ratios['period'] == period, 'gp'] = gp
        df_ratios.loc[df_ratios['period'] == period, 'np'] = netp
        df_ratios.loc[df_ratios['period'] == period, 'ebitda'] = ebitda
        df_ratios.loc[df_ratios['period'] == period, 'revenue'] = revenue
        plmeasures['plyearly'] = df_ratios
    return plmeasures


def settlement_days(invoices: list) -> int:
    col_days: list = []
    invoices = [inv for inv in invoices if not "CN" in inv]
    for invoice in invoices:
        inv_value: float = fCollection.loc[(fCollection['invoice_number'] == invoice), 'invoice_amount'].iloc[0]
        total_collection: float = fCollection.loc[(fCollection['invoice_number'] == invoice) & (
                fCollection['voucher_date'] <= end_date), 'voucher_amount'].sum()
        if (inv_value - total_collection) == 0:
            last_date: datetime = fCollection.loc[(fCollection['invoice_number'] == invoice) & (
                    fCollection['voucher_date'] <= end_date), 'voucher_date'].max()
            inv_date: datetime = fCollection.loc[(fCollection['invoice_number'] == invoice), 'invoice_date'].iloc[0]
            col_days.append(last_date - inv_date)

    return statistics.median(col_days) if col_days else timedelta(days=0)


def cust_ageing(customer: str) -> pd.DataFrame:
    ledgers: list = dCustomers.loc[(dCustomers['Cus_Name'] == customer), 'Ledger_Code'].tolist()
    credit_days: int = int(dCustomers.loc[dCustomers['Cus_Name'].isin([customer]), 'Credit_Days'].iloc[0])
    invoices: np.ndarray = fCollection.loc[fCollection['ledger_code'].isin(ledgers), 'invoice_number'].unique()
    cust_soa: pd.DataFrame = fCollection.loc[(fCollection['invoice_number'].isin(invoices)), ['invoice_date',
                                                                                              'invoice_amount',
                                                                                              'voucher_amount',
                                                                                              'invoice_number',
                                                                                              'voucher_date']]
    inv_value_list: list = []
    age_bracket_list: list = []
    ranges = [(0, 'Not Due'), (30, '1-30'), (60, '31-60'),
              (90, '61-90'), (120, '91-120'), (121, '121-150'),
              (151, '151-180'), (181, '181-210'), (211, '211-240'),
              (241, '241-270'), (271, '271-300'), (300, '301-330'),
              (331, '331-360'), (float('inf'), 'More than 361')]
    for invoice in invoices:
        total_collection: float = cust_soa.loc[
            (cust_soa['invoice_number'] == invoice) & (cust_soa['voucher_date'] <= end_date), 'voucher_amount'].sum()
        inv_value: float = cust_soa.loc[(cust_soa['invoice_number'] == invoice), 'invoice_amount'].iloc[0]
        if (inv_value - total_collection) != 0:
            inv_value_list.append(inv_value - total_collection)
            days_passed: int = (end_date - cust_soa.loc[(cust_soa['invoice_number'] == invoice), 'invoice_date'].iloc[
                0] - timedelta(days=credit_days)).days
            for threshold, label in ranges:
                if days_passed <= threshold:
                    age_bracket_list.append(label)
                    break
    outstanding_df: pd.DataFrame = pd.DataFrame(
        data={'Amount': inv_value_list, 'Age Bracket': age_bracket_list}).groupby(by='Age Bracket').sum()
    if not outstanding_df.empty:
        outstanding_df.reset_index(inplace=True)
        outstanding_df['Age Bracket'] = pd.Categorical(outstanding_df['Age Bracket'], categories=[i[1] for i in ranges],
                                                       ordered=True)
        outstanding_df.sort_values(by='Age Bracket', inplace=True)
        outstanding_df.set_index(keys='Age Bracket', drop=True, inplace=True)
    else:
        outstanding_df
    return outstanding_df


def customer_ratios(customers: list, fInvoices: pd.DataFrame, end_date: datetime, fCollection: pd.DataFrame,
                    dCustomer: pd.DataFrame, dEmployee: pd.DataFrame) -> dict:
    customer_info: dict = {}
    for customer in customers:
        customer_since: datetime = fInvoices.loc[
            (fInvoices['Cus_Name'] == customer), 'Invoice_Date'].min() if not pd.isna(
            fInvoices.loc[(fInvoices['Cus_Name'] == customer), 'Invoice_Date'].min()) else "Not Applicable"
        total_revenue: float = fInvoices.loc[
            (fInvoices['Cus_Name'] == customer) & (fInvoices['Invoice_Date'] <= end_date), 'Net_Amount'].sum()
        cust_invoices: list = fInvoices.loc[(fInvoices['Cus_Name'] == customer), 'Invoice_Number'].to_list()
        last_receipt_dt: datetime = fCollection.loc[
            fCollection['invoice_number'].isin(cust_invoices), 'voucher_date'].max() if not pd.isna(fCollection.loc[
                                                                                                        fCollection[
                                                                                                            'invoice_number'].isin(
                                                                                                            cust_invoices), 'voucher_date'].max()) else "Not Collected"
        last_receipt_number: str = "Not Collected" if last_receipt_dt == "Not Collected" else \
            fCollection.loc[(fCollection['invoice_number'].isin(cust_invoices)) & (
                    fCollection['voucher_date'] == last_receipt_dt), 'voucher_number'].tail(1).iloc[0]
        last_receipt_amt: float = "Not Collected" if last_receipt_dt == "Not Collected" else fCollection.loc[
            (fCollection['voucher_number'] == last_receipt_number), 'voucher_amount'].sum()
        cy_cp_rev: float = fInvoices.loc[
            (fInvoices['Cus_Name'] == customer) & (fInvoices['Invoice_Date'] <= end_date) & (
                    fInvoices['Invoice_Date'] >= datetime(year=end_date.year, month=end_date.month,
                                                          day=1)), 'Net_Amount'].sum()
        cy_pp_rev: float = fInvoices.loc[(fInvoices['Cus_Name'] == customer) & (
                fInvoices['Invoice_Date'] <= end_date.replace(day=1) - timedelta(days=1)) & (
                                                 fInvoices['Invoice_Date'] >= end_date + relativedelta(day=1,
                                                                                                       months=-1)), 'Net_Amount'].sum()
        cy_ytd_rev: float = fInvoices.loc[
            (fInvoices['Cus_Name'] == customer) & (fInvoices['Invoice_Date'] <= end_date) & (
                    fInvoices['Invoice_Date'] >= datetime(year=end_date.year, month=1, day=1)), 'Net_Amount'].sum()
        py_ytd_rev: float = fInvoices.loc[(fInvoices['Cus_Name'] == customer) & (
                fInvoices['Invoice_Date'] <= datetime(year=end_date.year - 1, month=end_date.month,
                                                      day=end_date.day)) & (
                                                  fInvoices['Invoice_Date'] >= datetime(year=end_date.year - 1, month=1,
                                                                                        day=1)), 'Net_Amount'].sum()
        py_cp_rev: float = fInvoices.loc[(fInvoices['Cus_Name'] == customer) & (
                fInvoices['Invoice_Date'] <= datetime(year=end_date.year - 1, month=end_date.month,
                                                      day=end_date.day)) & (
                                                 fInvoices['Invoice_Date'] >= datetime(year=end_date.year - 1,
                                                                                       month=end_date.month,
                                                                                       day=1)), 'Net_Amount'].sum()
        collection_median: float = "Not Collected" if last_receipt_dt == "Not Collected" else settlement_days(
            invoices=cust_invoices)
        credit_days: int = dCustomer.loc[dCustomers['Cus_Name'].isin([customer]), 'Credit_Days'].iloc[0]
        date_established: datetime = dCustomer.loc[dCustomer['Cus_Name'].isin([customer]), 'Date_Established'].iloc[0]
        outstanding_bal: float = fGL.loc[
            (fGL['Ledger_Code'].isin(dCustomer.loc[dCustomer['Cus_Name'].isin([customer]), 'Ledger_Code'].tolist())) & (
                    fGL['Voucher Date'] <= end_date), 'Amount'].sum()
        cy_cp_rev_contrib_pct: float = cy_cp_rev / fInvoices.loc[(fInvoices['Invoice_Date'] <= end_date) & (
                fInvoices['Invoice_Date'] >= datetime(year=end_date.year, month=end_date.month,
                                                      day=1)), 'Net_Amount'].sum() * 100
        cy_ytd_rev_contrib_pct: float = cy_ytd_rev / fInvoices.loc[(fInvoices['Invoice_Date'] <= end_date) & (
                fInvoices['Invoice_Date'] >= datetime(year=end_date.year, month=1,
                                                      day=1)), 'Net_Amount'].sum() * 100
        monthly_rev: pd.DataFrame = fInvoices.loc[
            (fInvoices['Cus_Name'] == customer) & (fInvoices['Invoice_Date'] <= end_date) & (
                    fInvoices['Invoice_Date'] >= datetime(year=end_date.year, month=1, day=1)), ['Invoice_Date',
                                                                                                 'Net_Amount']].groupby(
            by=['Invoice_Date']).sum()
        monthly_rev.reset_index(inplace=True)
        monthly_rev.rename(columns={'Invoice_Date': 'Month', 'Net_Amount': 'Amount'}, inplace=True)
        monthly_rev.set_index(keys='Month', drop=True, inplace=True)
        ageing: pd.DataFrame = cust_ageing(customer=customer)
        last_sales_person: str = fInvoices.loc[
            (fInvoices['Invoice_Date'] <= end_date) & (fInvoices['Cus_Name'] == customer), 'Employee_Code'].tail(
            1).iloc[0]
        last_sales_person = dEmployee.loc[last_sales_person, 'Employee_Name']

        stats: dict = {
            'customer_since': "Not Applicable" if customer_since == "Not Applicable" else customer_since.strftime(
                '%d-%m-%Y'), 'total_revenue': total_revenue, 'credit_score': 0,
            'last_receipt_amt': "Not Collected" if last_receipt_dt == "Not Collected" else last_receipt_amt,
            'cy_cp_rev': cy_cp_rev, 'cy_pp_rev': cy_pp_rev,
            'last_receipt_dt': "Not Collected" if last_receipt_dt == "Not Collected" else last_receipt_dt.strftime(
                '%d-%m-%Y'),
            'cy_ytd_rev': cy_ytd_rev, 'py_cp_rev': py_cp_rev, 'py_ytd_rev': py_ytd_rev,
            'collection_median': "Not Collected" if last_receipt_dt == "Not Collected" else collection_median.days,
            'credit_days': credit_days, 'last_sales_person': last_sales_person,
            'customer_gp_cp': 0, 'outstanding_bal': -outstanding_bal, 'ageing': ageing,
            'date_established': date_established.strftime('%d-%m-%Y'),
            'cy_cp_rev_contrib_pct': f'{round(cy_cp_rev_contrib_pct, 1)}%',
            'cy_ytd_rev_contrib_pct': f'{round(cy_ytd_rev_contrib_pct, 1)}%',
            'cy_cp_roi': 0, 'customer_gp_ytd': 0,
            'cy_ytd_roi': 0, 'monthly_rev': monthly_rev, 'remarks': 0}
        customer_info[customer] = stats
    return customer_info


def organic_sales(emp_id: str, mode: str) -> float:
    if mode.lower() == 'month':
        start_date: datetime = datetime(year=end_date.year, month=end_date.month, day=1)
    elif mode.lower() == 'ytd':
        start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    else:
        raise ValueError(f'Invalid mode{mode}')

    customers: list = list(fInvoices.loc[
                               (fInvoices['Invoice_Date'] <= end_date) & (fInvoices['Invoice_Date'] >= start_date) & (
                                       fInvoices['Employee_Code'] == emp_id), 'Customer_Code'].unique())
    for customer in customers:
        first_sales_person: str = fInvoices.loc[(fInvoices['Customer_Code'] == customer), 'Employee_Code'].tolist()[0]
        if first_sales_person != emp_id:
            customers.remove(customer)
    self_sales: float = fInvoices.loc[
        (fInvoices['Customer_Code'].isin(customers)) & (fInvoices['Invoice_Date'] <= end_date) & (
                fInvoices['Invoice_Date'] >= start_date), 'Net_Amount'].sum()
    return self_sales


def sales_person(emp_ids: list, dEmployee: pd.DataFrame, fInvoices: pd.DataFrame) -> dict:
    salesperson_stats: dict = {}
    for emp_id in emp_ids:
        doj: datetime = dEmployee.loc[emp_id, 'doj']
        cy_cp_rev: float = fInvoices.loc[(fInvoices['Invoice_Date'] <= end_date) & (
                fInvoices['Invoice_Date'] >= datetime(year=end_date.year, month=end_date.month, day=1)) & (
                                                 fInvoices['Employee_Code'] == emp_id), 'Net_Amount'].sum()
        cy_ytd_rev: float = fInvoices.loc[(fInvoices['Invoice_Date'] <= end_date) & (
                fInvoices['Invoice_Date'] >= datetime(year=end_date.year, month=1, day=1)) & (
                                                  fInvoices['Employee_Code'] == emp_id), 'Net_Amount'].sum()
        cy_cp_rev_org: float = organic_sales(emp_id=emp_id, mode='month')
        cy_ytd_rev_org: float = organic_sales(emp_id=emp_id, mode='ytd')
        cy_cp_customers: list = list(set(fInvoices.loc[(fInvoices['Employee_Code'] == emp_id) & (
                fInvoices['Invoice_Date'] >= datetime(year=end_date.year, month=end_date.month, day=1)) & (fInvoices[
                                                                                                               'Invoice_Date'] <= end_date), 'Customer_Code'].tolist()))
        customers_till: list = list(set(fInvoices.loc[(
                fInvoices['Invoice_Date'] <= end_date + relativedelta(day=31, months=-1)), 'Customer_Code'].tolist()))
        new_customers_added: int = len([customer for customer in cy_cp_customers if customer not in customers_till])

        ar_balance: pd.DataFrame = fGL.loc[
            (fGL['Ledger_Code'].isin(
                fInvoices.loc[fInvoices['Customer_Code'].isin(cy_cp_customers), 'Ledger_Code'].unique())) & (
                    fGL['Voucher Date'] <= end_date), ['Ledger_Code', 'Amount']].groupby(by='Ledger_Code').sum()
        monthly_rev: pd.DataFrame = fInvoices.loc[
            (fInvoices['Invoice_Date'] <= end_date) & (fInvoices['Employee_Code'] == emp_id) & (
                    fInvoices['Invoice_Date'] >= datetime(year=end_date.year, month=1, day=1)), ['Invoice_Date',
                                                                                                 'Net_Amount']].groupby(
            by='Invoice_Date').sum()
        cy_cp_rev_total: float = fInvoices.loc[(fInvoices['Invoice_Date'] <= end_date) & (
                fInvoices['Invoice_Date'] >= datetime(year=end_date.year, month=end_date.month,
                                                      day=1)), 'Net_Amount'].sum()
        cy_ytd_rev_total: float = fInvoices.loc[(fInvoices['Invoice_Date'] <= end_date) & (
                fInvoices['Invoice_Date'] >= datetime(year=end_date.year, month=1, day=1)), 'Net_Amount'].sum()
        cy_cp_rev_contrib_pct: float = cy_cp_rev / cy_cp_rev_total * 100
        cy_ytd_rev_contrib_pct: float = cy_ytd_rev / cy_ytd_rev_total * 100
        stats: dict = {'doj': doj.strftime('%d-%m-%Y'), 'cp_target': 0, 'cy_cp_rev': cy_cp_rev, 'ytd_target': 0,
                       'cy_ytd_rev': cy_ytd_rev, 'cy_cp_rev_org': cy_cp_rev_org,
                       'cy_ytd_rev_org': cy_ytd_rev_org,
                       'new_customers_added': new_customers_added, 'cy_cp_gp': 0, 'cy_ytd_gp': 0,
                       'ar_balance': ar_balance, 'monthly_rev': monthly_rev,
                       'cy_cp_rev_contrib_pct': f'{round(cy_cp_rev_contrib_pct, 1)}%',
                       'cy_ytd_rev_contrib_pct': f'{round(cy_ytd_rev_contrib_pct, 1)}%'}
        salesperson_stats[emp_id] = stats

    return salesperson_stats


def revenue(end_date: datetime, data: pd.DataFrame, fInvoices: pd.DataFrame) -> dict:
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    first_invoice_dates: pd.Series = fInvoices.groupby('Customer_Code')['Invoice_Date'].min()
    fInvoices: pd.DataFrame = fInvoices.loc[
        (fInvoices['Invoice_Date'] >= start_date) & (fInvoices['Invoice_Date'] <= end_date), ['Invoice_Number',
                                                                                              'Customer_Code',
                                                                                              'Net_Amount',
                                                                                              'Employee_Code',
                                                                                              'Invoice_Date', 'Type']]
    fInvoices.loc[:, 'new_or_old'] = fInvoices.apply(
        lambda row: 'Existing' if row['Invoice_Date'] > first_invoice_dates[row['Customer_Code']] else 'New', axis=1)
    new_or_old: pd.DataFrame = fInvoices.groupby(by=['Invoice_Date', 'new_or_old'], as_index=False)['Net_Amount'].sum()
    inv_emp: pd.DataFrame = fInvoices.groupby(by=['Invoice_Date', 'Employee_Code'], as_index=False)['Net_Amount'].sum()

    rev_filt = (data['Third_Level_Group_Name'] == 'Direct Income') & (
            data['Voucher Date'] <= end_date)
    rev_division: pd.DataFrame = data.loc[rev_filt, ['Voucher Date', 'Amount', 'Second_Level_Group_Name']].groupby(
        by=['Voucher Date', 'Second_Level_Group_Name'], as_index=False).sum()
    sales_invoices: np.ndarray = data.loc[rev_filt, 'Voucher Number'].unique()
    total_invoices: np.ndarray = fInvoices['Invoice_Number'].unique()
    worked_invoices: list = [
        inv for inv in sales_invoices if inv in total_invoices]
    rev_category: pd.DataFrame = data.loc[
        (data['Voucher Number'].isin(worked_invoices)) & (data['Third_Level_Group_Name'] == 'Direct Income'), [
            'Voucher Number', 'Amount', 'Voucher Date']].rename(
        columns={'Voucher Number': 'Invoice_Number'})
    rev_category: pd.DataFrame = pd.merge(left=rev_category, right=fInvoices[['Invoice_Number', 'Type']],
                                          on='Invoice_Number', how='left').drop(columns=['Invoice_Number']).groupby(
        by=['Voucher Date', 'Type'], as_index=False).sum()
    return {'rev_division': rev_division, 'rev_category': rev_category, 'new_or_old': new_or_old, 'inv_emp': inv_emp}


def closing_date(row, dCustomers: pd.DataFrame) -> datetime:
    """Add credit period (in days) to the voucher date and convert that date to end of the month

    Args:
        row (_type_): a row in dataframe

    Returns:
        datetime: last date of the month to which voucher becomes due
    """
    ledger_code: int = row['Ledger_Code']

    if ledger_code in dCustomers['Ledger_Code'].tolist():
        credit_days: int = int(dCustomers.loc[dCustomers['Ledger_Code'] == ledger_code, 'Credit_Days'].iloc[0])
        due_date = row['Voucher Date'] + timedelta(days=credit_days)
        return due_date + relativedelta(day=31)
    else:
        pass


def already_collected(row, fGL: pd.DataFrame, fCollection: pd.DataFrame) -> float:
    """Target collection for a given period is calculated by adding the credit period given to each customer.
    Invoices to which Target collection for a given period comprises may contain invoices which has been
    already collected prior they become due or before the beginning of target collection period. i.e. Invoice raised
    in 31/05/2024 which has 60 days credit period will become target collection in the period of 31/07/2024. But if
    such invoice has been collected on 15/06/2024, it should no longer be considered as Target collection for the period
    31/07/2024.

    Args:
        row (_type_): A row in the dataframe

    Returns:
        float: Amount already collected out of target collection
    """

    fGL = fGL.loc[(fGL['Transaction Type'].isin(VOUCHER_TYPES)) & (fGL['Ledger_Code'] >= 1000000000) & (
            fGL['Ledger_Code'] <= 1999999999)]
    fGL['Due Date'] = fGL.apply(closing_date, axis=1, args=[dCustomers])
    start_date: datetime = row['Due Date'].replace(day=1)
    due_inv_list: list = fGL.loc[
        (fGL['Due Date'] >= start_date) & (fGL['Due Date'] <= row['Due Date']), 'Voucher Number'].unique()
    collected_filt = (fCollection['invoice_number'].isin(due_inv_list)) & (fCollection['voucher_date'] < start_date)
    amount: float = fCollection.loc[collected_filt, 'voucher_amount'].sum()
    return amount


def collection(fCollection: pd.DataFrame, end_date: datetime, fGL: pd.DataFrame,
               dCustomers: pd.DataFrame) -> pd.DataFrame:
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    # filters the collection date based on the selection
    fCollection1 = fCollection.loc[
        (fCollection['voucher_date'] >= start_date) & (fCollection['voucher_date'] <= end_date)]
    # convert collection date to last date of the month, so it can be grouped to know total collected per period.
    fCollection1 = fCollection1.groupby(pd.Grouper(key='voucher_date', freq='ME'))[
        'voucher_amount'].sum().reset_index().rename(columns={'voucher_date': 'Due Date', 'voucher_amount': 'Actual'})
    fCollection1 = fCollection1.loc[(fCollection1['Due Date'] >= start_date) & (fCollection1['Due Date'] <= end_date)]

    # Reasons for Finance / Receipt total for a period not match with 'Actual' in this report
    # 1. Credit notes are part of 'Actual' in this report
    # 2. Receipts other than from customers i.e. Employee Receivable is not part of this report
    # 3. Receipts that were not allocated to invoices are not part of this report.
    # for 3 above check fCollection/Invoice Number Contains RV/CN and Payment Date ->Blank
    fGL1 = fGL.copy()
    fGL1 = fGL1.loc[(fGL1['Transaction Type'].isin(VOUCHER_TYPES)) & (fGL1['Ledger_Code'] >= 1000000000) & (
            fGL1['Ledger_Code'] <= 1999999999)]
    fGL1.loc[:, 'Amount'] = fGL1['Amount'] * -1
    fGL1.loc[:, 'Due Date'] = fGL1.apply(closing_date, axis=1, args=[dCustomers])
    fGL1 = fGL1.loc[(fGL1['Due Date'] >= start_date) & (fGL1['Due Date'] <= end_date)]
    fGL1 = fGL1.groupby(by=['Due Date'], as_index=False)['Amount'].sum()
    fGL1.loc[:, 'Already_Collected'] = fGL1.apply(already_collected, axis=1, args=[fGL, fCollection])
    fGL1['Amount'] = fGL1['Amount'] - fGL1['Already_Collected']

    fGL1.drop(columns=['Already_Collected'], inplace=True)
    fGL1.rename(columns={'Amount': 'Target'}, inplace=True)

    combined: pd.DataFrame = pd.concat([fGL1.set_index('Due Date'), fCollection1.set_index('Due Date')], axis=1,
                                       join='outer').reset_index()
    return combined


def topcustomers(fInvoices: pd.DataFrame, end_date: datetime, mode: str, div: str, type: str, cnt: int) -> pd.DataFrame:
    if mode.lower() == 'month':
        start_date: datetime = datetime(
            year=end_date.year, month=end_date.month, day=1)
    elif mode.lower() == 'ytd':
        start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    else:
        raise ValueError(f'Invalid mode{mode}')

    if div.lower() == 'guarding':
        pattern = 'CTR'
    elif div.lower() == 'elv':
        pattern = 'ORD|CRD'
    else:
        raise ValueError(f'Invalid div{div}')
    topfivecustomers: pd.DataFrame = fInvoices.loc[
        (fInvoices['Invoice_Date'] >= start_date) & (fInvoices['Invoice_Date'] <= end_date) & (
                fInvoices['Type'] == type) & (
            fInvoices['Order_ID'].str.contains(pat=pattern)), [
            'Net_Amount', 'Cus_Name']].groupby('Cus_Name').sum().sort_values(by='Net_Amount', ascending=False).head(
        cnt).reset_index().rename(columns={'Cus_Name': 'Customer', 'Net_Amount': 'Amount'})
    total_row: pd.DataFrame = pd.DataFrame(data={'Customer': ['Total'], 'Amount': [topfivecustomers['Amount'].sum()]})
    if not topfivecustomers.empty:
        topfivecustomers = pd.concat([topfivecustomers, total_row], ignore_index=True)
    return topfivecustomers


def revenue_change(fInvoices: pd.DataFrame, end_date: datetime, mode: str, order: bool) -> pd.DataFrame:
    start_date: datetime = datetime(year=end_date.year, month=end_date.month, day=1)
    pp_end: datetime = start_date - timedelta(days=1)
    pp_start: datetime = datetime(year=pp_end.year, month=pp_end.month, day=1)
    py_cp_start: datetime = datetime(year=end_date.year - 1, month=end_date.month, day=1)
    py_cp_end: datetime = py_cp_start + relativedelta(day=31)

    cy_cp: pd.Dataframe = fInvoices.loc[
        (fInvoices['Invoice_Date'] >= start_date) & (fInvoices['Invoice_Date'] <= end_date), ['Net_Amount',
                                                                                              'Cus_Name']].groupby(
        by='Cus_Name').sum().rename(columns={'Net_Amount': 'cycp'})
    cy_pp: pd.DataFrame = fInvoices.loc[
        (fInvoices['Invoice_Date'] >= pp_start) & (fInvoices['Invoice_Date'] <= pp_end), ['Net_Amount',
                                                                                          'Cus_Name']].groupby(
        by='Cus_Name').sum().rename(columns={'Net_Amount': 'cypp'})
    py_cp: pd.DataFrame = fInvoices.loc[
        (fInvoices['Invoice_Date'] >= py_cp_start) & (fInvoices['Invoice_Date'] <= py_cp_end), ['Net_Amount',
                                                                                                'Cus_Name']].groupby(
        by='Cus_Name').sum().rename(columns={'Net_Amount': 'pycp'})

    revenue_period: pd.DataFrame = pd.concat([cy_cp, cy_pp, py_cp], axis=1).fillna(0).reset_index().rename(
        columns={'Cus_Name': 'Customer'})

    revenue_period['Variance'] = revenue_period.apply(lambda x: x['cycp'] - x[f'{mode}'], axis=1)
    revenue_period.sort_values(by='Variance', ascending=order, inplace=True)
    revenue_period.drop(columns=['cycp', 'cypp', 'pycp'], inplace=True)
    total_row: pd.DataFrame = pd.DataFrame(
        data={'Customer': ['Total'], 'Variance': [revenue_period.head(5)['Variance'].sum()]})
    revenue_period = pd.concat([revenue_period.head(5), total_row], ignore_index=True)
    if order:
        revenue_period['Variance'] = revenue_period.apply(lambda x: x['Variance'] * -1, axis=1)
    return revenue_period


def number_format(num):
    if num == 0:
        return "-"
    elif num >= 0:
        return f'{num:,.0f}'
    else:
        return f'({abs(num):,.0f})'


def plotting_period(end_date: datetime, months: int) -> datetime:
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    delta = relativedelta(dt1=end_date, dt2=start_date)
    period: int = delta.months + (delta.years * 12) + 1
    if period < 6:
        start_date = end_date - relativedelta(months=months - 1)
        start_date = datetime(year=start_date.year, month=start_date.month, day=1)
    else:
        start_date
    return start_date


def change_orientation(doc, method):
    current_section = doc.sections[-1]
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    if method == 'l':  # simple letter "L"
        new_width, new_height = current_section.page_height, current_section.page_width
        new_section.orientation = WD_ORIENT.LANDSCAPE
    else:
        new_height, new_width, = current_section.page_width, current_section.page_height
        new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width = new_width
    new_section.page_height = new_height

    return new_section


def service_period(doj: datetime, end_date: datetime) -> str:
    days_passed: int = (end_date - doj).days
    service_ranges: list = [(365, '< One Year'), (730, '1-2 Years'), (1095, '2-3 Years'),
                            (1460, '3-4 Years'), (float('inf'), '4 Years +')]
    for threshold, label in service_ranges:
        if days_passed <= threshold:
            return label
            break


def emp_age(dob: datetime, end_date: datetime) -> str:
    age: int = end_date.year - dob.year - ((end_date.month, end_date.day) < (dob.month, dob.day))
    service_ranges: list = [(25, '< 25'), (35, '26-35 Years'), (45, '36-45 Years'), (float('inf'), '46 +')]
    for threshold, label in service_ranges:
        if age <= threshold:
            return label
            break


def employee_related(data: pd.DataFrame) -> dict:
    data.reset_index(inplace=True)
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    total_pie_slices: int = 5
    dEmployee: pd.DataFrame = data
    dEmployee['termination_date'] = pd.to_datetime(dEmployee['termination_date'])
    dEmployee['dob'] = pd.to_datetime(dEmployee['dob'])
    dEmployee['doj'] = pd.to_datetime(dEmployee['doj'])
    dEmployee = dEmployee.loc[(~dEmployee['Employee_Code'].isin(['ESS0015-OLD', 'ESS0016'])) & (
        dEmployee['Employee_Code'].str.contains('ESS')) & (dEmployee['doj'] <= end_date) & (
                                      (dEmployee['termination_date'] >= start_date) | (
                                  dEmployee['termination_date'].isna()))]
    emp_types: dict = {'MGMT': 'Staff', 'STAFF': 'Staff', 'ELV STAFF': 'Staff', 'LABOUR': 'Labour',
                       'LABOUR A': 'Labour',
                       'LABOUR A 2': 'Labour', 'LABOUR A 3': 'Labour', 'LABOUR A 4': 'Labour', 'ELV LABOUR': 'Labour'}
    current_emp: pd.DataFrame = dEmployee.loc[
        (dEmployee['termination_date'] > end_date) | (dEmployee['termination_date'].isna())]
    gender: dict = current_emp.value_counts(subset='Gender').to_dict()

    type: list = [emp_types[i] for i in current_emp['emp_type'].tolist()]
    type: dict = {item: type.count(item) for item in set(type)}

    dept: list = [i if i == 'ELV' else 'Guarding' for i in current_emp['Dept']]
    dept: dict = {item: dept.count(item) for item in set(dept)}

    designation: list = [i for i in current_emp['Designation'].tolist()]
    designation: dict = {item.title(): designation.count(item) for item in set(designation)}
    designation = dict(sorted(designation.items(), key=lambda item: item[1], reverse=True))
    d_sliced: dict = dict(islice(designation.items(), total_pie_slices - 1))
    d_sliced['Others'] = sum(
        [i[1] for i in dict(islice(designation.items(), total_pie_slices - 1, len(designation))).items()])

    nationality: list = [i for i in current_emp['nationality'].tolist()]
    nationality: dict = {item.title(): nationality.count(item) for item in set(nationality)}
    nationality = dict(sorted(nationality.items(), key=lambda item: item[1], reverse=True))
    n_sliced: dict = dict(islice(nationality.items(), total_pie_slices - 1))
    n_sliced['Others'] = sum(
        [i[1] for i in dict(islice(nationality.items(), total_pie_slices - 1, len(nationality))).items()])

    opening_emp = len(dEmployee.loc[(dEmployee['doj'] <= start_date - timedelta(days=1))])
    joined_emp: int = len(dEmployee.loc[(dEmployee['doj'] >= start_date)])
    resigned_emp: int = -len(dEmployee.loc[(dEmployee['termination_date'] <= end_date)])
    closing_emp: int = opening_emp + joined_emp + resigned_emp
    emp_movement: dict = {f'Employees at {start_date.date()}': opening_emp,
                          'New Joiners': joined_emp, 'Resigned Employees': resigned_emp,
                          f'Employees at {end_date.date()}': closing_emp}
    current_emp.loc[:, 'Service'] = current_emp.loc[:, 'doj'].apply(func=service_period, args=[end_date])
    current_emp.loc[:, 'Age'] = current_emp.loc[:, 'dob'].apply(func=emp_age, args=[end_date])

    service: list = current_emp['Service'].tolist()
    service: dict = {item: service.count(item) for item in set(service)}

    age: list = current_emp['Age'].tolist()
    age: dict = {item: age.count(item) for item in set(age)}

    df_new_joiner: pd.DataFrame = dEmployee.loc[dEmployee['doj'] >= start_date, ['doj', 'Employee_Code']].rename(
        columns={'doj': 'Period', 'Employee_Code': 'Joined'})
    new_joiners = df_new_joiner.groupby(pd.Grouper(key='Period', freq='ME')).count()

    df_resigned: pd.DataFrame = dEmployee.loc[
        dEmployee['termination_date'] <= end_date, ['Employee_Code', 'termination_date']].rename(
        columns={'termination_date': 'Period', 'Employee_Code': 'Resigned'})
    emp_resigned = df_resigned.groupby(pd.Grouper(key='Period', freq='ME')).count()

    total_employees: pd.DataFrame = pd.concat([new_joiners, emp_resigned], axis=1)
    total_employees['Total Employees'] = (total_employees['Joined'] - total_employees['Resigned'])
    total_employees.drop(columns=['Joined', 'Resigned'], inplace=True)
    total_employees['Total Employees'] = total_employees['Total Employees'].cumsum() + opening_emp

    employee_data: dict = {'Gender': gender, 'Type': type, 'Department': dept, 'Nationality': n_sliced,
                           'Employee Age': age, 'Service Period': service, 'Designation': d_sliced,
                           'Employee Movement': emp_movement,
                           'new_joiner': new_joiners, 'resigned_emp': emp_resigned, 'total_employees': total_employees}
    return employee_data


def operations(ftimesheet: pd.DataFrame, financial: pd.DataFrame, end_date: datetime) -> dict:
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)

    ftimesheet: pd.DataFrame = ftimesheet.loc[
        (ftimesheet['v_date'] >= start_date) & (ftimesheet['v_date'] <= end_date) & (
            ~ftimesheet['job_id'].isin(['discharged', 'not_joined']))]
    ftimesheet.loc[:, 'v_date'] = ftimesheet.apply(lambda x: x['v_date'] + relativedelta(day=31), axis=1)
    df_transport: pd.DataFrame = ftimesheet.copy()
    df_accommodation: pd.DataFrame = ftimesheet.copy()
    df_unproductive: pd.DataFrame = ftimesheet.copy()
    df_transport = df_transport.loc[
        ~df_transport['job_id'].isin(['AC-ACCOMODATION', 'Annual Leave', 'OF-Off', 'PS-PATROLING SUPERVISOR',
                                      'Paternity Leave', 'SB-STANDBY', 'Sick Leave - FP', 'UL-Unpaid Leave',
                                      'Unpaid Leave'])]
    df_unproductive = df_unproductive.loc[df_unproductive['job_id'].isin(
        ['AC-ACCOMODATION', 'Annual Leave', 'CI-CLIENT INTERVIEW', 'FP-FINGER PRINT', 'HO-HEAD OFFICE', 'ME-MOI Exam',
         'MM-MOI MEDICAL', 'MT-MOI Training',
         'OF-Off', 'OJ-ON JOB TRAINING', 'Paternity Leave', 'QM-QID MEDICAL', 'SB-STANDBY', 'Sick Leave - FP',
         'TN-TRAINING', 'UL-Unpaid Leave',
         'Unpaid Leave', 'WK-Worked'])]

    df_transport = df_transport.groupby(by=['v_date'])['cost_center'].count().reset_index().rename(
        columns={'cost_center': 'trpt_md'})
    df_accommodation = df_accommodation.groupby(by=['v_date'])['cost_center'].count().reset_index().rename(
        columns={'cost_center': 'acco_md'})
    df_unproductive = df_unproductive.groupby(by=['v_date'])['cost_center'].count().reset_index().rename(
        columns={'cost_center': 'unproductive_md'})
    financial: pd.DataFrame = financial.loc[
        financial['Description'].isin(['Transportation - Manpower', 'Accommodation - Manpower'])]
    trpt = financial.index[financial['Description'] == 'Transportation - Manpower'][0]
    acc = financial.index[financial['Description'] == 'Accommodation - Manpower'][0]
    financial = financial.transpose().reset_index().rename(columns={trpt: 'Transport', acc: 'Accommodation'}).drop(0)
    financial.loc[:, 'Transport'] = financial['Transport'] * -1
    financial.loc[:, 'Accommodation'] = financial['Accommodation'] * -1
    financial['Voucher Date'] = pd.to_datetime(financial['Voucher Date'], format='%m/%d/%Y %H:%M')
    operations: pd.DataFrame = pd.concat(
        [financial.set_index('Voucher Date'), df_transport.set_index('v_date'), df_accommodation.set_index('v_date'),
         df_unproductive.set_index('v_date')], axis=1)
    operations.loc[:, 'productive_md'] = operations['acco_md'] - operations['unproductive_md']
    return operations


def page_separator(head: str, document):
    text = document.add_paragraph()
    text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text = text.add_run(f'\n\n\n{head.upper()}')
    text.bold = True
    text.font.color.rgb = RGBColor(153, 37, 43)
    text.font.size = Pt(80)
    document.add_page_break()


def narration_refine(row):
    sample_text = row['Narration']
    start_index = sample_text.find('|')
    end_index = sample_text.find('|', start_index + 1)
    return sample_text[start_index + 1:end_index].title()


def abnormal_trn(fGL: pd.DataFrame, end_date: datetime, dCoAAdler: pd.DataFrame):
    start_date: datetime = datetime(year=end_date.year, month=end_date.month, day=1)
    fGL = fGL.loc[fGL['Narration'].str.contains(r'\|[^|]+\|', regex=True) & (fGL['Voucher Date'] >= start_date) & (
            fGL['Voucher Date'] <= end_date) & (fGL['Ledger_Code'] >= 5000000000), ['Voucher Date', 'Narration',
                                                                                    'Ledger_Code', 'Amount',
                                                                                    'Voucher Number']]
    fGL['Narration'] = fGL.apply(narration_refine, axis=1)
    fGL['Amount'] = fGL['Amount'] * -1
    fGL = fGL.groupby(by=['Narration', 'Ledger_Code', 'Voucher Number'], as_index=False)['Amount'].sum()
    fGL.sort_values(by='Ledger_Code', inplace=True)
    fGL = pd.merge(left=fGL, right=dCoAAdler[['Ledger_Name']], on='Ledger_Code', how='left').drop(
        columns=['Ledger_Code', 'Voucher Number']).rename(
        columns={'Narration': 'Description', 'Ledger_Name': 'Account'})
    return fGL


def cell_background(table, row: int, column: list, original: float, compare: float, good: str, bad: str):
    result = good if original >= compare else bad
    for idx, cell in enumerate(table.rows[row].cells):
        if idx in column:
            cell_xml_element = cell._tc
            table_cell_properties = cell_xml_element.get_or_add_tcPr()
            shade_obj = OxmlElement('w:shd')
            shade_obj.set(qn('w:fill'), result)
            table_cell_properties.append(shade_obj)


def job_profitability(fTimesheet:pd.DataFrame,fGL:pd.DataFrame,end_date:datetime,dEmployee:pd.DataFrame,dExclude:pd.DataFrame,fOT:pd.DataFrame,fInvoices:pd.DataFrame,cogs_map:dict,dJobs:pd.DataFrame)->pd.DataFrame:

    start_date:datetime = datetime(year=end_date.year,month=1,day=1)
    periods :list =  pd.date_range(start=start_date, end=end_date, freq='ME').to_pydatetime().tolist()
    fGL = fGL.loc[:,['Cost Center','Voucher Date','Ledger_Code','Amount','Third_Level_Group_Name','Second_Level_Group_Name']]
    fGL = fGL.loc[~fGL['Ledger_Code'].isin([5010101002,5010101003])]
    emp_list_full :list = dEmployee.index.tolist()
    driversandcleaners:list = dEmployee.loc[dEmployee['Designation'].isin(['HEAVY DRIVER','DRIVER','CAMP SUPERVISOR'])].index.tolist()
    emp_list :list = [i for i in emp_list_full if i not in driversandcleaners]
    timesheet_sum :dict = {'dc_emp_beni':None,'dc_trpt':None,'dc_out':None,'dc_sal':None}
    timesheet_jobs :dict = {'dc_emp_beni':None,'dc_trpt':None,'dc_out':None,'dc_sal':None}
    timesheet_grand_sum :dict = {'dc_emp_beni':None,'dc_trpt':None,'dc_out':None,'dc_sal':None}
    periodic_allocation :dict = {}

    for period in periods:
        st_date :datetime = period + relativedelta(day=1)
        fGL_fitlered :pd.DataFrame = fGL.loc[(fGL['Voucher Date']>=st_date) & (fGL['Voucher Date']<=period) & 
                    (fGL['Second_Level_Group_Name'] == 'Manpower Cost') ,['Cost Center','Voucher Date','Ledger_Code','Amount']]
        fGL_emp :pd.DataFrame = fGL_fitlered.loc[fGL_fitlered['Cost Center'].isin(emp_list)]
        fGL_other :pd.DataFrame = fGL_fitlered.loc[~fGL_fitlered['Cost Center'].isin(emp_list),['Amount','Ledger_Code']].groupby('Ledger_Code',as_index=False)['Amount'].sum()
        fGL_emp = fGL_emp.groupby(by=['Cost Center','Voucher Date','Ledger_Code'],as_index=False)['Amount'].sum()
        fGL_emp = fGL_emp.loc[fGL_emp['Amount']!=0]
        # TODO You may group this to cogs map using the ledger code. to be fixed. it will reduce the no of iteretion by approx 12.5%
        fTimesheet_filtered :pd.DataFrame = fTimesheet.loc[(fTimesheet['v_date'] >= st_date) & (fTimesheet['v_date']<=period)]
        fTimesheet_filtered = fTimesheet_filtered.groupby(['cost_center', 'job_id', 'v_date']).size().reset_index(name='count')
        billable_jobs:list = fTimesheet_filtered.loc[fTimesheet_filtered['job_id'].str.contains('ESS/CTR'),'job_id'].unique().tolist()
        
        for c in dExclude.columns:
            if c not in ['job_type','group']:
                valid_jobs :list = dExclude.loc[dExclude[c]==False]['job_type'].tolist() + billable_jobs
                timesheet_sum[c]  = fTimesheet_filtered.loc[fTimesheet_filtered['job_id'].isin(valid_jobs)].groupby(['cost_center','v_date'],as_index=False)['count'].sum()
                timesheet_jobs[c] = fTimesheet_filtered.loc[fTimesheet_filtered['job_id'].isin(valid_jobs)]
                timesheet_grand_sum[c]  = timesheet_sum[c]['count'].sum()
        allocation_dict :dict = {}
        unallocated_amount :float = 0
        for _,i in fGL_emp.iterrows():
            df_type :str = [(k,v) for k,v in cogs_map.items() if i['Ledger_Code'] in v][0][0]
            # TODO (a) YOU MAY FILTER df_sum/timesheet_sum and timesheet_detailed/timesheet_jobs only for those cost_centers apperiring in fGL_Emp. which will reduce the number of iterations.
            # Also filter by the ledger as well 
            df_sum :pd.DataFrame = timesheet_sum[df_type]
            timesheet_detailed:pd.DataFrame = timesheet_jobs[df_type]
            try:
                total_days: int = df_sum.loc[(df_sum['v_date'] == i['Voucher Date']) & (df_sum['cost_center'] == i['Cost Center']),'count'].iloc[0]
                timesheet_detailed = timesheet_detailed.loc[(timesheet_detailed['v_date']==i['Voucher Date']) & (timesheet_detailed['cost_center'] == i['Cost Center']),['job_id','count']]
                allocation_dict_init = {}
                for _,j in timesheet_detailed.iterrows():
                    # TODO (a) only those cost centers having a value will return a value from below. 
                    allocated :float =i['Amount']/total_days * j['count']
                    allocation_dict_init[j['job_id']] =  allocated
                allocation_dict = {k: allocation_dict_init.get(k,0) + allocation_dict.get(k,0) for k in set(allocation_dict)|set(allocation_dict_init)}
            except IndexError:
                unallocated_amount += i['Amount']
                allocation_dict['Un-Allocated'] = unallocated_amount
        fOT_filtered :pd.DataFrame = fOT.loc[(fOT['date'] >= st_date) & (fOT['date']<=period)]
        fOT_filtered :dict= fOT_filtered.groupby(by='job_id')['net'].sum().to_dict()
        allocation_dict = {k:allocation_dict.get(k,0) + fOT_filtered.get(k,0) for k in set(allocation_dict)|set(fOT_filtered)}
        inv_filtered_cust :dict= fInvoices.loc[(fInvoices['Invoice_Date'] >= st_date) & (fInvoices['Invoice_Date']<=period),['Order_ID','Net_Amount']].groupby('Order_ID')['Net_Amount'].sum().to_dict()
        allocation_dict = {k:allocation_dict.get(k,0) + inv_filtered_cust.get(k,0) for k in set(allocation_dict)|set(inv_filtered_cust)}
        for i in cogs_map:
            z:float = fGL_other.loc[fGL_other['Ledger_Code'].isin(cogs_map[i])]['Amount'].sum()
            if z != 0:
                for _,row in timesheet_jobs[i].groupby(by='job_id',as_index=False)['count'].sum().iterrows():
                    overhead_allocation :dict ={}
                    value:float = z / timesheet_grand_sum[i] * row['count']
                    overhead_allocation[row['job_id']] = value
                    allocation_dict = {k:allocation_dict.get(k,0) + overhead_allocation.get(k,0) for k in set(allocation_dict)|set(overhead_allocation)}
        acc_types :list = dExclude.loc[dExclude['group'].isin(['Accommodation']),'job_type'].tolist()
        accommodation_cost :float = sum([v for k,v in allocation_dict.items() if k in acc_types])
        non_accomo_sum :int = fTimesheet_filtered.loc[~fTimesheet_filtered['job_id'].isin(acc_types)]['count'].sum()
        non_accomo :pd.DataFrame = fTimesheet_filtered.loc[~fTimesheet_filtered['job_id'].isin(acc_types)]
        for _,row in non_accomo.iterrows():
            accommodation_allocation :dict = {}
            value :float = accommodation_cost/non_accomo_sum * row['count']
            accommodation_allocation[row['job_id']] = value
            allocation_dict = {k:allocation_dict.get(k,0) + accommodation_allocation.get(k,0) for k in set(allocation_dict)|set(accommodation_allocation)}
        del allocation_dict['AC-ACCOMODATION']
        del allocation_dict['AC']
        periodic_allocation[period] = allocation_dict


    cy_cp:pd.DataFrame = pd.DataFrame(list(periodic_allocation[end_date].items()),columns=['Order_ID','Amount'])
    cy_cp = pd.merge(left=cy_cp,right=dJobs[['Order_ID','Customer_Code','Employee_Code']],on='Order_ID',how='left')
    cy_cp_cus :pd.DataFrame = cy_cp.groupby(by='Customer_Code',as_index=False)['Amount'].sum()
    cy_cp_emp :pd.DataFrame= cy_cp.groupby(by='Employee_Code',as_index=False)['Amount'].sum()
    cy_ytd:pd.DataFrame = pd.DataFrame()
    for period in periods:
        month_df :pd.DataFrame = pd.DataFrame(list(periodic_allocation[period].items()),columns=['Order_ID','Amount'])
        cy_ytd = pd.concat([month_df,cy_ytd])
    cy_ytd = pd.merge(left=cy_ytd,right=dJobs[['Order_ID','Customer_Code','Employee_Code']],on='Order_ID',how='left')
    cy_ytd_cus:pd.DataFrame = cy_ytd.groupby(by='Customer_Code',as_index=False)['Amount'].sum()
    cy_ytd_emp:pd.DataFrame = cy_ytd.groupby(by='Employee_Code',as_index=False)['Amount'].sum()
    return {'periodic_allocation':periodic_allocation,'cy_cp_cus':cy_cp_cus,'cy_ytd_cus':cy_ytd_cus,'cy_cp_emp':cy_cp_emp,'cy_ytd_emp':cy_ytd_emp}

company_id = 0
end_date: datetime = datetime(year=2024, month=7, day=31)
start_date: datetime = datetime(year=end_date.year - 1, month=1, day=1)
sys_cut_off: datetime = datetime(year=2020, month=11, day=1)
VOUCHER_TYPES: list = ['Project Invoice',
                       'Contract Invoice', 'SERVICE INVOICE', 'Sales Invoice']

cleaned_data: dict = preprocessing(data=data_sources(company_id=0))
fGL: pd.DataFrame = cleaned_data['fGL']
dEmployee: pd.DataFrame = cleaned_data['dEmployee']
dCoAAdler: pd.DataFrame = cleaned_data['dCoAAdler']
merged: pd.DataFrame = pd.merge(
    left=fGL, right=dCoAAdler, on='Ledger_Code', how='left')
fInvoices: pd.DataFrame = cleaned_data['fInvoices']
fBudget: pd.DataFrame = cleaned_data['fBudget']
fCollection: pd.DataFrame = cleaned_data['fCollection']
dCustomers: pd.DataFrame = cleaned_data['dCustomers']
dJobs: pd.DataFrame = cleaned_data['dJobs']
fTimesheet: pd.DataFrame = cleaned_data['fTimesheet']
fOT: pd.DataFrame = cleaned_data['fOT']
dExclude: pd.DataFrame = cleaned_data['dExclude']


profitability: dict = job_profitability(fTimesheet=fTimesheet, fGL=merged, end_date=end_date, dEmployee=dEmployee,
                                        dExclude=dExclude, fOT=fOT, fInvoices=fInvoices, cogs_map=cogs_ledger_map,
                                        dJobs=dJobs)
cy_cp_profit_cus: pd.DataFrame = profitability['cy_cp_cus']
cy_ytd_profit_cus: pd.DataFrame = profitability['cy_ytd_cus']
cy_cp_profit_emp: pd.DataFrame = profitability['cy_cp_emp']
cy_ytd_profit_emp: pd.DataFrame = profitability['cy_ytd_emp']

financial_periods_bs: list = sorted(list(
    set([end_date, datetime(year=end_date.year - 1, month=end_date.month, day=end_date.day)] + list(
        pd.date_range(start=fGL['Voucher Date'].min(), end=end_date, freq='YE')))), reverse=True)
bscombined: pd.DataFrame = pd.DataFrame()
for f_year in financial_periods_bs:
    bs: pd.DataFrame = balancesheet(data=merged, end_date=f_year).rename(columns={'Amount': f'{f_year.date()}'})
    bscombined = pd.concat([bscombined, bs], axis=1)
bscombined = bscombined.reset_index().rename(columns={'index': 'Description'})
bscombined.fillna(value=0, inplace=True)

financial_periods_pl: list = sorted(list(
    set([end_date] + pd.date_range(start=fGL['Voucher Date'].min(), end=end_date, freq='YE').to_pydatetime().tolist())),
    reverse=True)
plcombined: pd.DataFrame = pd.DataFrame()
bu_plcombined = fGL['Bussiness Unit Name'].unique()
for f_year in financial_periods_pl:
    pl: dict = profitandloss(data=merged, end_date=f_year,
                             start_date=max(sys_cut_off, datetime(year=f_year.year, month=1, day=1)),
                             basic_pl=True, bu=bu_plcombined)
    pl_period: pd.DataFrame = pl['df_basic']['cy_ytd_basic'].rename(columns={'Amount': f'{f_year.date()}'}).set_index(
        keys='Description')
    plcombined = pd.concat([plcombined, pl_period], axis=1)
plcombined = plcombined.reset_index()

df_pl: dict = profitandloss(basic_pl=True, data=merged, start_date=start_date, end_date=end_date, full_pl=True,
                            bu=bu_plcombined)

cy_cp_basic: pd.DataFrame = df_pl['df_basic']['cy_cp_basic']
cy_ytd_basic: pd.DataFrame = df_pl['df_basic']['cy_ytd_basic']
cy_pp_basic: pd.DataFrame = df_pl['df_basic']['cy_pp_basic']
py_cp_basic: pd.DataFrame = df_pl['df_basic']['py_cp_basic']
py_ytd_basic: pd.DataFrame = df_pl['df_basic']['py_ytd_basic']
cy_cp_basic_bud: pd.DataFrame = df_pl['df_basic']['cy_cp_basic_bud']
cy_ytd_basic_bud: pd.DataFrame = df_pl['df_basic']['cy_ytd_basic_bud']
ratios_pandl: dict = plratios(df_pl=df_pl, plcombined=plcombined)

sort_order: list = coa_ordering(dCoAAdler=dCoAAdler)

cp_month: pd.DataFrame = pd.concat(
    [cy_cp_basic.set_index('Description'), cy_pp_basic.set_index('Description'), py_cp_basic.set_index('Description'),
     cy_cp_basic_bud.set_index('Description')],
    axis=1, join='outer').reset_index()
cp_month.fillna(value=0, inplace=True)
cp_month['Description'] = pd.Categorical(cp_month['Description'], categories=[k for k in sort_order.keys()],
                                         ordered=True)
cp_month.sort_values(by='Description', inplace=True)

document = Document()
doc = first_page(document=document, report_date=end_date)
document.add_page_break()

page_separator(head='Finance', document=document)

cy_cp_pl_company_title = document.add_paragraph().add_run('Elite Security Services W.L.L')
apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
cy_cp_pl_report_title = document.add_paragraph().add_run('Profit & Loss for the current period')
apply_style_properties(cy_cp_pl_report_title, style_picker(name='report_title'))

tbl_month_basic = document.add_table(rows=1, cols=5)
tbl_month_basic.columns[0].width = Cm(7.5)
heading_cells = tbl_month_basic.rows[0].cells
heading_cells[0].text = 'Description'
heading_cells[1].text = 'Current Month'
heading_cells[2].text = 'Previous Month'
heading_cells[3].text = 'SPLY'
heading_cells[4].text = 'Budget'

for _, row in cp_month.iterrows():
    cells = tbl_month_basic.add_row().cells
    cells[0].text = str(row['Description'])
    cells[1].text = number_format(row.iloc[1])
    cells[2].text = number_format(row.iloc[2])
    cells[3].text = number_format(row.iloc[3])
    cells[4].text = number_format(row.iloc[4])

plheads: list = ['Total Revenue', 'Gross Profit', 'Total Overhead', 'Net Profit']
table_formatter(table_name=tbl_month_basic, style_name='table_style_1', special=plheads)
document.add_page_break()

cy_cp_full: pd.DataFrame = df_pl['df_full']['cy_cp_full']
cy_pp_full: pd.DataFrame = df_pl['df_full']['cy_pp_full']
py_cp_full: pd.DataFrame = df_pl['df_full']['py_cp_full']
cy_cp_full_bud: pd.DataFrame = df_pl['df_full']['cy_cp_full_bud']

cp_month_full: pd.DataFrame = pd.concat(
    [cy_cp_full.set_index('Description'), cy_pp_full.set_index('Description'), py_cp_full.set_index('Description'),
     cy_cp_full_bud.set_index('Description')],
    axis=1, join='outer').reset_index()
cp_month_full.fillna(value=0, inplace=True)
cp_month_full['Description'] = pd.Categorical(cp_month_full['Description'], categories=[k for k in sort_order.keys()],
                                              ordered=True)
cp_month_full.sort_values(by='Description', inplace=True)

cy_cp_pl_company_title = document.add_paragraph().add_run('Elite Security Services W.L.L')
apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
cy_cp_pl_full_report_title = document.add_paragraph().add_run('Complete Profit & Loss for the current period')
apply_style_properties(cy_cp_pl_report_title, style_picker(name='report_title'))
tbl_month_full = document.add_table(rows=1, cols=5)
tbl_month_full.columns[0].width = Cm(11)
heading_cells = tbl_month_full.rows[0].cells
heading_cells[0].text = 'Description'
heading_cells[1].text = 'Current Month'
heading_cells[2].text = 'Previous Month'
heading_cells[3].text = 'SPLY'
heading_cells[4].text = 'Budget'

for _, row in cp_month_full.iterrows():
    cells = tbl_month_full.add_row().cells
    cells[0].text = str(row['Description'])
    cells[1].text = number_format(row.iloc[1])
    cells[2].text = number_format(row.iloc[2])
    cells[3].text = number_format(row.iloc[3])
    cells[4].text = number_format(row.iloc[4])
table_formatter(table_name=tbl_month_full, style_name='table_style_1', special=plheads)

abnormal_trn_title = document.add_paragraph().add_run('Explanations for Major Changes')
apply_style_properties(abnormal_trn_title, style_picker(name='report_title'))

abnormal_df: pd.DataFrame = abnormal_trn(fGL=fGL, end_date=end_date, dCoAAdler=dCoAAdler)

abnormal_tbl = document.add_table(rows=1, cols=3)
abnormal_tbl.columns[0].width = Cm(11)
heading_cells = abnormal_tbl.rows[0].cells
heading_cells[0].text = 'Description'
heading_cells[1].text = 'Account'
heading_cells[2].text = 'Amount'

for _, row in abnormal_df.iterrows():
    cells = abnormal_tbl.add_row().cells
    cells[0].text = str(row['Description'])
    cells[1].text = str(row['Account'])
    cells[2].text = number_format(row.iloc[1])

table_formatter(table_name=abnormal_tbl, style_name='table_style_1', special=[])

document.add_page_break()

cy_ytd_basic: pd.DataFrame = df_pl['df_basic']['cy_ytd_basic']
py_ytd_basic: pd.DataFrame = df_pl['df_basic']['py_ytd_basic']
cy_ytd_basic_bud: pd.DataFrame = df_pl['df_basic']['cy_ytd_basic_bud']

cp_ytd: pd.DataFrame = pd.concat(
    [cy_ytd_basic.set_index('Description'), py_ytd_basic.set_index('Description'),
     cy_ytd_basic_bud.set_index('Description')], axis=1, join='outer').reset_index()
cp_ytd.fillna(value=0, inplace=True)
cp_ytd['Description'] = pd.Categorical(cp_ytd['Description'], categories=[k for k in sort_order.keys()],
                                       ordered=True)
cp_ytd.sort_values(by='Description', inplace=True)
cy_cp_pl_company_title = document.add_paragraph().add_run('Elite Security Services W.L.L')
apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
cy_ytd_pl_report_title = document.add_paragraph().add_run('Profit & Loss for Year to Date')
apply_style_properties(cy_cp_pl_report_title, style_picker(name='report_title'))

tbl_ytd_basic = document.add_table(rows=1, cols=4)
tbl_ytd_basic.columns[0].width = Cm(11)
heading_cells = tbl_ytd_basic.rows[0].cells
heading_cells[0].text = 'Description'
heading_cells[1].text = 'YTD CY'
heading_cells[2].text = 'YTD PY'
heading_cells[3].text = 'Budget'

for _, row in cp_ytd.iterrows():
    cells = tbl_ytd_basic.add_row().cells
    cells[0].text = str(row['Description'])
    cells[1].text = number_format(row.iloc[1])
    cells[2].text = number_format(row.iloc[2])
    cells[3].text = number_format(row.iloc[3])

table_formatter(table_name=tbl_ytd_basic, style_name='table_style_1', special=plheads)
document.add_page_break()

plt.style.use('ggplot')
fig_pl, (ax1, ax2) = plt.subplots(nrows=2, ncols=1,figsize = (7.27,10))

ratiopl: pd.DataFrame = ratios_pandl['gp']['cy_ytd_basic_monthwise']
ax1.set_title(f'GP Vs NP VS EBITDA - {end_date.year}')
ax1.plot([i.strftime('%b') for i in ratiopl['Voucher Date']],
         (ratiopl['Gross Profit'] / ratiopl['Total Revenue'] * 100),
         label='GP')
ax1.plot([i.strftime('%b') for i in ratiopl['Voucher Date']], (ratiopl['EBITDA'] / ratiopl['Total Revenue'] * 100),
         label='EBITDA')
ax1.plot([i.strftime('%b') for i in ratiopl['Voucher Date']], (ratiopl['Net Profit'] / ratiopl['Total Revenue'] * 100),
         label='NP')

ax1.set_yticklabels(['{:,.0f}%'.format(i) for i in ax1.get_yticks()])
ax1.legend()

ratioplyearly: pd.DataFrame = ratios_pandl['plyearly']
ax2.set_title(f'GP Vs NP VS EBITDA ({sys_cut_off.year}-{end_date.year})')
ax2.plot([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in ratioplyearly['period']],
         (ratioplyearly['gp'] / ratioplyearly['revenue'] * 100),
         label='GP')
ax2.plot([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in ratioplyearly['period']],
         (ratioplyearly['ebitda'] / ratioplyearly['revenue'] * 100),
         label='EBITDA')
ax2.plot([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in ratioplyearly['period']],
         (ratioplyearly['np'] / ratioplyearly['revenue'] * 100),
         label='NP')

ax2.invert_xaxis()
ax2.set_yticklabels(['{:,.0f}%'.format(i) for i in ax2.get_yticks()])
ax2.legend()

pl_graph_buf = BytesIO()
plt.tight_layout(h_pad=3)
plt.savefig(pl_graph_buf, format='png')
plt.close(fig_pl)
pl_graph_buf.seek(0)
doc.add_picture(pl_graph_buf)
document.add_page_break()

cy_ytd_basic_monthwise: pd.DataFrame = df_pl['df_basic']['cy_ytd_basic_monthwise']
cy_ytd_basic_monthwise.fillna(value=0, inplace=True)
cy_ytd_basic_monthwise['Description'] = pd.Categorical(cy_ytd_basic_monthwise['Description'],
                                                       categories=[k for k in sort_order.keys()],
                                                       ordered=True)
cy_ytd_basic_monthwise.sort_values(by='Description', inplace=True)
change_orientation(doc=document, method='l')

cy_cp_pl_company_title = document.add_paragraph().add_run('Elite Security Services W.L.L')
apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
cy_ytd_pl_report_title = document.add_paragraph().add_run('Profit & Loss for Year to Date Month-Wise')
apply_style_properties(cy_cp_pl_report_title, style_picker(name='report_title'))

tbl_monthwise_basic = document.add_table(rows=1, cols=cy_ytd_basic_monthwise.shape[1])
tbl_monthwise_basic.columns[0].width = Cm(7.5)
heading_cells = tbl_monthwise_basic.rows[0].cells

for i in range(cy_ytd_basic_monthwise.shape[1]):
    if i == 0:
        heading_cells[i].text = 'Description'
    else:
        heading_cells[i].text = list(cy_ytd_basic_monthwise.columns)[i].strftime('%b')

for _, row in cy_ytd_basic_monthwise.iterrows():
    cells = tbl_monthwise_basic.add_row().cells
    for j in range(len(row)):
        if j == 0:
            cells[0].text = str(row['Description'])
        else:
            cells[j].text = number_format(row.iloc[j])
table_formatter(table_name=tbl_monthwise_basic, style_name='table_style_1', special=plheads)

document.add_page_break()

cy_cp_pl_company_title = document.add_paragraph().add_run('Elite Security Services W.L.L')
apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
cy_ytd_pl_guard_bu_title = document.add_paragraph().add_run('Profit & Loss for Guarding Division')
apply_style_properties(cy_ytd_pl_guard_bu_title, style_picker(name='report_title'))
df_pl_bu_guarding: dict = profitandloss(basic_pl=True, data=merged, start_date=start_date, end_date=end_date,
                                        full_pl=False, bu=['GUARDING-ESS'])
cy_cp_basic_guarding_bu: pd.DataFrame = df_pl_bu_guarding['df_basic']['cy_ytd_basic_monthwise']
cy_cp_basic_guarding_bu.fillna(value=0, inplace=True)
cy_cp_basic_guarding_bu['Description'] = pd.Categorical(cy_cp_basic_guarding_bu['Description'],
                                                        categories=[k for k in sort_order.keys()],
                                                        ordered=True)
cy_cp_basic_guarding_bu.sort_values(by='Description', inplace=True)

tbl_monthwise_basic_guarding_bu = document.add_table(rows=1, cols=cy_cp_basic_guarding_bu.shape[1])
tbl_monthwise_basic_guarding_bu.columns[0].width = Cm(7.5)
heading_cells = tbl_monthwise_basic_guarding_bu.rows[0].cells

for i in range(cy_cp_basic_guarding_bu.shape[1]):
    if i == 0:
        heading_cells[i].text = 'Description'
    else:
        heading_cells[i].text = list(cy_cp_basic_guarding_bu.columns)[i].strftime('%b')

for _, row in cy_cp_basic_guarding_bu.iterrows():
    cells = tbl_monthwise_basic_guarding_bu.add_row().cells
    for j in range(len(row)):
        if j == 0:
            cells[0].text = str(row['Description'])
        else:
            cells[j].text = number_format(row.iloc[j])
table_formatter(table_name=tbl_monthwise_basic_guarding_bu, style_name='table_style_1', special=plheads)

document.add_page_break()

cy_cp_pl_company_title = document.add_paragraph().add_run('Elite Security Services W.L.L')
apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
cy_ytd_pl_elv_bu_title = document.add_paragraph().add_run('Profit & Loss for ELV Division')
apply_style_properties(cy_ytd_pl_elv_bu_title, style_picker(name='report_title'))
df_pl_bu_elv: dict = profitandloss(basic_pl=True, data=merged, start_date=start_date, end_date=end_date, full_pl=False,
                                   bu=['ELV-ESS'])
cy_cp_basic_elv_bu: pd.DataFrame = df_pl_bu_elv['df_basic']['cy_ytd_basic_monthwise']
cy_cp_basic_elv_bu.fillna(value=0, inplace=True)
cy_cp_basic_elv_bu['Description'] = pd.Categorical(cy_cp_basic_elv_bu['Description'],
                                                   categories=[k for k in sort_order.keys()],
                                                   ordered=True)
cy_cp_basic_elv_bu.sort_values(by='Description', inplace=True)

tbl_monthwise_basic_elv_bu = document.add_table(rows=1, cols=cy_cp_basic_elv_bu.shape[1])
tbl_monthwise_basic_elv_bu.columns[0].width = Cm(7.5)
heading_cells = tbl_monthwise_basic_elv_bu.rows[0].cells

for i in range(cy_cp_basic_elv_bu.shape[1]):
    if i == 0:
        heading_cells[i].text = 'Description'
    else:
        heading_cells[i].text = list(cy_cp_basic_elv_bu.columns)[i].strftime('%b')

for _, row in cy_cp_basic_elv_bu.iterrows():
    cells = tbl_monthwise_basic_elv_bu.add_row().cells
    for j in range(len(row)):
        if j == 0:
            cells[0].text = str(row['Description'])
        else:
            cells[j].text = number_format(row.iloc[j])
table_formatter(table_name=tbl_monthwise_basic_elv_bu, style_name='table_style_1', special=plheads)
document.add_page_break()

change_orientation(doc=document, method='p')

cy_cp_pl_company_title = document.add_paragraph().add_run('Elite Security Services W.L.L')
apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
pl_report_title = document.add_paragraph().add_run('Historical Profit and Loss Comparison')
apply_style_properties(pl_report_title, style_picker(name='report_title'))

plcombined.reset_index(inplace=True)
tbl_yearly_pl = document.add_table(rows=1, cols=plcombined.shape[1])
heading_cells = tbl_yearly_pl.rows[0].cells
plcombined.fillna(value=0, inplace=True)
plcombined['Description'] = pd.Categorical(plcombined['Description'], categories=[k for k in sort_order.keys()],
                                           ordered=True)
plcombined.sort_values(by='Description', inplace=True)

for i in range(plcombined.shape[1]):
    if i == 0:
        heading_cells[i].text = 'Description'
    else:
        heading_cells[i].text = list(plcombined.columns)[i]

for _, row in plcombined.iterrows():
    cells = tbl_yearly_pl.add_row().cells
    for j in range(len(row)):
        if j == 0:
            cells[0].text = str(row['Description'])
        else:
            cells[j].text = f"{row.iloc[j]:,.0f}" if row.iloc[j] >= 0 else f"({abs(row.iloc[j]):,.0f})"

table_formatter(table_name=tbl_yearly_pl, style_name='table_style_1', special=plheads)
document.add_page_break()

change_orientation(doc=document, method='l')
cy_cp_pl_company_title = document.add_paragraph().add_run('Elite Security Services W.L.L')
apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
cy_mw_bs_report_title = document.add_paragraph().add_run('Statement of Financial Position (Balance Sheet)')
apply_style_properties(cy_mw_bs_report_title, style_picker(name='report_title'))
bscombined['Description'] = pd.Categorical(bscombined['Description'],
                                           categories=[k for k in sort_order.keys()],
                                           ordered=True)
bscombined.sort_values(by='Description', inplace=True)

tbl_yearly_bs = document.add_table(rows=1, cols=bscombined.shape[1])
heading_cells = tbl_yearly_bs.rows[0].cells
for i in range(bscombined.shape[1]):
    if i == 0:
        heading_cells[i].text = 'Description'
    else:
        heading_cells[i].text = list(bscombined.columns)[i]

for _, row in bscombined.iterrows():
    cells = tbl_yearly_bs.add_row().cells
    for j in range(len(row)):
        if j == 0:
            cells[0].text = str(row['Description'])
        else:
            cells[j].text = f"{row.iloc[j]:,.0f}" if row.iloc[j] >= 0 else f"({abs(row.iloc[j]):,.0f})"

table_formatter(table_name=tbl_yearly_bs, style_name='table_style_1',
                special=['Current Liabilities', 'Non Current Liabilities', 'Liabilities', 'Equity', 'Current Assets',
                         'Non Current Assets', 'Assets', 'Total Equity & Liabilities'])
document.add_page_break()

change_orientation(doc=document, method='p')
interco: dict = interco_bal(data=merged, end_date=end_date)
rpr_df: pd.DataFrame = interco.get('rpr_df')
rpr_total_row: pd.DataFrame = pd.DataFrame(data={'Amount': [rpr_df['Amount'].sum()], 'Description': 'Total'}, index=[
    '9999'])
rpr_df = pd.concat([rpr_df, rpr_total_row], ignore_index=False)

cy_cp_pl_company_title = document.add_paragraph().add_run('Elite Security Services W.L.L')
apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
rpr_report_title = document.add_paragraph().add_run('Break-up of Related-Party Receiavable')
apply_style_properties(rpr_report_title, style_picker(name='report_title'))

tbl_rpr = document.add_table(rows=1, cols=2)
heading_cells = tbl_rpr.rows[0].cells
heading_cells[0].text = 'Description'
heading_cells[1].text = 'Amount'

for _, row in rpr_df.iterrows():
    cells = tbl_rpr.add_row().cells
    cells[0].text = str(row['Description'])
    cells[1].text = number_format(-row.iloc[1])

table_formatter(table_name=tbl_rpr, style_name='table_style_1', special=['Total'])

rpp_df: pd.DataFrame = interco.get('rpp_df')
rpp_total_row: pd.DataFrame = pd.DataFrame(data={'Amount': [rpp_df['Amount'].sum()], 'Description': 'Total'}, index=[
    '9999'])
rpp_df = pd.concat([rpp_df, rpp_total_row], ignore_index=False)

rpp_report_title = document.add_paragraph().add_run('\n\nBreak-up of Related-Party Payables')
apply_style_properties(rpp_report_title, style_picker(name='report_title'))
tbl_rpp = document.add_table(rows=1, cols=2)
heading_cells = tbl_rpp.rows[0].cells
heading_cells[0].text = 'Description'
heading_cells[1].text = 'Amount'

for _, row in rpp_df.iterrows():
    cells = tbl_rpp.add_row().cells
    cells[0].text = str(row['Description'])
    cells[1].text = number_format(row.iloc[1])

table_formatter(table_name=tbl_rpp, style_name='table_style_1', special=['Total'])

document.add_page_break()

plt.style.use('ggplot')
fig_bs, (ax1, ax2, ax3) = plt.subplots(nrows=3, ncols=1, sharex=True,figsize = (7.27,10))

bs_ratios_df: pd.DataFrame = bsratios(bsdata=bscombined, pldata=plcombined, periods=financial_periods_bs,
                                      end_date=end_date)

ax1.set_title('Current Ratio')
ax1.plot([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in bs_ratios_df['period']], bs_ratios_df['cr'])
ax1.set_yticklabels(['{:,.2f}'.format(i) for i in ax1.get_yticks()])

ax2.set_title('Assets Turnover Ratio')
ax2.plot([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in bs_ratios_df['period']], bs_ratios_df['ato'])
ax2.set_yticklabels(['{:,.2f}'.format(i) for i in ax2.get_yticks()])

ax3.set_title('Return on Equity')
ax3.plot([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in bs_ratios_df['period']], bs_ratios_df['roe'])
ax3.set_yticklabels(['{:,.0f}%'.format(i) for i in ax3.get_yticks()])
ax3.invert_xaxis()

bs_graph_buf = BytesIO()
plt.tight_layout(h_pad=3)
plt.savefig(bs_graph_buf, format='png')
plt.close(fig_bs)
bs_graph_buf.seek(0)
doc.add_picture(bs_graph_buf)
document.add_page_break()

page_separator(head='Sales', document=document)

change_orientation(doc=document, method='l')
rev_summary = plt.figure()
rev_summary.set_figheight(7)
rev_summary.set_figwidth(10.5)

ini_shape = (4, 5)
ax1 = plt.subplot2grid(shape=ini_shape, loc=(0, 0), colspan=4)
ax2 = plt.subplot2grid(shape=ini_shape, loc=(1, 0), colspan=2)
ax3 = plt.subplot2grid(shape=ini_shape, loc=(1, 2), colspan=2)
ax4 = plt.subplot2grid(shape=ini_shape, loc=(2, 0), colspan=4)
ax5 = plt.subplot2grid(shape=ini_shape, loc=(3, 0), colspan=2)
ax6 = plt.subplot2grid(shape=ini_shape, loc=(3, 2), colspan=2)
ax7 = plt.subplot2grid(shape=ini_shape, loc=(0, 4), colspan=1)
ax8 = plt.subplot2grid(shape=ini_shape, loc=(1, 4), colspan=1)
ax9 = plt.subplot2grid(shape=ini_shape, loc=(2, 4), colspan=1)
ax10 = plt.subplot2grid(shape=ini_shape, loc=(3, 4), colspan=1)

df_rev: dict = revenue(end_date=end_date, data=merged, fInvoices=fInvoices)
rev_division: pd.DataFrame = df_rev['rev_division']

rev_division_plot: pd.DataFrame = rev_division.copy()
rev_division = rev_division.loc[(rev_division['Voucher Date'] <= end_date) & (
        rev_division['Voucher Date'] >= plotting_period(end_date=end_date, months=6))].pivot_table(
    index='Second_Level_Group_Name', columns='Voucher Date', values='Amount',
    aggfunc='sum', fill_value=0).reset_index().rename(columns={'Second_Level_Group_Name': 'Description'})

rev_division_line: pd.DataFrame = rev_division_plot.loc[(rev_division_plot['Voucher Date'] <= end_date) & (
        rev_division_plot['Voucher Date'] >= plotting_period(end_date=end_date, months=6))].pivot_table(
    index='Voucher Date', columns='Second_Level_Group_Name', values='Amount',
    aggfunc='sum', fill_value=0).reset_index().rename(columns={'Voucher Date': 'Period'}).set_index(keys='Period')

rev_division_pie_ytd: pd.DataFrame = rev_division_plot.loc[(rev_division_plot['Voucher Date'] <= end_date) & (
        rev_division_plot['Voucher Date'] >= datetime(year=end_date.year, month=1, day=1)), [
    'Second_Level_Group_Name', 'Amount']].groupby(by='Second_Level_Group_Name').sum().reset_index().rename(
    columns={'Second_Level_Group_Name': 'Category'}).set_index(keys='Category')

rev_division_pie_month: pd.DataFrame = rev_division_plot.loc[(rev_division_plot['Voucher Date'] <= end_date) & (
        rev_division_plot['Voucher Date'] >= datetime(year=end_date.year, month=end_date.month, day=1)), [
    'Second_Level_Group_Name', 'Amount']].rename(columns={'Second_Level_Group_Name': 'Category'}).set_index(
    keys='Category')

rev_category: pd.DataFrame = df_rev['rev_category']

rev_category_plot: pd.DataFrame = rev_category.copy()

rev_category = rev_category.loc[(rev_category['Voucher Date'] <= end_date) & (
        rev_category['Voucher Date'] >= plotting_period(end_date=end_date, months=6))].pivot_table(index='Type',
                                                                                                   columns='Voucher '
                                                                                                           'Date',
                                                                                                   values='Amount',
                                                                                                   aggfunc='sum',
                                                                                                   fill_value=0).reset_index().rename(
    columns={'Type': 'Description'})

rev_category_line: pd.DataFrame = rev_category_plot.loc[(rev_category_plot['Voucher Date'] <= end_date) & (
        rev_category_plot['Voucher Date'] >= plotting_period(end_date=end_date, months=6))].pivot_table(
    index='Voucher Date', columns='Type', values='Amount',
    aggfunc='sum', fill_value=0).reset_index().rename(columns={'Voucher Date': 'Period'}).set_index(keys='Period')

rev_category_pie: pd.DataFrame = df_rev['rev_category']
rev_category_pie_ytd: pd.DataFrame = rev_category_pie.loc[(rev_category_pie['Voucher Date'] <= end_date) & (
        rev_category_pie['Voucher Date'] >= datetime(year=end_date.year, month=1, day=1)), ['Type',
                                                                                            'Amount']].groupby(
    by='Type').sum()

rev_category_pie_month: pd.DataFrame = rev_category_pie.loc[(rev_category_pie['Voucher Date'] <= end_date) & (
        rev_category_pie['Voucher Date'] >= datetime(year=end_date.year, month=end_date.month, day=1)), ['Type',
                                                                                                         'Amount']].groupby(
    by='Type').sum()

ax1.table(cellText=[[j[0]] + [f'{i:,.0f}' for i in j if isinstance(i, float)] for j in rev_category.values],
          colLabels=['Description'] + [i.strftime('%b') for i in rev_category.columns if i != 'Description'],
          cellLoc='left', loc='best')
ax1.set_title('Market/Related-party sales')
ax1.axis('off')
ax2.plot([i.strftime('%b') for i in rev_category_line.index], rev_category_line['Market'])
ax2.set_yticklabels(['{:,}'.format(int(i)) for i in ax2.get_yticks()])
ax2.set_title('Market Sales')
ax3.plot([i.strftime('%b') for i in rev_category_line.index], rev_category_line['Related'])
ax3.set_yticklabels(['{:,}'.format(int(i)) for i in ax3.get_yticks()])
ax3.set_title('Related Sales')

ax4.table(cellText=[[j[0]] + [f'{i:,.0f}' for i in j if isinstance(i, float)] for j in rev_division.values],
          colLabels=['Description'] + [i.strftime('%b') for i in rev_division.columns if i != 'Description'],
          cellLoc='left', loc='best')

ax4.set_title('Division wise monthly sales')
ax4.axis('off')

ax5.plot([i.strftime('%b') for i in rev_division_line.index], rev_division_line['Manpower'])
ax5.set_yticklabels(['{:,}'.format(int(i)) for i in ax5.get_yticks()])
ax5.set_title('Manpower Sales')
ax6.plot([i.strftime('%b') for i in rev_division_line.index], rev_division_line['Projects'])
ax6.set_yticklabels(['{:,}'.format(int(i)) for i in ax6.get_yticks()])
ax6.set_title('Projects Sales')
ax7.pie(x=rev_category_pie_month['Amount'], labels=rev_category_pie_month.index, autopct='%.0f%%', labeldistance=1,
        pctdistance=0.3)
ax7.set_title('Month')

ax8.pie(x=rev_category_pie_ytd['Amount'], labels=rev_category_pie_ytd.index, autopct='%.0f%%', labeldistance=1,
        pctdistance=0.3)
ax8.set_title('Year')

ax9.pie(x=rev_division_pie_month['Amount'], labels=rev_division_pie_month.index, autopct='%.1f%%', labeldistance=1,
        pctdistance=0.5)
ax9.set_title('Month')

ax10.pie(x=rev_division_pie_ytd['Amount'], labels=rev_division_pie_ytd.index, autopct='%.1f%%', labeldistance=1,
         pctdistance=0.5)
ax10.set_title('Year')

plt.tight_layout()

buf_revenue = BytesIO()
plt.tight_layout()
plt.savefig(buf_revenue, format='png', dpi=2400)
plt.close(rev_summary)
buf_revenue.seek(0)
doc.add_picture(buf_revenue)
doc.add_page_break()

plt.style.use('ggplot')
fig_sales, (new_existing, salesman_wise, col_graph) = plt.subplots(nrows=3, ncols=1, sharex=True)
fig_sales.set_figheight(7)
fig_sales.set_figwidth(10.5)

new_or_old: pd.DataFrame = df_rev['new_or_old'].groupby(['Invoice_Date', 'new_or_old'])['Net_Amount'].sum().unstack(
    fill_value=0)

new_or_old.index = [i.strftime('%b') for i in new_or_old.index]
new_or_old.plot(kind='bar', stacked=True, ax=new_existing)
new_existing.set_title('Revenue by Existing / New Customers')
new_existing.legend()
new_existing.set_yticklabels(['{:,}'.format(int(i)) for i in new_existing.get_yticks()])

inv_emp: pd.DataFrame = df_rev['inv_emp']
demp: pd.DataFrame = dEmployee.copy().reset_index()
inv_emp = pd.merge(left=inv_emp, right=demp[['Employee_Name', 'Employee_Code']], on='Employee_Code', how='left')
inv_emp['Employee_Name'] = inv_emp.apply(lambda x: ' '.join(x['Employee_Name'].split(sep=' ')[:2]).title(), axis=1)

inv_emp.drop(columns=['Employee_Code'], inplace=True)
inv_emp: pd.DataFrame = inv_emp.groupby(['Invoice_Date', 'Employee_Name'])['Net_Amount'].sum().unstack(fill_value=0)
inv_emp.index = [i.strftime('%b') for i in inv_emp.index]
inv_emp.plot(kind='bar', stacked=True, ax=salesman_wise)
salesman_wise.set_title('Revenue by Sales Person')
salesman_wise.legend(loc='best')
salesman_wise.set_yticklabels(['{:,}'.format(int(i)) for i in salesman_wise.get_yticks()])

monthly_collection: pd.DataFrame = collection(dCustomers=dCustomers, end_date=end_date, fGL=fGL,
                                              fCollection=fCollection)
col_graph.set_title('Target Collection Vs Actual Collection')
col_graph.plot([i.strftime('%b') for i in monthly_collection['Due Date']], monthly_collection['Target'],
               label='Target')
col_graph.plot([i.strftime('%b') for i in monthly_collection['Due Date']], monthly_collection['Actual'],
               label='Actual')
col_graph.set_yticklabels(['{:,}'.format(int(i)) for i in col_graph.get_yticks()])
col_graph.legend()

buf_sales = BytesIO()
plt.tight_layout()
plt.savefig(buf_sales, format='png', dpi=2400)
plt.close(fig_sales)
buf_sales.seek(0)
doc.add_picture(buf_sales)

doc.add_page_break()
change_orientation(doc=document, method='p')

customer_list: list = sorted(fInvoices.loc[(fInvoices['Invoice_Date'] >= datetime(year=end_date.year,
                                                                                  month=end_date.month, day=1)) & (
                                                   fInvoices['Invoice_Date'] <= end_date), 'Cus_Name'].unique())
customer_info: dict = customer_ratios(customers=customer_list, fInvoices=fInvoices, end_date=end_date,
                                      fCollection=fCollection, dCustomer=dCustomers, dEmployee=dEmployee)

heading_format = {'fontfamily': 'Georgia', 'color': 'k', 'fontweight': 'bold', 'fontsize': 10}

cy_cp_pl_company_title = document.add_paragraph().add_run('Elite Security Services W.L.L')
apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
cust_info_toc = document.add_paragraph().add_run('Key data about customers')
apply_style_properties(cust_info_toc, style_picker(name='report_title'))
tbl_cust_toc = document.add_table(rows=1, cols=2)
heading_cells = tbl_cust_toc.rows[0].cells
heading_cells[0].text = 'Customer Name'
heading_cells[1].text = 'Page #'

for idx, row in enumerate(customer_list):
    cells = tbl_cust_toc.add_row().cells
    cells[0].text = str(row.upper())
    cells[1].text = str(idx + 1)

table_formatter(table_name=tbl_cust_toc, style_name='table_style_1', special=[])
document.add_page_break()

for customer in customer_list:
    cus_code: list = dCustomers.loc[(dCustomers['Cus_Name'] == customer), 'Customer_Code'].tolist()
    cy_cp_pl_company_title = document.add_paragraph().add_run(customer.upper())
    apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
    tbl_cust_main = document.add_table(rows=2, cols=4)
    tbl_cust_main_th = tbl_cust_main.rows[0]
    tbl_cust_main_th.cells[0].text = 'Date of Establishment'
    tbl_cust_main_th.cells[1].text = 'Customer Since'
    tbl_cust_main_th.cells[2].text = 'Salesperson'
    tbl_cust_main_th.cells[3].text = 'Balance'
    tbl_cust_main_td = tbl_cust_main.rows[1]
    tbl_cust_main_td.cells[0].text = str(customer_info[customer]['date_established'])
    tbl_cust_main_td.cells[1].text = str(customer_info[customer]['customer_since'])
    tbl_cust_main_td.cells[2].text = ' '.join(
        str(customer_info[customer]['last_sales_person']).split(sep=' ')[:2]).title()
    tbl_cust_main_td.cells[3].text = number_format(num=customer_info[customer]['outstanding_bal'])
    table_formatter(table_name=tbl_cust_main, style_name='table_style_1', special=[])

    tbl_cust_rev_1 = document.add_table(rows=2, cols=4)
    tbl_cust_rev_th_1 = tbl_cust_rev_1.rows[0]
    tbl_cust_rev_th_1.cells[0].text = f'CY CP Revenue\n({end_date.strftime("%B")} Month)'
    tbl_cust_rev_th_1.cells[1].text = 'CY YTD Revenue'
    tbl_cust_rev_th_1.cells[2].text = 'CY CP Rev \nContribution'
    tbl_cust_rev_th_1.cells[
        3].text = f'CY PP Revenue\n({(end_date.replace(day=1) - timedelta(days=1)).strftime("%B")} Month)'
    tbl_cust_rev_td_1 = tbl_cust_rev_1.rows[1]
    arrow_type_cy: str = '↑' if customer_info[customer]['cy_cp_rev'] >= customer_info[customer]['py_cp_rev'] else '↓'
    arrow_type_ytd: str = '↑' if customer_info[customer]['cy_ytd_rev'] >= customer_info[customer]['py_ytd_rev'] else '↓'
    tbl_cust_rev_td_1.cells[0].text = f"{number_format(num=customer_info[customer]['cy_cp_rev'])} {arrow_type_cy}"
    tbl_cust_rev_td_1.cells[1].text = f"{number_format(num=customer_info[customer]['cy_ytd_rev'])} {arrow_type_ytd}"
    tbl_cust_rev_td_1.cells[2].text = str(customer_info[customer]['cy_cp_rev_contrib_pct'])
    tbl_cust_rev_td_1.cells[3].text = number_format(num=customer_info[customer]['cy_pp_rev'])
    table_formatter(table_name=tbl_cust_rev_1, style_name='table_style_1', special=[])

    cell_background(table=tbl_cust_rev_1, row=1, column=[0], original=customer_info[customer]['cy_cp_rev'],
                    compare=customer_info[customer]['py_cp_rev'], good='d2f3cf', bad='ffd8d5')
    cell_background(table=tbl_cust_rev_1, row=1, column=[1], original=customer_info[customer]['cy_ytd_rev'],
                    compare=customer_info[customer]['py_ytd_rev'], good='d2f3cf', bad='ffd8d5')

    tbl_cust_rev_2 = document.add_table(rows=2, cols=4)
    tbl_cust_rev_th_2 = tbl_cust_rev_2.rows[0]
    tbl_cust_rev_th_2.cells[0].text = 'PY CP Revenue'
    tbl_cust_rev_th_2.cells[1].text = 'PY YTD Revenue'
    tbl_cust_rev_th_2.cells[2].text = 'CY YTD Rev \nContribution'
    tbl_cust_rev_th_2.cells[3].text = 'Total Revenue Made'
    tbl_cust_rev_td_2 = tbl_cust_rev_2.rows[1]

    tbl_cust_rev_td_2.cells[0].text = number_format(num=customer_info[customer]['py_cp_rev'])
    tbl_cust_rev_td_2.cells[1].text = number_format(num=customer_info[customer]['py_ytd_rev'])
    tbl_cust_rev_td_2.cells[2].text = str(customer_info[customer]['cy_ytd_rev_contrib_pct'])
    tbl_cust_rev_td_2.cells[3].text = number_format(num=customer_info[customer]['total_revenue'])
    table_formatter(table_name=tbl_cust_rev_2, style_name='table_style_1', special=[])

    tbl_cust_col = document.add_table(rows=2, cols=4)
    tbl_cust_col_th = tbl_cust_col.rows[0]
    tbl_cust_col_th.cells[0].text = 'Credit Score'
    tbl_cust_col_th.cells[1].text = 'Credit Days'
    tbl_cust_col_th.cells[2].text = 'Median Collection\nDays'
    tbl_cust_col_th.cells[3].text = 'Last Collection Date\n and Amount'
    tbl_cust_col_td = tbl_cust_col.rows[1]
    tbl_cust_col_td.cells[0].text = number_format(customer_info[customer]['credit_score'])
    tbl_cust_col_td.cells[1].text = number_format(num=customer_info[customer]['credit_days'])
    tbl_cust_col_td.cells[2].text = str(customer_info[customer]['collection_median'])
    tbl_cust_col_td.cells[
        3].text = f"{number_format(num=customer_info[customer]['last_receipt_amt'])}\n{str(customer_info[customer]['last_receipt_dt'])}"
    table_formatter(table_name=tbl_cust_col, style_name='table_style_1', special=[])

    tbl_cust_gp = document.add_table(rows=2, cols=4)
    tbl_cust_gp_th = tbl_cust_gp.rows[0]
    tbl_cust_gp_th.cells[0].text = 'GP Month'
    tbl_cust_gp_th.cells[1].text = 'GP YTD'
    tbl_cust_gp_th.cells[2].text = 'ROI Month'
    tbl_cust_gp_th.cells[3].text = 'ROI YTD'
    tbl_cust_gp_td = tbl_cust_gp.rows[1]
    cp_gp_pct = round(cy_cp_profit_cus.loc[cy_cp_profit_cus['Customer_Code'].isin(cus_code), 'Amount'].sum() / customer_info[customer]['cy_cp_rev'] * 100,2)
    ytd_gp_pct = round(cy_ytd_profit_cus.loc[cy_ytd_profit_cus['Customer_Code'].isin(cus_code), 'Amount'].sum() / customer_info[customer]['cy_ytd_rev'] * 100,2)
    tbl_cust_gp_td.cells[0].text = f"{number_format(num=cy_cp_profit_cus.loc[cy_cp_profit_cus['Customer_Code'].isin(cus_code), 'Amount'].sum())} | {cp_gp_pct}%"
    tbl_cust_gp_td.cells[1].text = f"{number_format(num=cy_ytd_profit_cus.loc[cy_ytd_profit_cus['Customer_Code'].isin(cus_code), 'Amount'].sum())} | {ytd_gp_pct}%"
    tbl_cust_gp_td.cells[2].text = str(customer_info[customer]['cy_cp_roi'])
    tbl_cust_gp_td.cells[3].text = str(customer_info[customer]['cy_ytd_roi'])
    table_formatter(table_name=tbl_cust_gp, style_name='table_style_1', special=[])

    fig, ((age_tbl, age_pie), (rev_tbl, rev_bar)) = plt.subplots(nrows=2, ncols=2)

    ageing: pd.DataFrame = customer_info[customer]['ageing']

    ageing.reset_index(inplace=True)
    monthly_rev: pd.DataFrame = customer_info[customer]['monthly_rev']
    monthly_rev.reset_index(inplace=True)
    if not ageing.empty:
        age_tbl.set_title('Receivable Ageing', loc='left', **heading_format)
        age_tbl.table(cellText=[[i[0], f'{i[1]:,.0f}'] for i in ageing.values], colLabels=ageing.columns,
                      cellLoc='center', loc='center')
        age_tbl.axis('off')

        age_pie.pie(x=ageing['Amount'], labels=ageing['Age Bracket'], autopct='%.1f%%')
        age_pie.axis('off')
    else:
        age_tbl.text(s='Zero Balance', x=0.5, y=0.5, ha='center', va='center', fontsize=28)
        age_tbl.axis('off')
        age_pie.text(s='Zero Balance', x=0.5, y=0.5, ha='center', va='center', fontsize=28)
        age_pie.axis('off')
    rev_tbl.set_title('Monthly Sales', loc='left', **heading_format)
    rev_tbl.table(cellText=[[i[0].strftime('%B'), f'{i[1]:,.0f}'] for i in monthly_rev.values],
                  colLabels=monthly_rev.columns, cellLoc='center', loc='center')
    rev_tbl.axis('off')
    rev_bar.bar([i.strftime('%b') for i in monthly_rev['Month']], monthly_rev['Amount'])
    rev_bar.set_yticklabels(['{:,}'.format(int(i)) for i in rev_bar.get_yticks()])

    buf = BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format='png')
    plt.close(fig)
    buf.seek(0)
    doc.add_picture(buf)

    document.add_page_break()

salesperson_list: list = fInvoices.loc[(fInvoices['Invoice_Date'] <= end_date) & (
        fInvoices['Invoice_Date'] >= datetime(year=end_date.year, month=1, day=1)), 'Employee_Code'].unique()

salesperson_stats: dict = sales_person(emp_ids=salesperson_list, dEmployee=dEmployee, fInvoices=fInvoices)

cy_cp_pl_company_title = document.add_paragraph().add_run('Elite Security Services W.L.L')
apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
salesman_info_toc = document.add_paragraph().add_run('Key data about sales person')
apply_style_properties(salesman_info_toc, style_picker(name='report_title'))
tbl_salesman_toc = document.add_table(rows=1, cols=2)
heading_cells = tbl_salesman_toc.rows[0].cells
heading_cells[0].text = 'Salesperson Name'
heading_cells[1].text = 'Page #'

for idx, row in enumerate(salesperson_list):
    cells = tbl_salesman_toc.add_row().cells
    cells[0].text = ' '.join(dEmployee.loc[row, 'Employee_Name'].split(sep=' ')[:2]).title()
    cells[1].text = str(idx + 1)

table_formatter(table_name=tbl_salesman_toc, style_name='table_style_1', special=[])
document.add_page_break()

for idx, salesperson in enumerate(salesperson_list):
    if (idx + 1) % 2 == 0:
        document.add_paragraph('\n\n\n')
    salesperson_name: str = ' '.join(dEmployee.loc[salesperson, 'Employee_Name'].split(sep=' ')[:2]).title()
    salutation: str = "Mr." if dEmployee.loc[salesperson, 'Gender'] == 'Male' else "Ms."
    full_name: str = f'{salutation}{salesperson_name}'
    cy_cp_pl_company_title = document.add_paragraph().add_run(full_name)
    apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
    tbl_salesman_main_1 = document.add_table(rows=2, cols=2)
    tbl_salesman_main_th_1 = tbl_salesman_main_1.rows[0]
    tbl_salesman_main_th_1.cells[0].text = 'Date of Join'
    tbl_salesman_main_th_1.cells[1].text = 'New Customers Added'

    tbl_salesman_main_td_1 = tbl_salesman_main_1.rows[1]
    tbl_salesman_main_td_1.cells[0].text = str(salesperson_stats[salesperson]['doj'])
    tbl_salesman_main_td_1.cells[1].text = number_format(num=salesperson_stats[salesperson]['new_customers_added'])
    table_formatter(table_name=tbl_salesman_main_1, style_name='table_style_1', special=[])

    tbl_salesman_main_2 = document.add_table(rows=2, cols=2)
    tbl_salesman_main_th_2 = tbl_salesman_main_2.rows[0]
    tbl_salesman_main_th_2.cells[0].text = 'CP Target'
    tbl_salesman_main_th_2.cells[1].text = 'YTD Target'

    tbl_salesman_main_td_2 = tbl_salesman_main_2.rows[1]
    tbl_salesman_main_td_2.cells[0].text = number_format(num=salesperson_stats[salesperson]['cp_target'])
    tbl_salesman_main_td_2.cells[1].text = number_format(num=salesperson_stats[salesperson]['ytd_target'])
    table_formatter(table_name=tbl_salesman_main_2, style_name='table_style_1', special=[])

    tbl_salesman_rev_1 = document.add_table(rows=2, cols=2)
    tbl_salesman_rev_th_1 = tbl_salesman_rev_1.rows[0]
    tbl_salesman_rev_th_1.cells[0].text = 'CY CP Revenue'
    tbl_salesman_rev_th_1.cells[1].text = 'CY YTD Revenue'

    tbl_salesman_rev_td_1 = tbl_salesman_rev_1.rows[1]
    tbl_salesman_rev_td_1.cells[0].text = number_format(salesperson_stats[salesperson]['cy_cp_rev'])
    tbl_salesman_rev_td_1.cells[1].text = number_format(salesperson_stats[salesperson]['cy_ytd_rev'])
    table_formatter(table_name=tbl_salesman_rev_1, style_name='table_style_1', special=[])

    tbl_salesman_rev_2 = document.add_table(rows=2, cols=2)
    tbl_salesman_rev_th_2 = tbl_salesman_rev_2.rows[0]
    tbl_salesman_rev_th_2.cells[0].text = 'CY CP Own\nRevenue'
    tbl_salesman_rev_th_2.cells[1].text = 'CY YTD Own\nRevenue'

    tbl_salesman_rev_td_2 = tbl_salesman_rev_2.rows[1]
    tbl_salesman_rev_td_2.cells[0].text = number_format(num=salesperson_stats[salesperson]['cy_cp_rev_org'])
    tbl_salesman_rev_td_2.cells[1].text = number_format(salesperson_stats[salesperson]['cy_ytd_rev_org'])
    table_formatter(table_name=tbl_salesman_rev_2, style_name='table_style_1', special=[])

    tbl_salesman_gp_1 = document.add_table(rows=2, cols=2)
    tbl_salesman_gp_th_1 = tbl_salesman_gp_1.rows[0]
    tbl_salesman_gp_th_1.cells[0].text = 'CY CP GP'
    tbl_salesman_gp_th_1.cells[1].text = 'CY YTD GP'

    cy_cp_gp: float = cy_cp_profit_emp.loc[cy_cp_profit_emp['Employee_Code'] == salesperson, 'Amount'].sum()
    cy_ytd_gp: float = cy_ytd_profit_emp.loc[cy_ytd_profit_emp['Employee_Code'] == salesperson, 'Amount'].sum()

    tbl_salesman_gp_td_1 = tbl_salesman_gp_1.rows[1]
    tbl_salesman_gp_td_1.cells[0].text = number_format(num=cy_cp_gp)
    tbl_salesman_gp_td_1.cells[1].text = number_format(num=cy_ytd_gp)
    table_formatter(table_name=tbl_salesman_gp_1, style_name='table_style_1', special=[])

    tbl_salesman_gp_2 = document.add_table(rows=2, cols=2)
    tbl_salesman_gp_th_2 = tbl_salesman_gp_2.rows[0]
    tbl_salesman_gp_th_2.cells[0].text = 'CY CP Revenue\nContribution'
    tbl_salesman_gp_th_2.cells[1].text = 'CY YTD Revenue\nContribution'

    tbl_salesman_gp_td_2 = tbl_salesman_gp_2.rows[1]
    tbl_salesman_gp_td_2.cells[0].text = str(salesperson_stats[salesperson]['cy_cp_rev_contrib_pct'])
    tbl_salesman_gp_td_2.cells[1].text = str(salesperson_stats[salesperson]['cy_ytd_rev_contrib_pct'])
    table_formatter(table_name=tbl_salesman_gp_2, style_name='table_style_1', special=[])
    if (idx + 1) % 2 == 0:
        document.add_page_break()

cp_in_guard_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='month', div='guarding',
                                            type='Related', cnt=5)
cp_in_elv_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='month', div='elv',
                                          type='Related', cnt=5)
cp_ex_guard_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='month', div='guarding',
                                            type='Market', cnt=5)
cp_ex_elv_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='month', div='elv',
                                          type='Market', cnt=5)
ytd_in_guard_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='ytd', div='guarding',
                                             type='Related', cnt=5)
ytd_in_elv_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='ytd', div='elv',
                                           type='Related', cnt=5)
ytd_ex_guard_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='ytd', div='guarding',
                                             type='Market', cnt=5)
ytd_ex_elv_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='ytd', div='elv', type='Market',
                                           cnt=5)
cus_info_1 = {0: {0: {'name': 'Current Month Internal Guarding', 'df': cp_in_guard_df},
                  1: {'name': 'Current Month Internal ELV', 'df': cp_in_elv_df}},
              1: {0: {'name': 'Current Month External Guarding', 'df': cp_ex_guard_df},
                  1: {'name': 'Current Month External ELV', 'df': cp_ex_elv_df}},
              2: {0: {'name': 'Year to Date Internal Guarding', 'df': ytd_in_guard_df},
                  1: {'name': 'Year to Date Internal ELV', 'df': ytd_in_guard_df}},
              3: {0: {'name': 'Year to Date External Guarding', 'df': ytd_ex_elv_df},
                  1: {'name': 'Year to Date External ELV', 'df': ytd_ex_elv_df}}}
rows_report_1: int = 4
cols_report_1: int = 2
keydatacus1 = document.add_table(rows=rows_report_1, cols=cols_report_1)

for row in range(rows_report_1):
    for col in range(cols_report_1):
        row_0 = keydatacus1.rows[row].cells
        row_0[col].text = cus_info_1[row][col]['name']
        df: pd.DataFrame = cus_info_1[row][col]['df']
        inner_tbl_rows = df.shape[0] + 1
        inner_tbl_cols = df.shape[1]
        inner_tbl = row_0[col].add_table(rows=1, cols=2)
        inner_tbl_hdr = inner_tbl.rows[0].cells
        inner_tbl_hdr[0].text = 'Customer'
        inner_tbl_hdr[1].text = 'Amount'
        for _, j in df.iterrows():
            cells = inner_tbl.add_row().cells
            cells[0].text = str(j['Customer'])
            cells[1].text = number_format(j.iloc[1])
        table_formatter(table_name=inner_tbl, style_name='table_style_1', special=['Total'])
document.add_page_break()

inc_pp: pd.DataFrame = revenue_change(fInvoices=fInvoices, end_date=end_date, mode='cypp', order=True)
dec_pp: pd.DataFrame = revenue_change(fInvoices=fInvoices, end_date=end_date, mode='cypp', order=False)
inc_py: pd.DataFrame = revenue_change(fInvoices=fInvoices, end_date=end_date, mode='pycp', order=True)
dec_py: pd.DataFrame = revenue_change(fInvoices=fInvoices, end_date=end_date, mode='pycp', order=False)

cus_info_2 = {0: {0: {'name': 'Top 5 Customers with Incresed\nRevenue compared to previous month', 'df': inc_pp},
                  1: {'name': 'Top 5 Customers with Decreased\nRevenue compared to previous month', 'df': dec_pp}},
              1: {0: {'name': 'Top 5 Customers with Increased\nRevenue compared to previous year', 'df': inc_py},
                  1: {'name': 'Top 5 Customers with Decreased\nRevenue compared to previous year', 'df': dec_py}}}
rows_report_2: int = 2
cols_report_2: int = 2
keydatacus2 = document.add_table(rows=rows_report_1, cols=cols_report_1)  # r4,c2

for row in range(rows_report_2):
    for col in range(cols_report_2):
        row_0 = keydatacus2.rows[row].cells
        row_0[col].text = cus_info_2[row][col]['name']
        df: pd.DataFrame = cus_info_2[row][col]['df']
        inner_tbl_rows = df.shape[0] + 1
        inner_tbl_cols = df.shape[1]
        inner_tbl = row_0[col].add_table(rows=1, cols=2)
        inner_tbl_hdr = inner_tbl.rows[0].cells
        inner_tbl_hdr[0].text = 'Customer'
        inner_tbl_hdr[1].text = 'Amount'
        for _, j in df.iterrows():
            cells = inner_tbl.add_row().cells
            cells[0].text = str(j['Customer'])
            cells[1].text = number_format(j.iloc[1])
        table_formatter(table_name=inner_tbl, style_name='table_style_1', special=['Total'])
document.add_page_break()

page_separator(head='HR', document=document)

emp_data: dict = employee_related(data=dEmployee)
plt.style.use('ggplot')
hr_fig_1, ((gender, type), (dept, nationality)) = plt.subplots(nrows=2, ncols=2, figsize=(10.5, 7))

gender.set_title('Gender')
gender.pie(x=list(emp_data['Gender'].values()), labels=list(emp_data['Gender'].keys()), autopct='%.0f%%',
           labeldistance=1, pctdistance=0.3)
gender.axis('off')

type.set_title('Category')
type.pie(x=list(emp_data['Type'].values()), labels=list(emp_data['Type'].keys()), autopct='%.0f%%', labeldistance=1,
         pctdistance=0.3)
type.axis('off')

dept.set_title('Department')
dept.pie(x=list(emp_data['Department'].values()), labels=list(emp_data['Department'].keys()), autopct='%.0f%%',
         labeldistance=1, pctdistance=0.3)
dept.axis('off')

nationality.set_title('Nationality')
nationality.pie(x=list(emp_data['Nationality'].values()), labels=list(emp_data['Nationality'].keys()), autopct='%.0f%%',
                labeldistance=1, pctdistance=0.3)
nationality.axis('off')

change_orientation(doc=document, method='l')
hr_graph_1_buf = BytesIO()
plt.tight_layout()
plt.savefig(hr_graph_1_buf, format='png')
plt.close(hr_fig_1)
hr_graph_1_buf.seek(0)
doc.add_picture(hr_graph_1_buf)
document.add_page_break()

hr_fig_2, ((age, service), (designation, movement)) = plt.subplots(nrows=2, ncols=2, figsize=(10.5, 7))
age.set_title('Age')
age.pie(x=list(emp_data['Employee Age'].values()), labels=list(emp_data['Employee Age'].keys()), autopct='%.0f%%',
        labeldistance=1, pctdistance=0.3)
age.axis('off')

service.set_title('Service')
service.pie(x=list(emp_data['Service Period'].values()), labels=list(emp_data['Service Period'].keys()),
            autopct='%.0f%%', labeldistance=1, pctdistance=0.3)
service.axis('off')

designation.set_title('Designation')
designation.pie(x=list(emp_data['Designation'].values()), labels=list(emp_data['Designation'].keys()), autopct='%.0f%%',
                labeldistance=1, pctdistance=0.7)
designation.axis('off')

emp_move = pd.DataFrame(list(emp_data['Employee Movement'].items()), columns=['Description', '# of Emp'])

movement.set_title('Movement')
movement.table(cellText=[i for i in emp_move.values], colLabels=emp_move.columns, cellLoc='center', loc='center')
movement.axis('off')

hr_graph_2_buf = BytesIO()
plt.tight_layout()
plt.savefig(hr_graph_2_buf, format='png')
plt.close(hr_fig_2)
hr_graph_2_buf.seek(0)
doc.add_picture(hr_graph_2_buf)
document.add_page_break()

hr_fig_3, (jo_re, total_staff) = plt.subplots(nrows=2, ncols=1, sharex=True,figsize = (7.27,10))

jo_re.set_title('New Joiners and Leavers')
jo_re.plot([i.strftime('%b') for i in emp_data['new_joiner'].index], emp_data['new_joiner']['Joined'],
           label='New Joiners')
jo_re.plot([i.strftime('%b') for i in emp_data['new_joiner'].index], emp_data['resigned_emp']['Resigned'],
           label='Resigned')
jo_re.legend()

total_staff.set_title('Total Manpower')
total_staff.plot([i.strftime('%b') for i in emp_data['new_joiner'].index],
                 emp_data['total_employees']['Total Employees'], label='Total Employees')

change_orientation(doc=document, method='p')
hr_graph_3_buf = BytesIO()
plt.tight_layout()
plt.savefig(hr_graph_3_buf, format='png')
plt.close(hr_fig_3)
hr_graph_3_buf.seek(0)
doc.add_picture(hr_graph_3_buf)
document.add_page_break()

ops_data: pd.DataFrame = operations(ftimesheet=fTimesheet, financial=cy_ytd_basic_monthwise, end_date=end_date)

page_separator(head='Operations', document=document)

plt.style.use('ggplot')
fig_ops_1, (cost_line, ph_line) = plt.subplots(nrows=2, ncols=1, sharex=True,figsize = (7.27,10))

cost_line.set_title('Transportation and Accommodation Expenses')
cost_line.plot([i.strftime('%b') for i in ops_data.index], ops_data['Transport'], label='Transport')
cost_line.plot([i.strftime('%b') for i in ops_data.index], ops_data['Accommodation'], label='Accommodation')
cost_line.set_yticklabels(['{:,}'.format(int(i)) for i in cost_line.get_yticks()])
cost_line.legend()

ph_line.set_title('Transportation and Accommodation Per Head')
ph_line.plot([i.strftime('%b') for i in ops_data.index], (ops_data['Transport'] / ops_data['trpt_md']) * 30,
             label='Transport')
ph_line.plot([i.strftime('%b') for i in ops_data.index], (ops_data['Accommodation'] / ops_data['acco_md']) * 30,
             label='Accommodation')
ph_line.set_yticklabels(['{:,}'.format(int(i)) for i in ph_line.get_yticks()])
ph_line.legend()

ops_graph_1_buf = BytesIO()
plt.tight_layout()
plt.savefig(ops_graph_1_buf, format='png')
plt.close(fig_ops_1)
ops_graph_1_buf.seek(0)
doc.add_picture(ops_graph_1_buf)
document.add_page_break()

fig_ops_2, (bill_nonbil, efficiency,non_billable) = plt.subplots(nrows=3, ncols=1, figsize = (7.73,10.63), sharex=True,gridspec_kw={'height_ratios': [1,1,2]})

bill_nonbil.set_title('Billable Vs Non-Billable Mandays')
bill_nonbil.plot([i.strftime('%b') for i in ops_data.index], ops_data['productive_md'], label='Productive')
bill_nonbil.plot([i.strftime('%b') for i in ops_data.index], ops_data['unproductive_md'], label='Un-productive')
bill_nonbil.set_yticklabels(['{:,}'.format(int(i)) for i in bill_nonbil.get_yticks()])
bill_nonbil.legend()

efficiency.set_title('Manpower Utilization Efficiency')
efficiency.plot([i.strftime('%b') for i in ops_data.index], (ops_data['productive_md'] / ops_data['acco_md']) * 100,
                label='Efficiency')
efficiency.set_yticklabels(['{:,.0f}%'.format(i) for i in efficiency.get_yticks()])
efficiency.legend()


periods = pd.date_range(start=start_date, end=end_date, freq='ME').to_pydatetime().tolist()

c = {}
exclude_dict = dExclude.groupby('group')['job_type'].apply(set).to_dict()
for t in periods:
    period_allocation = profitability['periodic_allocation'].get(t, {})
    a = {}    
    for group, job_types in exclude_dict.items():
        for job_type in job_types:
            if job_type in period_allocation:
                a[group] = a.get(group, 0) + period_allocation[job_type]
    a = {k: v for k, v in a.items() if v != 0}
    c[t] = a
results_df = pd.DataFrame.from_dict(c, orient='index').fillna(0) * -1

non_billable.set_title('Non-Billable Cost')
for p in results_df.columns:
    non_billable.plot([i.strftime('%b') for i in results_df.index], results_df[p], label=p)
non_billable.set_yticklabels(['{:,}'.format(int(i)) for i in non_billable.get_yticks()])
non_billable.legend()


ops_graph_2_buf = BytesIO()
plt.tight_layout()
plt.savefig(ops_graph_2_buf, format='png')
plt.close(fig_ops_2)
ops_graph_2_buf.seek(0)
doc.add_picture(ops_graph_2_buf)
document.add_page_break()

credit = document.add_paragraph(
    '\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\nNadun Jayathunga\n')
credit.add_run('Chief Accountant\nNasser Bin Nawaf & Partners Holding W.L.L\n')
credit.add_run('mail:njayathunga@nbn.qa\nTel:+974 4403 0407').italic = True

document.core_properties.author = "Nadun Jayathunga"
document.core_properties.keywords = ("Chief Accountant\nNasser Bin Nawaf and Partners Holdings "
                                     "W.L.L\nmail:njayathunga@nbn.qa\nTele:+974 4403 0407")

doc.save('Monthly FS.docx')
convert('Monthly FS.docx')
os.unlink('Monthly FS.docx')

# TODO COHART to show
