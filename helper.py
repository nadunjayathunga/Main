import sys

sys.path.append(r'C:\Users\NadunJayathunga\OneDrive - NBN Holdings\Financials\Other\Programmes\Dashboards\Main')
import pandas as pd
from dateutil.relativedelta import relativedelta
from datetime import datetime, timedelta
from sqlalchemy import create_engine
from colorama import Fore, init
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.ticker import FixedLocator, FixedFormatter
from matplotlib.gridspec import GridSpec
from io import BytesIO
import itertools
from itertools import islice
from docx2pdf import convert
import os
import statistics

from data import company_info, db_info, doc_styles, table_style, SYSTEM_CUT_OFF, company_data, cogs_ledger_map, \
    VOUCHER_TYPES

init()


def welcome_page() -> dict:
    print('\n')
    for idx, company in enumerate(company_info):
        print(f'{idx}\t{company["data"]["long_name"]}')
    print('\nPress Q to Quit')
    proceed: bool = True
    no_of_companies: int = len(company_info)
    msg: str = f'Enter a Digit from 0 - {no_of_companies - 1} or Q to Quit'
    while proceed:
        try:
            user_input: str = input('Please enter company ID>> ').lower()
            if user_input == 'q':
                break
            else:
                user_input: int = int(user_input)
            if user_input < 0 or user_input > no_of_companies - 1:
                print(msg)
            else:
                company_id: int = user_input
                print('\n')
                database: str = company_info[company_id]['data']['database']
                abbr: str = company_info[company_id]['data']['abbr']
                long_name: str = company_info[company_id]['data']['long_name']
                rev_cats: list = company_info[company_id]['data']['rev_cat']
                proceed = False

        except ValueError:
            print(msg)
    engine = create_engine(
        f'postgresql://{db_info["USERNAME"]}:{db_info["PWD"]}@{db_info["HOSTNAME"]}:{db_info["PORT_ID"]}/{database}')

    gl_start: datetime = pd.read_sql_query('SELECT MIN(voucher_date) FROM "fGL"', con=engine).squeeze()
    gl_end: datetime = pd.read_sql_query('SELECT MAX(voucher_date) FROM "fGL"', con=engine).squeeze()

    while not proceed:
        user_input = input("Enter a date (YYYY-MM-DD): ")
        try:
            end_date = datetime.strptime(user_input, '%Y-%m-%d')
            if gl_start <= end_date <= gl_end:
                proceed = True
            else:
                print(f"{Fore.RED}Date must be between {gl_start.date()} and {gl_end.date()}.{Fore.RESET}")
        except ValueError:
            print(f"{Fore.RED}Invalid date format. Please use 'YYYY-MM-DD'.{Fore.RESET}")

    return {'database': database, 'engine': engine, 'end_date': end_date, 'abbr': abbr, 'long_name': long_name,
            'rev_cats': rev_cats}


def data_sources(engine, database: str) -> dict:
    fGL: pd.DataFrame = pd.read_sql_query(sql=f'SELECT * FROM merged', con=engine)
    fInvoices: pd.DataFrame = pd.read_sql_query(sql=f'SELECT * FROM finvoices', con=engine)
    dJobs: pd.DataFrame = pd.read_sql_query(sql=f'SELECT * FROM djobs', con=engine)
    dEmployee: pd.DataFrame = pd.read_sql_table(table_name='dEmployee', con=engine, )
    dCoAAdler: pd.DataFrame = pd.read_sql_table(table_name='dCoAAdler', con=engine)
    dCustomer: pd.DataFrame = pd.read_sql_table(table_name='dCustomer', con=engine)
    fCreditNote: pd.DataFrame = pd.read_sql_table(table_name='fCreditNote', con=engine)
    fBudget: pd.DataFrame = pd.read_sql_table(table_name='fBudget', con=engine)
    fCollection: pd.DataFrame = pd.read_sql_table(table_name='fCollection', con=engine)
    fAP: pd.DataFrame = pd.read_sql_table(table_name='fAP', con=engine)
    common: dict = {'fGL': fGL, 'dEmployee': dEmployee, 'dCoAAdler': dCoAAdler, 'fCreditNote': fCreditNote,
                    'dCustomer': dCustomer, 'fBudget': fBudget, 'fCollection': fCollection, 'fAP': fAP,
                    'fInvoices': fInvoices, 'dJobs': dJobs}
    if database == 'elite_security':
        ftimesheet: pd.DataFrame = pd.read_sql_table(table_name='ftimesheet', con=engine)
        fOT: pd.DataFrame = pd.read_sql_table(table_name='fOT', con=engine)
        dExclude: pd.DataFrame = pd.read_sql_table(table_name='dExclude', con=engine)
        ess_specific: dict = {'ftimesheet': ftimesheet, 'fOT': fOT, 'dExclude': dExclude}
        common = common | ess_specific
    elif database == 'nbn_logistics':
        fLogInv: pd.DataFrame = pd.read_sql_table(table_name='fLogInv', con=engine)
        dLogContract: pd.DataFrame = pd.read_sql_table(table_name='dLogContract', con=engine)
        fData: pd.DataFrame = pd.read_sql_table(table_name='fData', con=engine)
        fNotInvoiced: pd.DataFrame = pd.read_sql_table(table_name='fNotInvoiced', con=engine)
        fSalesTill2020: pd.DataFrame = pd.read_sql_table(table_name='fSalesTill2020', con=engine)
        nbnl_specific: dict = {'fLogInv': fLogInv, 'dLogContract': dLogContract, 'fData': fData,
                               'fNotInvoiced': fNotInvoiced, 'fSalesTill2020': fSalesTill2020}
        common = common | nbnl_specific
    elif database == 'nbn_realestate':
        nbnre_specific: dict = {}
        common = common | nbnre_specific
    else:
        ph_specific: dict = {}
        common = common | ph_specific
    return common


def business_unit(row, dEmployee: pd.DataFrame, dCoAAdler: pd.DataFrame, database: str) -> str:
    if database == 'elite_security':
        elv_groups: list = dCoAAdler.loc[
            dCoAAdler['first_level'].isin(['Material Parts & Consumables - Projects',
                                           'Maintenance - Projects',
                                           'Others - Projects', 'Projects Revenue']), 'ledger_code'].tolist()

        ledger_code: str = row['ledger_code']
        cc: str = row['cost_center']

        if ledger_code in elv_groups:
            return 'ELV-ESS'
        if pd.isna(cc) or cc == '':
            return 'GUARDING-ESS'
        dept = dEmployee.loc[dEmployee['emp_id'] == cc, 'dept'].squeeze() if cc in dEmployee[
            'emp_id'].tolist() else 'GUARDING-ESS'
        return 'ELV-ESS' if dept == 'ELV' else 'GUARDING-ESS'


def receipts_recorded(data: pd.DataFrame, database: str) -> pd.DataFrame:
    # data['invoice_date'] = pd.to_datetime(data['invoice_date'], errors='coerce')
    # Voucher date is null mean the subject invoice has not been paid at all
    if database == 'nbn_logistics':
        data: pd.DataFrame = data.loc[~data['invoice_number'].str.contains('NBL/JV200072|NBL/JV200073')]
        correction = [{'ledger_code': 1020404001, 'invoice_number': 'NBL/JV200073-1', 'invoice_date': '12/31/2020 0:00',
                       'invoice_amount': 5667, 'voucher_number': 'NBL/RV210074-5667.00',
                       'voucher_date': pd.Timestamp('2021-03-01 00:00:00')},
                      {'ledger_code': 1020404002, 'invoice_number': 'NBL/JV200073-2', 'invoice_date': '12/31/2020 0:00',
                       'invoice_amount': 35417, 'voucher_number': 'NBL/RV210244-35417.00',
                       'voucher_date': pd.Timestamp('2021-08-16 00:00:00')},
                      {'ledger_code': 1020404007, 'invoice_number': 'NBL/JV200073-4', 'invoice_date': '12/31/2020 0:00',
                       'invoice_amount': 11333, 'voucher_number': 'NBL/RV210064-11333.00',
                       'voucher_date': pd.Timestamp('2021-02-15 00:00:00')},
                      {'ledger_code': 1020404012, 'invoice_number': 'NBL/JV200073-5', 'invoice_date': '12/31/2020 0:00',
                       'invoice_amount': 1417, 'voucher_number': 'NBL/RV210068-1417.00',
                       'voucher_date': pd.Timestamp('2021-02-15 00:00:00')},
                      {'ledger_code': 1020404010, 'invoice_number': 'NBL/JV200073-7', 'invoice_date': '12/31/2020 0:00',
                       'invoice_amount': 36833, 'voucher_number': 'NBL/RV210072-36833.00',
                       'voucher_date': pd.Timestamp('2021-03-01 00:00:00')},
                      {'ledger_code': 1020406003, 'invoice_number': 'NBL/JV200073-8', 'invoice_date': '12/31/2020 0:00',
                       'invoice_amount': 8334, 'voucher_number': 'NBL/RV210500-8334.00',
                       'voucher_date': pd.Timestamp('2021-06-01 00:00:00')},
                      {'ledger_code': 1020405002, 'invoice_number': 'NBL/JV200073-9', 'invoice_date': '12/31/2020 0:00',
                       'invoice_amount': 25500, 'voucher_number': 'NBL/RV210479-25500.00',
                       'voucher_date': pd.Timestamp('2021-03-31 00:00:00')},
                      {'ledger_code': 1020404001, 'invoice_number': 'NBL/JV200072-1', 'invoice_date': '11/30/2020 0:00',
                       'invoice_amount': 5667, 'voucher_number': 'NBL/RV210074-5667.00',
                       'voucher_date': pd.Timestamp('2021-03-01 00:00:00')},
                      {'ledger_code': 1020404002, 'invoice_number': 'NBL/JV200072-2', 'invoice_date': '11/30/2020 0:00',
                       'invoice_amount': 35417, 'voucher_number': 'NBL/RV200105-35417.00',
                       'voucher_date': pd.Timestamp('2020-12-20 00:00:00')},
                      {'ledger_code': 1020404012, 'invoice_number': 'NBL/JV200072-4', 'invoice_date': '11/30/2020 0:00',
                       'invoice_amount': 1417, 'voucher_number': 'NBL/RV210068-1417.00',
                       'voucher_date': pd.Timestamp('2021-02-15 00:00:00')},
                      {'ledger_code': 1020404007, 'invoice_number': 'NBL/JV200072-5', 'invoice_date': '11/30/2020 0:00',
                       'invoice_amount': 11333, 'voucher_number': 'NBL/RV210064-11333.00',
                       'voucher_date': pd.Timestamp('2021-02-15 00:00:00')},
                      {'ledger_code': 1020404010, 'invoice_number': 'NBL/JV200072-7', 'invoice_date': '11/30/2020 0:00',
                       'invoice_amount': 36833, 'voucher_number': 'NBL/RV210072-36833.00',
                       'voucher_date': pd.Timestamp('2021-03-01 00:00:00')},
                      {'ledger_code': 1020406003, 'invoice_number': 'NBL/JV200072-8', 'invoice_date': '11/30/2020 0:00',
                       'invoice_amount': 8334, 'voucher_number': 'NBL/RV210500-8334.00',
                       'voucher_date': pd.Timestamp('2021-06-01 00:00:00')},
                      {'ledger_code': 1020405002, 'invoice_number': 'NBL/JV200072-9', 'invoice_date': '11/30/2020 0:00',
                       'invoice_amount': 25500, 'voucher_number': 'NBL/RV200094-25500.00',
                       'voucher_date': pd.Timestamp('2020-12-31 00:00:00')},
                      {'ledger_code': 1020404003, 'invoice_number': 'NBL/JV200073-3', 'invoice_date': '12/31/2020 0:00',
                       'invoice_amount': 11333, 'voucher_number': 'NBL/JV210461-4333.00;NBL/RV220211-7000.00',
                       'voucher_date': '31-Jan-2021,18-Apr-2022'},
                      {'ledger_code': 1020404009, 'invoice_number': 'NBL/JV200073-6', 'invoice_date': '12/31/2020 0:00',
                       'invoice_amount': 14617, 'voucher_number': 'NBL/RV210464-3617.00;NBL/RV210462-11000.00',
                       'voucher_date': '31-Jul-2021,06-Oct-2021'},
                      {'ledger_code': 1020404003, 'invoice_number': 'NBL/JV200072-3', 'invoice_date': '11/30/2020 0:00',
                       'invoice_amount': 11333, 'voucher_number': 'NBL/JV210461-4333.00;NBL/RV220211-7000.00',
                       'voucher_date': '31-Jan-2021,18-Apr-2022'},
                      {'ledger_code': 1020404009, 'invoice_number': 'NBL/JV200072-6', 'invoice_date': '11/30/2020 0:00',
                       'invoice_amount': 14617, 'voucher_number': 'NBL/RV210464-3617.00;NBL/RV210462-11000.00',
                       'voucher_date': '31-Jul-2021,06-Oct-2021'}]
        correction = pd.DataFrame(data=correction)
        data = pd.concat([data, correction])
    nulldf: pd.DataFrame = data.loc[data['voucher_date'].isna()]
    # certain invoices have been settled in multiple receipts. mode of settlement can be receipt/ credit note or even a JV
    multidates: pd.DataFrame = data.loc[~data['voucher_date'].isna() & data['voucher_date'].str.contains(pat=',')]
    # This will convert voucher_date column to datetime format, those entries which does not follow the format will be returned as null
    data['voucher_date'] = pd.to_datetime(data['voucher_date'], errors='coerce', format='%Y-%m-%d %H:%M:%S')
    # This will filter entries which has been settled only once. (fully or partially)
    singledate: pd.DataFrame = data.loc[~data['voucher_date'].isna()]
    # For those entries which has only one receipt will look like below ESS/DN230001-2650.00. Split voucher_number into 2 parts convert them to 
    # columns and assign them into voucher_number nad voucher_amount
    singledate[['voucher_number', 'voucher_amount']] = singledate['voucher_number'].apply(
        lambda name: pd.Series(name.split("-", 1)))
    final_collection_df: pd.DataFrame = pd.DataFrame(columns=[
        'invoice_number', 'ledger_code', 'invoice_date', 'invoice_amount',
        'voucher_number', 'voucher_amount', 'voucher_date'])

    for _, row in multidates.iterrows():
        pv_number = row['voucher_number']
        #  for each entry which has multiple receipts, this will grab the voucher_number for each individual
        #  settlement ESS/CN200006-5370.62;ESS/RV210579-22094.62 -> ['ESS/CN200006','ESS/RV210579']
        voucher_number = [voucher.split(sep='-')[0] for voucher in
                          pv_number.split(sep=';')]
        #  for each entry which has multiple receipts, this will grab the voucher_amount for each individual
        #  settlement ESS/CN200006-5370.62;ESS/RV210579-22094.62 -> [5370.62,22094.62]
        voucher_amount = [float(voucher.split(sep='-')[1]) for voucher in
                          pv_number.split(sep=';')]
        # 31-Dec-2020,24-Feb-2021 -> [31-Dec-2020,24-Feb-2021]
        voucher_date = [datetime.strptime(i, '%d-%b-%Y') for i in row['voucher_date'].split(sep=',')]
        # invoice_number will be repeated for no of times a invoice has been settled. 
        invoice_number = [row['invoice_number'] for _ in range(len(voucher_number))]
        ledger_code = [row['ledger_code'] for _ in range(len(voucher_number))]
        invoice_date = [row['invoice_date'] for _ in range(len(voucher_number))]
        invoice_amount = [row['invoice_amount'] for _ in range(len(voucher_number))]
        collection_df: pd.DataFrame = pd.DataFrame(
            data={'invoice_number': invoice_number, 'ledger_code': ledger_code, 'invoice_date': invoice_date,
                  'invoice_amount': invoice_amount,
                  'voucher_number': voucher_number, 'voucher_amount': voucher_amount, 'voucher_date': voucher_date})
        final_collection_df = pd.concat(i for i in [final_collection_df, collection_df] if not i.empty)
    final_collection_df = pd.concat(
        i for i in [final_collection_df, nulldf, singledate] if not i.empty and not i.isnull().all().all())
    final_collection_df.loc[:, 'voucher_date'] = pd.to_datetime(final_collection_df['voucher_date'], errors='coerce')
    final_collection_df.loc[:, 'voucher_amount'] = final_collection_df['voucher_amount'].fillna(value=0)
    final_collection_df.loc[:, 'voucher_amount'] = pd.to_numeric(final_collection_df['voucher_amount'])
    final_collection_df['invoice_date'] = pd.to_datetime(final_collection_df['invoice_date'])
    return final_collection_df


def preprocessing(data: dict, database: str) -> dict:
    fGL: pd.DataFrame = data['fGL']
    dEmployee: pd.DataFrame = data['dEmployee']
    dCoAAdler: pd.DataFrame = data['dCoAAdler']
    dCustomer: pd.DataFrame = data['dCustomer']
    fBudget: pd.DataFrame = data['fBudget']
    fCollection: pd.DataFrame = data['fCollection']
    fCreditNote: pd.DataFrame = data['fCreditNote']
    dJobs: pd.DataFrame = data['dJobs']
    if database == 'elite_security':
        fGL['business_unit_name'] = fGL.apply(
            business_unit, axis=1, args=[dEmployee, dCoAAdler, database])
    fGL['narration'] = fGL['narration'].fillna('')
    fGL['first_level'] = fGL['first_level'].replace(to_replace='Depreciation - Projects', value='Depreciation')
    fCreditNote['amount'] = fCreditNote['amount'] * -1
    fCreditNote = pd.merge(
        left=fCreditNote, right=dCustomer, on='ledger_code', how='left')
    fBudget = pd.melt(fBudget, id_vars=[
        'fy', 'ledger_code'], var_name='Month', value_name='amount')
    fBudget['voucher_date'] = fBudget.apply(
        lambda x: pd.to_datetime(f'{x["fy"]}-{x["Month"]}-01') + relativedelta(day=31), axis=1)
    fBudget.drop(columns=['fy', 'Month'], inplace=True)
    fBudget = fBudget.loc[fBudget['amount'] != 0]
    fBudget = pd.merge(left=fBudget, right=dCoAAdler,
                       on='ledger_code', how='left')
    fAP: pd.DataFrame = data['fAP']
    interco_lgr: set = set(itertools.chain.from_iterable(
        [dCoAAdler.loc[dCoAAdler['ledger_name'].isin(company_data[i]['names']), 'ledger_code'].tolist() for i in
         company_data]))
    fAP = fAP.loc[~fAP['ledger_code'].isin(interco_lgr)]
    fAP = fAP.dropna(how='any')
    fAP = pd.merge(left=fAP, right=dCoAAdler[['ledger_name', 'ledger_code']], on='ledger_code', how='left').drop(
        columns=['ledger_code'])
    fAP['amount'] = fAP['amount'].astype(float)
    fAP.loc[:, 'amount'] = fAP['amount'] * -1
    fAP = fAP.groupby(by=['ledger_name', 'bracket'], as_index=False)['amount'].sum()
    fAP = fAP.pivot_table(index='ledger_name', columns='bracket', values='amount')
    fAP.loc[:, 'Total'] = fAP.loc[:, '0-30':].sum(axis=1)
    fAP = fAP.loc[fAP['Total'] > 0]
    fAP.dropna(how='all', axis=1, inplace=True)
    fAP.fillna(0, inplace=True)
    fAP.reset_index(inplace=True)
    fInvoices: pd.DataFrame = data['fInvoices']
    for i in fAP.columns:
        if i != 'ledger_name':
            fAP.loc['Total', i] = fAP[i].sum()
    fAP.loc['Total', 'ledger_name'] = 'Total'
    if database == 'elite_security':
        fBudget['business_unit_name'] = 'GUARDING-ESS'
    elif database == 'nbn_logistics':
        fBudget['business_unit_name'] = 'NBN Logistics'
    elif database == 'premium':
        fBudget['business_unit_name'] = 'Premium Hospitality'
    else:
        fBudget['business_unit_name'] = 'NBN Real Estate'

    fBudget.loc[fBudget['forth_level'] == 'Expenses', 'amount'] *= -1
    fCollection = receipts_recorded(data=fCollection, database=database)
    dEmployee.loc[:, 'travel_cost'] = dEmployee['travel_cost'].fillna(0)
    dEmployee['termination_date'] = pd.to_datetime(dEmployee['termination_date'])
    common: dict = {'fGL': fGL, 'dEmployee': dEmployee, 'dCoAAdler': dCoAAdler, 'fBudget': fBudget,
                    'dCustomer': dCustomer, 'fCollection': fCollection, 'fAP': fAP, 'fInvoices': fInvoices,
                    'dJobs': dJobs}

    if database == 'elite_security':
        ftimesheet: pd.DataFrame = data['ftimesheet']
        fOT: pd.DataFrame = data['fOT']
        dExclude: pd.DataFrame = data['dExclude']
        fOT.loc[:, 'voucher_date'] = fOT['voucher_date'].str.split(' ', expand=True)[4].str.strip()
        fOT.loc[:, 'voucher_date'] = pd.to_datetime(fOT['voucher_date'], format='%b-%Y') + pd.offsets.MonthEnd(0)
        fOT.fillna(0, inplace=True)
        fOT = fOT.loc[fOT['amount'] != 0]
        fOT.loc[:, 'amount'] = fOT['amount'] * -1
        ftimesheet = ftimesheet.loc[~ftimesheet['order_id'].isin(['discharged', 'not_joined'])]
        ftimesheet.loc[:, 'v_date'] = ftimesheet['v_date'] + pd.offsets.MonthEnd(0)

        ess_specific: dict = {'ftimesheet': ftimesheet, 'fOT': fOT,
                              'dExclude': dExclude}
        common = common | ess_specific
    elif database == 'premium':
        ph_specific: dict = {}
        common = common | ph_specific
    elif database == 'nbn_logistics':
        fLogInv: pd.DataFrame = data['fLogInv']
        fLogInv['invoice_date'] = fLogInv['invoice_date'] + pd.offsets.MonthEnd(0)
        dLogContract: pd.DataFrame = data['dLogContract']
        fData: pd.DataFrame = data['fData']
        fNotInvoiced: pd.DataFrame = data['fNotInvoiced']
        fSalesTill2020: pd.DataFrame = data['fSalesTill2020']
        nbnl_specific: dict = {'fLogInv': fLogInv, 'dLogContract': dLogContract, 'fData': fData,
                               'fNotInvoiced': fNotInvoiced, 'fSalesTill2020': fSalesTill2020}
        common = common | nbnl_specific
    else:
        nbnre_specific: dict = {}
        common = common | nbnre_specific

    return common


def first_page(document, report_date: datetime, abbr: str, long_name: str):
    new_section = document.sections[-1]
    new_section.left_margin = Inches(0.4)
    new_section.right_margin = Inches(0.4)
    new_section.top_margin = Inches(0.3)
    new_section.bottom_margin = Inches(0.1)
    new_section.header_distance = Inches(0.1)
    new_section.footer_distance = Inches(0.1)
    logo = document.add_picture(
        f'C:\Masters\images\logo\{abbr}-logo.png')
    logo = document.paragraphs[-1]
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first = document.add_paragraph()
    first.add_run('\n\n\n')
    first_run = first.add_run(long_name.upper())
    first_run.bold = True
    first_run.font.size = Pt(24)
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER

    second = document.add_paragraph()
    second_run = second.add_run(
        f'For the period ended {report_date.strftime("%Y-%b-%d")}')
    second_run.font.size = Pt(24)
    second.alignment = WD_ALIGN_PARAGRAPH.CENTER

    third = document.add_paragraph()
    third_run = third.add_run('COMPREHENSIVE FINANCIAL STATEMENT ANALYSIS')
    third_run.font.size = Pt(24)
    third.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return document


def closing(document, abbr: str):
    document.save(f"{abbr}-Monthly FS.docx")
    convert(f"{abbr}-Monthly FS.docx")
    os.unlink(f"{abbr}-Monthly FS.docx")


def page_separator(head: str, document):
    text = document.add_paragraph()
    text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text = text.add_run(f'\n\n\n{head.upper()}')
    text.bold = True
    text.font.color.rgb = RGBColor(153, 37, 43)
    text.font.size = Pt(80)
    document.add_page_break()


def apply_style_properties(run, properties):
    if 'bold' in properties:
        run.bold = properties['bold']
    if 'size' in properties:
        run.font.size = Pt(properties['size'])
    if 'name' in properties:
        run.font.name = properties['name']
    if 'color' in properties:
        run.font.color.rgb = RGBColor(*properties['color'])


def style_picker(name: str) -> dict:
    return [i[name] for i in doc_styles if name in i][0]


def header(title: str, company: str, document):
    cy_cp_pl_company_title = document.add_paragraph().add_run(f'{company}')
    apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
    cy_cp_pl_report_title = document.add_paragraph().add_run(title)
    apply_style_properties(cy_cp_pl_report_title, style_picker(name='report_title'))


def number_format(num):
    if num == 0:
        return "-"
    elif num >= 0:
        return f'{num:,.0f}'
    else:
        return f'({abs(num):,.0f})'


def table_formatter(table_name, style_name: dict, special: list):
    # Set the table style
    table_name.style = 'Table Grid'

    # Get the style configuration
    style = table_style[style_name]

    for element in ['th', 'td']:
        if element == 'th':
            # Format header cells
            for th_row in table_name.rows[0].cells:
                for paragraph in th_row.paragraphs:
                    run = paragraph.runs[0]
                    run.font.size = Pt(style[f'{element}_style']['font_size'])
                    run.font.name = style[f'{element}_style']['font_name']
                    run.font.color.rgb = RGBColor(*style[f'{element}_style']['font_color'])
                    run.bold = style[f'{element}_style']['bold']

                # Set header cell background color
                cell_xml_element = th_row._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement('w:shd')
                shade_obj.set(qn('w:fill'), style[f'{element}_style']['cell_color'])
                table_cell_properties.append(shade_obj)

        else:
            # Format table data cells
            for row_index in range(1, len(table_name.rows)):
                row_has_special = False

                # Check if any cell in the row meets the special criterion
                for cell in table_name.rows[row_index].cells:
                    if cell.text.strip() in special:
                        row_has_special = True
                        break

                # Apply formatting to the entire row if a special cell is found
                for cell in table_name.rows[row_index].cells:
                    cell_style = style[f'{element}_sp_style'] if row_has_special else style[f'{element}_style']

                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(cell_style['font_size'])
                            run.font.name = cell_style['font_name']
                            run.font.color.rgb = RGBColor(*cell_style['font_color'])
                            run.bold = cell_style['bold']

                    # Set cell background color
                    cell_xml_element = cell._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement('w:shd')
                    shade_obj.set(qn('w:fill'), cell_style['cell_color'])
                    table_cell_properties.append(shade_obj)


def profitandlossheads(data: pd.DataFrame, start_date: datetime, end_date: datetime, bu: list) -> pd.DataFrame:
    gp_filt = (data['third_level'].isin(['Cost of Sales', 'Direct Income'])) & (
            data['voucher_date'] >= start_date) & (data['voucher_date'] <= end_date) & (
                  data['business_unit_name'].isin(bu))
    gp: float = data.loc[gp_filt, 'amount'].sum()
    oh_filt = (data['third_level'].isin(['Overhead', 'Finance Cost'])) & (
            data['voucher_date'] >= start_date) & (data['voucher_date'] <= end_date) & (
                  data['business_unit_name'].isin(bu))
    overhead: float = data.loc[oh_filt, 'amount'].sum()
    np_filt = (data['forth_level'].isin(['Expenses', 'Income'])) & (data['voucher_date'] >= start_date) & (
            data['voucher_date'] <= end_date) & (data['business_unit_name'].isin(bu))
    np: float = data.loc[np_filt, 'amount'].sum()
    rev_filt = (data['third_level'].isin(['Direct Income'])) & (
            data['voucher_date'] >= start_date) & (data['voucher_date'] <= end_date) & (
                   data['business_unit_name'].isin(bu))
    rev: float = data.loc[rev_filt, 'amount'].sum()
    gp_row: pd.DataFrame = pd.DataFrame(data={'amount': [gp], 'voucher_date': [end_date]}, index=['Gross Profit'])
    oh_row: pd.DataFrame = pd.DataFrame(data={'amount': [overhead], 'voucher_date': [end_date]},
                                        index=['Total Overhead'])
    np_row: pd.DataFrame = pd.DataFrame(data={'amount': [np], 'voucher_date': [end_date]}, index=['Net Profit'])
    rev_row: pd.DataFrame = pd.DataFrame(data={'amount': [rev], 'voucher_date': [end_date]}, index=['Total Revenue'])
    pl_summary: pd.DataFrame = pd.concat([gp_row, oh_row, np_row, rev_row], ignore_index=False)
    return pl_summary


def profitandloss(data: pd.DataFrame, fBudget: pd.DataFrame, bu: list, end_date: datetime, start_date: datetime,
                  basic_pl: bool = False,
                  mid_pl: bool = False, full_pl: bool = False, ) -> dict:
    df_basic: pd.DataFrame = pd.DataFrame(data={'voucher_date': [], 'amount': []})
    df_basic_bud: pd.DataFrame = pd.DataFrame(data={'voucher_date': [], 'amount': []})
    df_mid: pd.DataFrame = pd.DataFrame(data={'voucher_date': [], 'amount': []})
    df_mid_bud: pd.DataFrame = pd.DataFrame(data={'voucher_date': [], 'amount': []})
    df_full: pd.DataFrame = pd.DataFrame(data={'voucher_date': [], 'amount': []})
    df_full_bud: pd.DataFrame = pd.DataFrame(data={'voucher_date': [], 'amount': []})
    basic: pd.DataFrame = pd.DataFrame()
    basic_bud: pd.DataFrame = pd.DataFrame()
    mid: pd.DataFrame = pd.DataFrame()
    mid_bud: pd.DataFrame = pd.DataFrame()
    full: pd.DataFrame = pd.DataFrame()
    full_bud: pd.DataFrame = pd.DataFrame()
    month_end_dates = pd.date_range(start=start_date, end=end_date, freq='ME')
    for end in month_end_dates:
        start: datetime = end + relativedelta(day=1)
        indirect_inc_filt = data['third_level'].isin(['Indirect Income']) & (
                data['voucher_date'] >= start) & (data['voucher_date'] <= end) & (
                                data['business_unit_name'].isin(bu))
        indirect_inc_brief: pd.DataFrame = data.loc[
            indirect_inc_filt, ['first_level', 'voucher_date', 'amount']].groupby(
            by=['voucher_date', 'first_level'], as_index=False).sum().rename(
            columns={'first_level': 'Description'})
        indirect_inc_brief = indirect_inc_brief.loc[indirect_inc_brief['amount'] != 0]
        indirect_inc_filt_bud = fBudget['third_level'].isin(['Indirect Income']) & (
                fBudget['voucher_date'] >= start) & (fBudget['voucher_date'] <= end) & (
                                    fBudget['business_unit_name'].isin(bu))
        indirect_inc_brief_bud: pd.DataFrame = fBudget.loc[
            indirect_inc_filt_bud, ['first_level', 'voucher_date', 'amount']].groupby(
            by=['voucher_date', 'first_level'], as_index=False).sum().rename(
            columns={'first_level': 'Description'})
        indirect_inc_brief_bud = indirect_inc_brief_bud.loc[indirect_inc_brief_bud['amount'] != 0]

        overhead_brief_filt = (data['third_level'].isin(['Overhead', 'Finance Cost'])) & (
                data['voucher_date'] >= start) & (data['voucher_date'] <= end) & (
                                  data['business_unit_name'].isin(bu))
        overhead_brief_filt_bud = (fBudget['third_level'].isin(['Overhead', 'Finance Cost'])) & (
                fBudget['voucher_date'] >= start) & (fBudget['voucher_date'] <= end) & (
                                      fBudget['business_unit_name'].isin(bu))
        summary_actual: pd.DataFrame = profitandlossheads(data=data, start_date=start, end_date=end, bu=bu)
        summary_budget: pd.DataFrame = profitandlossheads(data=fBudget, start_date=start, end_date=end, bu=bu)
        trade_account_filt = data['third_level'].isin(['Cost of Sales', 'Direct Income']) & (
                data['voucher_date'] >= start) & (data['voucher_date'] <= end) & (
                                 data['business_unit_name'].isin(bu))
        # basic version
        if basic_pl:
            trade_account_filt_bud = fBudget['third_level'].isin(['Cost of Sales', 'Direct Income']) & (
                    fBudget['voucher_date'] >= start) & (fBudget['voucher_date'] <= end) & (
                                         fBudget['business_unit_name'].isin(bu))
            overhead_brief_basic: pd.DataFrame = data.loc[
                overhead_brief_filt, ['first_level', 'voucher_date', 'amount']].groupby(
                by=['voucher_date', 'first_level'], as_index=False).sum().rename(
                columns={'first_level': 'Description'})
            overhead_brief_basic_bud: pd.DataFrame = fBudget.loc[
                overhead_brief_filt_bud, ['first_level', 'voucher_date', 'amount']].groupby(
                by=['voucher_date', 'first_level'], as_index=False).sum().rename(
                columns={'first_level': 'Description'})
            trade_account_brief: pd.DataFrame = data.loc[
                trade_account_filt, ['first_level', 'voucher_date', 'amount']].groupby(
                by=['voucher_date', 'first_level'], as_index=False).sum().rename(
                columns={'first_level': 'Description'})
            trade_account_brief_bud: pd.DataFrame = fBudget.loc[
                trade_account_filt_bud, ['first_level', 'voucher_date', 'amount']].groupby(
                by=['voucher_date', 'first_level'], as_index=False).sum().rename(
                columns={'first_level': 'Description'})
            basic: pd.DataFrame = pd.concat(
                i for i in [trade_account_brief, indirect_inc_brief, overhead_brief_basic] if not i.empty).rename(
                columns={'first_level': 'Description'})
            basic_bud: pd.DataFrame = pd.concat(
                [trade_account_brief_bud, indirect_inc_brief_bud, overhead_brief_basic_bud]).rename(
                columns={'first_level': 'Description'})
            basic = basic.loc[basic['amount'] != 0].set_index(keys='Description')
            basic_bud = basic_bud.loc[basic_bud['amount'] != 0].set_index(keys='Description')
            df_basic = pd.concat(i for i in [basic, summary_actual, df_basic] if not i.empty)
            df_basic_bud = pd.concat(i for i in [basic_bud, summary_budget, df_basic_bud] if not i.empty)

        # mid version
        if mid_pl:
            trade_account_mid: pd.DataFrame = data.loc[
                trade_account_filt, ['second_level', 'voucher_date', 'amount']].groupby(
                by=['voucher_date', 'second_level'], as_index=False).sum().rename(
                columns={'second_level': 'Description'})
            trade_account_mid_bud: pd.DataFrame = fBudget.loc[
                trade_account_filt_bud, ['second_level', 'voucher_date', 'amount']].groupby(
                by=['voucher_date', 'second_level'], as_index=False).sum().rename(
                columns={'second_level': 'Description'})
            overhead_brief_mid: pd.DataFrame = data.loc[
                overhead_brief_filt, ['first_level', 'voucher_date', 'amount']].groupby(
                by=['voucher_date', 'first_level'], as_index=False).sum().rename(
                columns={'first_level': 'Description'})
            overhead_brief_mid_bud: pd.DataFrame = fBudget.loc[
                overhead_brief_filt_bud, ['first_level', 'voucher_date', 'amount']].groupby(
                by=['voucher_date', 'first_level'], as_index=False).sum().rename(
                columns={'first_level': 'Description'})
            mid = pd.concat(
                [trade_account_mid, indirect_inc_brief, overhead_brief_mid])
            mid_bud = pd.concat(
                [trade_account_mid_bud, indirect_inc_brief_bud, overhead_brief_mid_bud])
            mid = mid.loc[mid['amount'] != 0].set_index(keys='Description')
            mid_bud = mid_bud.loc[mid_bud['amount'] != 0].set_index(keys='Description')
            df_mid = pd.concat([mid, summary_actual, df_mid])
            df_mid_bud = pd.concat([mid_bud, summary_budget, df_mid_bud])

        # full version
        if full_pl:
            detailed_filt = data['forth_level'].isin(['Income', 'Expenses']) & (
                    data['voucher_date'] >= start) & (data['voucher_date'] <= end) & (
                                data['business_unit_name'].isin(bu))
            detailed_filt_bud = fBudget['third_level'].isin(
                ['Indirect Income', 'Overhead', 'Finance Cost', 'Direct Income', 'Cost of Sales']) & (
                                        fBudget['voucher_date'] >= start) & (fBudget['voucher_date'] <= end) & (
                                    fBudget['business_unit_name'].isin(bu))
            full = data.loc[detailed_filt, ['ledger_name', 'voucher_date', 'amount']].groupby(
                by=['voucher_date', 'ledger_name'], as_index=False).sum().rename(columns={'ledger_name': 'Description'})
            full_bud = fBudget.loc[detailed_filt_bud, ['ledger_name', 'voucher_date', 'amount']].groupby(
                by=['voucher_date', 'ledger_name'], as_index=False).sum().rename(columns={'ledger_name': 'Description'})
            full = full.loc[full['amount'] != 0].set_index(keys='Description')
            full_bud = full_bud.loc[full_bud['amount'] != 0].set_index(keys='Description')
            df_full = pd.concat(i for i in [df_full, summary_actual, full] if not i.empty)
            df_full_bud = pd.concat(i for i in [df_full_bud, summary_budget, full_bud] if not i.empty)
    cy_cp_basic: pd.DataFrame = df_basic.loc[df_basic['voucher_date'] == end_date].drop(
        columns=['voucher_date']).reset_index().rename(columns={'index': 'Description'})
    cy_cp_basic_bud: pd.DataFrame = df_basic_bud.loc[df_basic_bud['voucher_date'] == end_date].drop(
        columns=['voucher_date']).reset_index().rename(columns={'index': 'Description'})

    cy_pp_basic: pd.DataFrame = df_basic.loc[
        df_basic['voucher_date'] == end_date - relativedelta(months=1) + relativedelta(day=31)].drop(
        columns=['voucher_date']).reset_index().rename(columns={'index': 'Description'})

    cy_ytd_basic: pd.DataFrame = df_basic.loc[(df_basic['voucher_date'] <= end_date) & (
            df_basic['voucher_date'] >= datetime(year=end_date.year, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['voucher_date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()
    last12: pd.DataFrame = df_basic.loc[(df_basic['voucher_date'] <= end_date) & (
            df_basic['voucher_date'] >= (end_date-relativedelta(years=1)+timedelta(days=1)))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['voucher_date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()
    cy_ytd_basic_bud: pd.DataFrame = df_basic_bud.loc[(df_basic_bud['voucher_date'] <= end_date) & (
            df_basic_bud['voucher_date'] >= datetime(year=end_date.year, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['voucher_date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()

    py_cp_basic: pd.DataFrame = df_basic.loc[
        df_basic['voucher_date'] == datetime(year=end_date.year - 1, month=end_date.month, day=end_date.day)].drop(
        columns=['voucher_date']).reset_index().rename(columns={'index': 'Description'})

    py_ytd_basic: pd.DataFrame = df_basic.loc[
        (df_basic['voucher_date'] <= end_date - relativedelta(years=1) + relativedelta(day=31)) & (
                df_basic['voucher_date'] >= datetime(year=end_date.year - 1, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['voucher_date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()

    cy_ytd_basic_monthwise: pd.DataFrame = df_basic.loc[(df_basic['voucher_date'] <= end_date) & (
            df_basic['voucher_date'] >= datetime(year=end_date.year, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).pivot_table(index='Description', columns='voucher_date', values='amount',
                                                      aggfunc='sum', fill_value=0).reset_index()

    cy_cp_mid: pd.DataFrame = df_mid.loc[df_mid['voucher_date'] == end_date].drop(
        columns=['voucher_date']).reset_index().rename(columns={'index': 'Description'})
    cy_cp_mid_bud: pd.DataFrame = df_mid_bud.loc[df_mid['voucher_date'] == end_date].drop(
        columns=['voucher_date']).reset_index().rename(columns={'index': 'Description'})

    cy_pp_mid: pd.DataFrame = df_mid.loc[
        df_mid['voucher_date'] == end_date - relativedelta(months=1) + relativedelta(day=31)].drop(
        columns=['voucher_date']).reset_index().rename(columns={'index': 'Description'})

    cy_ytd_mid: pd.DataFrame = df_mid.loc[
        (df_mid['voucher_date'] <= end_date) & (
                df_mid['voucher_date'] >= datetime(year=end_date.year, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['voucher_date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()
    cy_ytd_mid_bud: pd.DataFrame = df_mid_bud.loc[
        (df_mid_bud['voucher_date'] <= end_date) & (
                df_mid_bud['voucher_date'] >= datetime(year=end_date.year, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['voucher_date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()

    py_cp_mid: pd.DataFrame = df_mid.loc[
        df_mid['voucher_date'] == datetime(year=end_date.year - 1, month=end_date.month, day=end_date.day)].drop(
        columns=['voucher_date']).reset_index().rename(columns={'index': 'Description'})

    py_ytd_mid: pd.DataFrame = df_mid.loc[
        (df_mid['voucher_date'] <= end_date - relativedelta(years=1) + relativedelta(day=31)) & (
                df_mid['voucher_date'] >= datetime(year=end_date.year - 1, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['voucher_date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()

    cy_cp_full: pd.DataFrame = df_full.loc[df_full['voucher_date'] == end_date].drop(
        columns=['voucher_date']).reset_index().rename(columns={'index': 'Description'})
    cy_cp_full_bud: pd.DataFrame = df_full_bud.loc[df_full_bud['voucher_date'] == end_date].drop(
        columns=['voucher_date']).reset_index().rename(columns={'index': 'Description'})

    cy_pp_full: pd.DataFrame = df_full.loc[
        df_full['voucher_date'] == end_date - relativedelta(months=1) + relativedelta(day=31)].drop(
        columns=['voucher_date']).reset_index().rename(columns={'index': 'Description'})

    cy_ytd_full: pd.DataFrame = df_full.loc[(df_full['voucher_date'] <= end_date) & (
            df_full['voucher_date'] >= datetime(year=end_date.year, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['voucher_date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()
    cy_ytd_full_bud: pd.DataFrame = df_full_bud.loc[(df_full_bud['voucher_date'] <= end_date) & (
            df_full_bud['voucher_date'] >= datetime(year=end_date.year, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['voucher_date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()

    py_cp_full: pd.DataFrame = df_full.loc[
        df_full['voucher_date'] == datetime(year=end_date.year - 1, month=end_date.month, day=end_date.day)].drop(
        columns=['voucher_date']).reset_index().rename(columns={'index': 'Description'})

    py_ytd_full: pd.DataFrame = df_full.loc[
        (df_full['voucher_date'] <= end_date - relativedelta(years=1) + relativedelta(day=31)) & (
                df_full['voucher_date'] >= datetime(year=end_date.year - 1, month=1, day=1))].reset_index().rename(
        columns={'index': 'Description'}).drop(columns=['voucher_date']).groupby(by=['Description'],
                                                                                 as_index=False).sum()

    return {'df_basic': {'cy_cp_basic': cy_cp_basic, 'cy_pp_basic': cy_pp_basic, 'cy_ytd_basic': cy_ytd_basic,
                         'py_cp_basic': py_cp_basic, 'py_ytd_basic': py_ytd_basic, 'cy_cp_basic_bud': cy_cp_basic_bud,
                         'cy_ytd_basic_bud': cy_ytd_basic_bud, 'cy_ytd_basic_monthwise': cy_ytd_basic_monthwise,'last12':last12},

            'df_mid': {'cy_cp_mid': cy_cp_mid, 'cy_pp_mid': cy_pp_mid, 'cy_ytd_mid': cy_ytd_mid, 'py_cp_mid': py_cp_mid,
                       'py_ytd_mid': py_ytd_mid, 'cy_cp_mid_bud': cy_cp_mid_bud, 'cy_ytd_mid_bud': cy_ytd_mid_bud},

            'df_full': {'cy_cp_full': cy_cp_full, 'cy_pp_full': cy_pp_full, 'cy_ytd_full': cy_ytd_full,
                        'py_cp_full': py_cp_full, 'py_ytd_full': py_ytd_full, 'cy_cp_full_bud': cy_cp_full_bud,
                        'cy_ytd_full_bud': cy_ytd_full_bud}}


def plratios(df_pl: pd.DataFrame, plcombined: pd.DataFrame) -> dict:
    data :pd.DataFrame = plcombined.copy()
    data.drop(columns=['last12'],inplace=True)

    plmeasures: dict = {
        'gp': {'cy_cp_basic': 0, 'cy_ytd_basic': 0, 'cy_pp_basic': 0, 'py_cp_basic': 0, 'py_ytd_basic': 0,
               'cy_cp_basic_bud': 0, 'cy_ytd_basic_bud': 0, 'cy_ytd_basic_monthwise': 0},
        'np': {'cy_cp_basic': 0, 'cy_ytd_basic': 0, 'cy_pp_basic': 0, 'py_cp_basic': 0, 'py_ytd_basic': 0,
               'cy_cp_basic_bud': 0, 'cy_ytd_basic_bud': 0, 'cy_ytd_basic_monthwise': 0},
        'ebitda': {'cy_cp_basic': 0, 'cy_ytd_basic': 0, 'cy_pp_basic': 0, 'py_cp_basic': 0, 'py_ytd_basic': 0,
                   'cy_cp_basic_bud': 0, 'cy_ytd_basic_bud': 0, 'cy_ytd_basic_monthwise': 0}}

    for measure in plmeasures.keys():
        for k, v in df_pl['df_basic'].items():
            if k == 'cy_ytd_basic_monthwise':
                financial: pd.DataFrame = v.loc[v['Description'].isin(['Gross Profit', 'Net Profit', 'Total Revenue',
                                                                       'Depreciation',
                                                                       'Interest Expenses'])]
                gp = financial.index[financial['Description'] == 'Gross Profit'][0]
                netp = financial.index[financial['Description'] == 'Net Profit'][0]
                rev = financial.index[financial['Description'] == 'Total Revenue'][0]
                dep = financial.index[financial['Description'] == 'Depreciation'][0]
                interest = financial.index[financial['Description'] == 'Interest Expenses'][0]
                financial = financial.transpose().reset_index().rename(columns={gp: 'Gross Profit', netp: 'Net Profit',
                                                                                rev: 'Total Revenue',
                                                                                'index': 'Description',
                                                                                dep: 'Depreciation',
                                                                                interest: 'Interest Expenses'}).drop(0)
                financial.loc[:, 'EBITDA'] = financial['Net Profit'] - financial['Depreciation'] - financial[
                    'Interest Expenses']
                financial.drop(columns=['Depreciation', 'Interest Expenses'], inplace=True)
                plmeasures[measure][k] = financial
            else:
                df: pd.DataFrame = v.set_index('Description')
                if measure == 'gp':
                    ratio: float = df.loc['Gross Profit', 'amount'] / df.loc['Total Revenue', 'amount'] * 100
                if measure == 'np':
                    ratio: float = df.loc['Net Profit', 'amount'] / df.loc['Total Revenue', 'amount'] * 100
                if measure == 'ebitda':
                    ratio: float = (df.loc['Net Profit', 'amount'] -
                                    df.loc['Depreciation', 'amount'] if 'Depreciation' in df.index else 0 - df.loc[
                        'Interest Expenses', 'amount'] if 'Interest Expenses' in df.index else 0) / df.loc[
                                       'Total Revenue', 'amount'] * 100
                plmeasures[measure][k] = ratio

    data.fillna(0, inplace=True)
    data.set_index('Description', inplace=True)
    values: list = [np.nan] * len(data.columns)

    df_ratios = pd.DataFrame(
        data={'period': data.columns.tolist(), 'gp': values, 'np': values, 'ebitda': values, 'revenue': values})

    for period in df_ratios['period']:
        revenue: float = data.loc['Total Revenue', period]
        gp: float = data.loc['Gross Profit', period]
        netp: float = data.loc['Net Profit', period]
        interest: float = data.loc['Interest Expenses', period]
        dep: float = data.loc['Depreciation', period]

        ebitda: float = netp - dep - interest
        df_ratios.loc[df_ratios['period'] == period, 'gp'] = gp
        df_ratios.loc[df_ratios['period'] == period, 'np'] = netp
        df_ratios.loc[df_ratios['period'] == period, 'ebitda'] = ebitda
        df_ratios.loc[df_ratios['period'] == period, 'revenue'] = revenue
        plmeasures['plyearly'] = df_ratios
    return plmeasures


def pl_month_brief(document, data: pd.DataFrame, special: list):
    tbl_month_basic = document.add_table(rows=1, cols=5)
    tbl_month_basic.columns[0].width = Cm(7.5)
    heading_cells = tbl_month_basic.rows[0].cells
    heading_cells[0].text = 'Description'
    heading_cells[1].text = 'Current Month'
    heading_cells[2].text = 'Previous Month'
    heading_cells[3].text = 'SPLY'
    heading_cells[4].text = 'Budget'

    for _, row in data.iterrows():
        cells = tbl_month_basic.add_row().cells
        cells[0].text = str(row['Description'])
        cells[1].text = number_format(row.iloc[1])
        cells[2].text = number_format(row.iloc[2])
        cells[3].text = number_format(row.iloc[3])
        cells[4].text = number_format(row.iloc[4])

    table_formatter(table_name=tbl_month_basic, style_name='table_style_1', special=special)
    document.add_page_break()


def pl_month_detailed(document, data: pd.DataFrame, special: list):
    tbl_month_full = document.add_table(rows=1, cols=5)
    tbl_month_full.columns[0].width = Cm(11)
    heading_cells = tbl_month_full.rows[0].cells
    heading_cells[0].text = 'Description'
    heading_cells[1].text = 'Current Month'
    heading_cells[2].text = 'Previous Month'
    heading_cells[3].text = 'SPLY'
    heading_cells[4].text = 'Budget'

    for _, row in data.iterrows():
        cells = tbl_month_full.add_row().cells
        cells[0].text = str(row['Description'])
        cells[1].text = number_format(row.iloc[1])
        cells[2].text = number_format(row.iloc[2])
        cells[3].text = number_format(row.iloc[3])
        cells[4].text = number_format(row.iloc[4])
    table_formatter(table_name=tbl_month_full, style_name='table_style_1', special=special)
    document.add_page_break()


def coa_ordering(dCoAAdler: pd.DataFrame) -> list:
    other_income_df: pd.DataFrame = dCoAAdler.loc[dCoAAdler['third_level'] == 'Indirect Income'].copy()
    coa_df: pd.DataFrame = dCoAAdler.loc[dCoAAdler['third_level'] != 'Indirect Income'].copy()

    coa_df.sort_values(by='ledger_code', inplace=True)
    coa_df.set_index(keys='ledger_code', inplace=True)
    coa_df.reset_index(inplace=True)
    coa_list: list = coa_df['ledger_code'].tolist()

    other_income_df.sort_values(by='ledger_code', inplace=True)
    other_income_df.set_index(keys='ledger_code', inplace=True)
    other_income_df.reset_index(inplace=True)
    other_inc: list = other_income_df['ledger_code'].tolist()

    coa_sort_order: dict = coa_df['ledger_name'].reset_index().reset_index().set_index(keys='ledger_name')[
        'index'].to_dict()

    first_level: np.ndarray = coa_df['first_level'].unique()
    for i in first_level:
        coa_sort_order[i] = coa_list.index(
            coa_df.loc[(coa_df['first_level'] == i), 'ledger_code'].max()) + 0.1
    second_level: np.ndarray = coa_df['second_level'].unique()
    for i in second_level:
        coa_sort_order[i] = coa_list.index(
            coa_df.loc[(coa_df['second_level'] == i), 'ledger_code'].max()) + 0.2
    third_level: np.ndarray = coa_df['third_level'].unique()
    for i in third_level:
        coa_sort_order[i] = coa_list.index(
            coa_df.loc[(coa_df['third_level'] == i), 'ledger_code'].max()) + 0.3
    forth_level: np.ndarray = coa_df['forth_level'].unique()
    for i in forth_level:
        coa_sort_order[i] = coa_list.index(
            coa_df.loc[(coa_df['forth_level'] == i), 'ledger_code'].max()) + 0.4
    coa_sort_order['Gross Profit'] = coa_sort_order['Cost of Sales'] + 0.1

    for i, j in enumerate(other_inc):
        coa_sort_order[other_income_df.loc[other_income_df['ledger_code'] == j, 'ledger_name'].iloc[0]] = \
            coa_sort_order['Gross Profit'] + i / 10
    first_level_other_inc: np.ndarray = other_income_df['first_level'].unique()
    for i in first_level_other_inc:
        coa_sort_order[i] = coa_sort_order[other_income_df.loc[other_income_df['ledger_code'] == other_income_df.loc[
            (other_income_df['first_level'] == i), 'ledger_code'].max(), 'ledger_name'].iloc[0]] + 0.1
    second_level_other_inc: np.ndarray = other_income_df['second_level'].unique()
    for i in second_level_other_inc:
        coa_sort_order[i] = coa_sort_order[other_income_df.loc[other_income_df['ledger_code'] == other_income_df.loc[
            (other_income_df['second_level'] == i), 'ledger_code'].max(), 'ledger_name'].iloc[0]] + 0.2
    third_level_other_inc: np.ndarray = other_income_df['third_level'].unique()
    for i in third_level_other_inc:
        coa_sort_order[i] = coa_sort_order[other_income_df.loc[other_income_df['ledger_code'] == other_income_df.loc[
            (other_income_df['third_level'] == i), 'ledger_code'].max(), 'ledger_name'].iloc[0]] + 0.3
    forth_level_other_inc: np.ndarray = other_income_df['forth_level'].unique()
    for i in forth_level_other_inc:
        coa_sort_order[i] = coa_sort_order[other_income_df.loc[other_income_df['ledger_code'] == other_income_df.loc[
            (other_income_df['forth_level'] == i), 'ledger_code'].max(), 'ledger_name'].iloc[0]] + 0.4
    coa_sort_order['Total Overhead'] = coa_sort_order['Expenses'] + 0.1
    coa_sort_order['Net Profit'] = coa_sort_order['Expenses'] + 0.1

    value = coa_sort_order.pop('Direct Income')
    coa_sort_order['Total Revenue'] = value

    for i in ['Due From Related Parties', 'Due To Related Parties', 'Total Equity & Liabilities']:
        if i == 'Due From Related Parties':
            coa_sort_order['Due From Related Parties'] = coa_sort_order.get('Current Assets') - 0.01
        elif i == 'Due To Related Parties':
            coa_sort_order['Due To Related Parties'] = coa_sort_order.get('Current Liabilities') - 0.01
        else:
            coa_sort_order['Total Equity & Liabilities'] = coa_sort_order.get('Equity') + 0.01

    sorted_data = dict(sorted(coa_sort_order.items(), key=lambda item: item[1]))
    return sorted_data


def logistic_div(data: pd.DataFrame, ctgr: list, end_date: datetime) -> dict:
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    logistics_df = data.loc[(data['voucher_date'] >= start_date) & (data['voucher_date'] <= end_date) & (
        data['first_level'].isin(['Service Cost - Logistics', 'Staff Cost - Logistics', 'Logistics Revenue'])), [
        'voucher_date', 'amount', 'ledger_name', 'first_level']]
    ctgr_dict = {}
    for i in ctgr:
        df = logistics_df.loc[logistics_df['ledger_name'].str.contains(pat=i)]
        df = pd.pivot_table(data=df, index='first_level', columns=pd.Grouper(key='voucher_date', freq='ME'),
                            aggfunc='sum', values='amount', fill_value=0).reset_index().rename(
            columns={'first_level': 'Description'})
        ctgr_dict[i] = df
    return ctgr_dict


def log_div_profit(profit: dict, document):
    for k, v in profit.items():
        document.add_paragraph(f'\nProfitability of {k} Division')
        tbl_div_profit = document.add_table(rows=1, cols=v.shape[1])
        tbl_div_profit.columns[0].width = Cm(7.5)
        heading_cells = tbl_div_profit.rows[0].cells

        for i in range(v.shape[1]):
            if i == 0:
                heading_cells[i].text = 'Description'
            else:
                heading_cells[i].text = list(v.columns)[i].strftime('%b')

        for _, row in v.iterrows():
            cells = tbl_div_profit.add_row().cells
            for j in range(len(row)):
                if j == 0:
                    cells[0].text = str(row['Description'])
                else:
                    cells[j].text = number_format(row.iloc[j])
        table_formatter(table_name=tbl_div_profit, style_name='table_style_1', special=[])


def change_orientation(document, method):
    current_section = document.sections[-1]
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    if method == 'l':  # simple letter "L"
        new_width, new_height = current_section.page_height, current_section.page_width
        new_section.orientation = WD_ORIENT.LANDSCAPE
    else:
        new_height, new_width, = current_section.page_width, current_section.page_height
        new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width = new_width
    new_section.page_height = new_height

    return new_section


def pl_ytd_brief(document, data: pd.DataFrame, special: list):
    tbl_ytd_basic = document.add_table(rows=1, cols=4)
    tbl_ytd_basic.columns[0].width = Cm(11)
    heading_cells = tbl_ytd_basic.rows[0].cells
    heading_cells[0].text = 'Description'
    heading_cells[1].text = 'YTD CY'
    heading_cells[2].text = 'YTD PY'
    heading_cells[3].text = 'Budget'

    for _, row in data.iterrows():
        cells = tbl_ytd_basic.add_row().cells
        cells[0].text = str(row['Description'])
        cells[1].text = number_format(row.iloc[1])
        cells[2].text = number_format(row.iloc[2])
        cells[3].text = number_format(row.iloc[3])

    table_formatter(table_name=tbl_ytd_basic, style_name='table_style_1', special=special)


def gpnpebitda_graph(end_date: datetime, document, ratios):
    plt.style.use('ggplot')
    fig_pl, (ax1, ax2) = plt.subplots(nrows=2, ncols=1, figsize=(7.27, 10))

    ratiopl: pd.DataFrame = ratios['gp']['cy_ytd_basic_monthwise']
    ax1.set_title(f'GP Vs NP VS EBITDA - {end_date.year}')
    ax1.plot([i.strftime('%b') for i in ratiopl['voucher_date']],
             (ratiopl['Gross Profit'] / ratiopl['Total Revenue'] * 100),
             label='GP')
    for xy in zip([i.strftime('%b') for i in ratiopl['voucher_date']],
                  (ratiopl['Gross Profit'] / ratiopl['Total Revenue'] * 100).tolist()):
        ax1.annotate('{:,.0f}%'.format(xy[1]), xy=xy)
    ax1.plot([i.strftime('%b') for i in ratiopl['voucher_date']], (ratiopl['EBITDA'] / ratiopl['Total Revenue'] * 100),
             label='EBITDA')
    ax1.plot([i.strftime('%b') for i in ratiopl['voucher_date']],
             (ratiopl['Net Profit'] / ratiopl['Total Revenue'] * 100),
             label='NP')
    for xy in zip([i.strftime('%b') for i in ratiopl['voucher_date']],
                  (ratiopl['Net Profit'] / ratiopl['Total Revenue'] * 100).tolist()):
        ax1.annotate('{:,.0f}%'.format(xy[1]), xy=xy)

    tick_locations = ax1.get_yticks()
    ax1.yaxis.set_major_locator(FixedLocator(tick_locations))
    ax1.yaxis.set_major_formatter(FixedFormatter(['{:,.0f}%'.format(int(i)) for i in tick_locations]))
    ax1.legend()

    ratioplyearly: pd.DataFrame = ratios['plyearly']
    ax2.set_title(f'GP Vs NP VS EBITDA ({SYSTEM_CUT_OFF.year}-{end_date.year})')
    ax2.plot([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in ratioplyearly['period']],
             (ratioplyearly['gp'] / ratioplyearly['revenue'] * 100),
             label='GP')
    for xy in zip([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in ratioplyearly['period']],
                  (ratioplyearly['gp'] / ratioplyearly['revenue'] * 100).tolist()):
        ax2.annotate('{:,.0f}%'.format(xy[1]), xy=xy)
    ax2.plot([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in ratioplyearly['period']],
             (ratioplyearly['ebitda'] / ratioplyearly['revenue'] * 100),
             label='EBITDA')
    for xy in zip([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in ratioplyearly['period']],
                  (ratioplyearly['ebitda'] / ratioplyearly['revenue'] * 100).tolist()):
        ax2.annotate('{:,.0f}%'.format(xy[1]), xy=xy)
    ax2.plot([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in ratioplyearly['period']],
             (ratioplyearly['np'] / ratioplyearly['revenue'] * 100),
             label='NP')

    ax2.invert_xaxis()
    tick_locations = ax2.get_yticks()
    ax2.yaxis.set_major_locator(FixedLocator(tick_locations))
    ax2.yaxis.set_major_formatter(FixedFormatter(['{:,.0f}%'.format(int(i)) for i in tick_locations]))
    ax2.legend()

    pl_graph_buf = BytesIO()
    plt.tight_layout(h_pad=3)
    plt.savefig(pl_graph_buf, format='png')
    plt.close(fig_pl)
    pl_graph_buf.seek(0)
    document.add_picture(pl_graph_buf)


def plmonthwise(document, data, special):
    tbl_monthwise_basic = document.add_table(rows=1, cols=data.shape[1])
    tbl_monthwise_basic.columns[0].width = Cm(7.5)
    heading_cells = tbl_monthwise_basic.rows[0].cells

    for i in range(data.shape[1]):
        if i == 0:
            heading_cells[i].text = 'Description'
        else:
            heading_cells[i].text = list(data.columns)[i].strftime('%b')

    for _, row in data.iterrows():
        cells = tbl_monthwise_basic.add_row().cells
        for j in range(len(row)):
            if j == 0:
                cells[0].text = str(row['Description'])
            else:
                cells[j].text = number_format(row.iloc[j])
    table_formatter(table_name=tbl_monthwise_basic, style_name='table_style_1', special=special)


def narration_refine(row):
    sample_text = row['narration']
    start_index = sample_text.find('|')
    end_index = sample_text.find('|', start_index + 1)
    return sample_text[start_index + 1:end_index].title()


def abnormal_trn(data: pd.DataFrame, end_date: datetime):
    start_date: datetime = datetime(year=end_date.year, month=end_date.month, day=1)
    data = data.loc[data['narration'].str.contains(r'\|[^|]+\|', regex=True) & (data['voucher_date'] >= start_date) & (
            data['voucher_date'] <= end_date) & (data['ledger_code'] >= 5000000000), ['voucher_date', 'narration',
                                                                                      'ledger_name', 'amount',
                                                                                      'voucher_number']]
    data.loc[:, 'narration'] = data.apply(narration_refine, axis=1)
    data.loc[:, 'amount'] = data['amount'] * -1
    data = data.groupby(by=['narration', 'ledger_name', 'voucher_number'], as_index=False)['amount'].sum()
    data.sort_values(by='ledger_name', inplace=True)
    data = data.drop(columns=['voucher_number']).rename(columns={'narration': 'Description', 'ledger_name': 'Account'})
    return data


def excpdetails(document, data: pd.DataFrame, end_date: datetime):
    abnormal_df: pd.DataFrame = abnormal_trn(data=data, end_date=end_date)

    abnormal_tbl = document.add_table(rows=1, cols=3)
    abnormal_tbl.columns[0].width = Cm(11)
    heading_cells = abnormal_tbl.rows[0].cells
    heading_cells[0].text = 'Description'
    heading_cells[1].text = 'Account'
    heading_cells[2].text = 'amount'
    for _, row in abnormal_df.iterrows():
        cells = abnormal_tbl.add_row().cells
        cells[0].text = str(row['Description'])
        cells[1].text = str(row['Account'])
        cells[2].text = number_format(row['amount'])

    table_formatter(table_name=abnormal_tbl, style_name='table_style_1', special=[])


def guardingbumonthwise(document, end_date: datetime, special: list, fBudget: pd.DataFrame, merged: pd.DataFrame,
                        sort_order: list):
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    df_pl_bu_guarding: dict = profitandloss(basic_pl=True, data=merged, start_date=start_date, end_date=end_date,
                                            full_pl=False, bu=['GUARDING-ESS'], fBudget=fBudget)
    cy_cp_basic_guarding_bu: pd.DataFrame = df_pl_bu_guarding['df_basic']['cy_ytd_basic_monthwise']
    cy_cp_basic_guarding_bu.fillna(value=0, inplace=True)
    cy_cp_basic_guarding_bu['Description'] = pd.Categorical(cy_cp_basic_guarding_bu['Description'],
                                                            categories=[k for k in sort_order.keys()],
                                                            ordered=True)
    cy_cp_basic_guarding_bu.sort_values(by='Description', inplace=True)

    tbl_monthwise_basic_guarding_bu = document.add_table(rows=1, cols=cy_cp_basic_guarding_bu.shape[1])
    tbl_monthwise_basic_guarding_bu.columns[0].width = Cm(7.5)
    heading_cells = tbl_monthwise_basic_guarding_bu.rows[0].cells

    for i in range(cy_cp_basic_guarding_bu.shape[1]):
        if i == 0:
            heading_cells[i].text = 'Description'
        else:
            heading_cells[i].text = list(cy_cp_basic_guarding_bu.columns)[i].strftime('%b')

    for _, row in cy_cp_basic_guarding_bu.iterrows():
        cells = tbl_monthwise_basic_guarding_bu.add_row().cells
        for j in range(len(row)):
            if j == 0:
                cells[0].text = str(row['Description'])
            else:
                cells[j].text = number_format(row.iloc[j])
    table_formatter(table_name=tbl_monthwise_basic_guarding_bu, style_name='table_style_1', special=special)


def elvbumonthwise(document, end_date: datetime, special: list, fBudget: pd.DataFrame, merged: pd.DataFrame,
                   sort_order: list):
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    df_pl_bu_elv: dict = profitandloss(basic_pl=True, data=merged, start_date=start_date, end_date=end_date,
                                       full_pl=False,
                                       bu=['ELV-ESS'], fBudget=fBudget)
    cy_cp_basic_elv_bu: pd.DataFrame = df_pl_bu_elv['df_basic']['cy_ytd_basic_monthwise']
    cy_cp_basic_elv_bu.fillna(value=0, inplace=True)
    cy_cp_basic_elv_bu['Description'] = pd.Categorical(cy_cp_basic_elv_bu['Description'],
                                                       categories=[k for k in sort_order.keys()],
                                                       ordered=True)
    cy_cp_basic_elv_bu.sort_values(by='Description', inplace=True)

    tbl_monthwise_basic_elv_bu = document.add_table(rows=1, cols=cy_cp_basic_elv_bu.shape[1])
    tbl_monthwise_basic_elv_bu.columns[0].width = Cm(7.5)
    heading_cells = tbl_monthwise_basic_elv_bu.rows[0].cells

    for i in range(cy_cp_basic_elv_bu.shape[1]):
        if i == 0:
            heading_cells[i].text = 'Description'
        else:
            heading_cells[i].text = list(cy_cp_basic_elv_bu.columns)[i].strftime('%b')

    for _, row in cy_cp_basic_elv_bu.iterrows():
        cells = tbl_monthwise_basic_elv_bu.add_row().cells
        for j in range(len(row)):
            if j == 0:
                cells[0].text = str(row['Description'])
            else:
                cells[j].text = number_format(row.iloc[j])
    table_formatter(table_name=tbl_monthwise_basic_elv_bu, style_name='table_style_1', special=special)


def credits(document):
    credit = document.add_paragraph(
        '\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\nNadun Jayathunga\n')
    credit.add_run('Chief Accountant\nNasser Bin Nawaf & Partners Holding W.L.L\n')
    credit.add_run('mail:njayathunga@nbn.qa\nTel:+974 4403 0407').italic = True
    credit.add_run(f"\nReport generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.core_properties.author = "Nadun Jayathunga"
    document.core_properties.keywords = ("Chief Accountant\nNasser Bin Nawaf and Partners Holdings "
                                         "W.L.L\nmail:njayathunga@nbn.qa\nTele:+974 4403 0407")


def plhistorical(document, special: list, data: pd.DataFrame, sort_order: list):
    df:pd.DataFrame = data.copy()
    df.reset_index(inplace=True,drop=True)

    df.drop(columns=['last12'],inplace=True)

    tbl_yearly_pl = document.add_table(rows=1, cols=df.shape[1])
    heading_cells = tbl_yearly_pl.rows[0].cells
    df.fillna(value=0, inplace=True)
    df['Description'] = pd.Categorical(df['Description'], categories=[k for k in sort_order.keys()],
                                         ordered=True)
    df.sort_values(by='Description', inplace=True)


    for i in range(df.shape[1]):
        if i == 0:
            heading_cells[i].text = 'Description'
        else:
            heading_cells[i].text = list(df.columns)[i]

    for _, row in df.iterrows():
        cells = tbl_yearly_pl.add_row().cells
        for j in range(len(row)):
            if j == 0:
                cells[0].text = str(row['Description'])
            else:

                cells[j].text = f"{row.iloc[j]:,.0f}" if row.iloc[j] >= 0 else f"({abs(row.iloc[j]):,.0f})"

    table_formatter(table_name=tbl_yearly_pl, style_name='table_style_1', special=special)


def bshistorical(document, data: pd.DataFrame, special: list, sort_order: list):
    data['Description'] = pd.Categorical(data['Description'],
                                         categories=[k for k in sort_order.keys()],
                                         ordered=True)
    data.sort_values(by='Description', inplace=True)

    tbl_yearly_bs = document.add_table(rows=1, cols=data.shape[1])
    heading_cells = tbl_yearly_bs.rows[0].cells
    for i in range(data.shape[1]):
        if i == 0:
            heading_cells[i].text = 'Description'
        else:
            heading_cells[i].text = list(data.columns)[i]

    for _, row in data.iterrows():
        cells = tbl_yearly_bs.add_row().cells
        for j in range(len(row)):
            if j == 0:
                cells[0].text = str(row['Description'])
            else:
                cells[j].text = f"{row.iloc[j]:,.0f}" if row.iloc[j] >= 0 else f"({abs(row.iloc[j]):,.0f})"

    table_formatter(table_name=tbl_yearly_bs, style_name='table_style_1', special=special)


def interco_bal(data: pd.DataFrame, end_date: datetime, dCoAAdler: pd.DataFrame) -> dict:
    interco_final: pd.DataFrame = pd.DataFrame()
    for entity in company_data:
        interco_ids: list = dCoAAdler.loc[dCoAAdler['ledger_name'].isin(
            company_data[entity]['names']), 'ledger_code'].tolist()
        interco_filt = (data['voucher_date'] <= end_date) & (
            data['ledger_code'].isin(interco_ids))
        interco_df: pd.DataFrame = data.loc[interco_filt, ['amount']]
        interco_df['Description'] = entity
        interco_df = interco_df.groupby(
            by=['Description'], as_index=False).sum()
        interco_final = pd.concat([interco_final, interco_df])
    interco_final = interco_final.loc[interco_final['amount'] != 0]
    rpr: float = interco_final.loc[interco_final['amount'] < 0, 'amount'].sum()
    rpp: float = interco_final.loc[interco_final['amount'] > 0, 'amount'].sum()
    rpr_df: pd.DataFrame = interco_final.loc[interco_final['amount'] < 0, [
        'Description', 'amount']].sort_values(by='amount', ascending=True)
    rpp_df: pd.DataFrame = interco_final.loc[interco_final['amount'] > 0, [
        'Description', 'amount']].sort_values(by='amount', ascending=False)
    return {'rpr': rpr, 'rpp': rpp, 'rpr_df': rpr_df, 'rpp_df': rpp_df}


def balancesheet(data: pd.DataFrame, end_date: datetime, dCoAAdler: pd.DataFrame, database: str,
                 company_info: dict) -> pd.DataFrame:
    offset_accounts: list = [i['data']['offset_accounts'] for i in company_info if i['data']['database'] == database][0]
    # Sum total of offset_accounts is zero. i.e. PDC
    interco_acc_names: list = [i for j in [
        company_data[i]['names'] for i in company_data] for i in j]
    interco_acc_codes: list = dCoAAdler.loc[dCoAAdler['ledger_name'].isin(
        interco_acc_names), 'ledger_code'].tolist()

    exclude_bs_codes: list = offset_accounts + interco_acc_codes

    bs_filt = (data['voucher_date'] <= end_date) & (~data['ledger_code'].isin(exclude_bs_codes)) & (
        data['forth_level'].isin(['Assets', 'Liabilities', 'Equity']))
    is_filt = (data['voucher_date'] <= end_date) & (
        data['forth_level'].isin(['Income', 'Expenses']))

    dr_in_ap = data.loc[
        (data['second_level'] == 'Accounts Payables') & (data['voucher_date'] <= end_date), ['ledger_code',
                                                                                             'amount']].groupby(
        by='ledger_code').sum()
    # returns negative figure
    dr_in_ap = dr_in_ap.loc[dr_in_ap['amount'] < 0, 'amount'].sum()
    cr_in_ar = data.loc[
        (data['second_level'] == 'Trade Receivables') & (data['voucher_date'] <= end_date), ['ledger_code',
                                                                                             'amount']].groupby(
        by='ledger_code').sum()
    # returns positive figure
    cr_in_ar = cr_in_ar.loc[cr_in_ar['amount'] > 0, 'amount'].sum()

    bs_data: pd.DataFrame = data.loc[bs_filt, ['second_level', 'amount']].groupby(
        by=['second_level'], as_index=False).sum().rename(
        columns={'second_level': 'Description'}).set_index(keys='Description')
    cum_profit: float = data.loc[is_filt, 'amount'].sum()
    rounding_diff: float = data.loc[data['voucher_date'] <= end_date, 'amount'].sum()
    interco: dict = interco_bal(data=data, end_date=end_date, dCoAAdler=dCoAAdler)
    rpr: float = interco.get('rpr')
    rpp: float = interco.get('rpp')
    rpr_row: pd.DataFrame = pd.DataFrame(data={'amount': [rpr]}, index=[
        'Due From Related Parties'])
    rpp_row: pd.DataFrame = pd.DataFrame(
        data={'amount': [rpp]}, index=['Due To Related Parties'])
    bs_data.loc['Accounts Payables',
    'amount'] = bs_data.loc['Accounts Payables', 'amount'] - dr_in_ap
    bs_data.loc['Other Receivable',
    'amount'] = bs_data.loc['Other Receivable', 'amount'] + dr_in_ap
    bs_data.loc['Trade Receivables',
    'amount'] = bs_data.loc['Trade Receivables', 'amount'] + cr_in_ar
    bs_data.loc['Accruals & Other Payables', 'amount'] = bs_data.loc[
                                                             'Accruals & Other Payables', 'amount'] - cr_in_ar - rounding_diff
    bs_data.loc['Retained Earnings',
    'amount'] = bs_data.loc['Retained Earnings', 'amount'] + cum_profit
    bs_data = pd.concat([bs_data, rpr_row, rpp_row], ignore_index=False)

    ca: float = bs_data.loc['Cash & Cash Equivalents', 'amount'] + bs_data.loc['Inventory', 'amount'] + bs_data.loc[
        'Other Receivable', 'amount'] + bs_data.loc['Trade Receivables', 'amount'] + bs_data.loc[
                    'Due From Related Parties', 'amount']
    nca: float = (bs_data.loc['Intangible Assets', 'amount'] if 'Intangible Assets' in bs_data.index else 0) + \
                 bs_data.loc[
                     'Property, Plant  & Equipment', 'amount'] + \
                 bs_data.loc[
                     'Right of use Asset', 'amount']
    equity: float = bs_data.loc['Retained Earnings', 'amount'] + bs_data.loc['Share Capital', 'amount'] + (bs_data.loc[
                                                                                                               'Statutory Reserves', 'amount'] if 'Statutory Reserves' in bs_data.index else 0)
    cl: float = bs_data.loc['Accounts Payables', 'amount'] + bs_data.loc['Accruals & Other Payables', 'amount'] + \
                bs_data.loc[
                    'Due To Related Parties', 'amount'] + (bs_data.loc[
                                                               'Short Term Bank Facilities', 'amount'] if 'Short Term Bank Facilities' in bs_data.index else 0)
    ncl: float = bs_data.loc['Provisions', 'amount'] + bs_data.loc[
        'Lease Liabilities', 'amount'] if 'Lease Liabilities' in bs_data.index else 0

    ta: float = ca + nca
    tl: float = cl + ncl
    tle: float = tl + equity

    cl_row: pd.DataFrame = pd.DataFrame(
        data={'amount': [cl]}, index=['Current Liabilities'])
    ncl_row: pd.DataFrame = pd.DataFrame(data={'amount': [ncl]}, index=[
        'Non Current Liabilities'])
    tl_row: pd.DataFrame = pd.DataFrame(
        data={'amount': [tl]}, index=['Liabilities'])
    ca_row: pd.DataFrame = pd.DataFrame(
        data={'amount': [ca]}, index=['Current Assets'])
    nca_row: pd.DataFrame = pd.DataFrame(
        data={'amount': [nca]}, index=['Non Current Assets'])
    ta_row: pd.DataFrame = pd.DataFrame(
        data={'amount': [ta]}, index=['Assets'])
    equity_row: pd.DataFrame = pd.DataFrame(
        data={'amount': [equity]}, index=['Equity'])
    tle_row: pd.DataFrame = pd.DataFrame(data={'amount': [tle]}, index=[
        'Total Equity & Liabilities'])

    bs_data = pd.concat([bs_data, cl_row, ncl_row, tl_row, equity_row, ca_row, nca_row, ta_row, tle_row],
                        ignore_index=False)

    bs_data = bs_data.loc[bs_data['amount'] != 0]
    bs_data['amount'] = bs_data['amount'] * -1
    return bs_data


def rpbalances(document, end_date: datetime, data: pd.DataFrame, dCoAAdler: pd.DataFrame):
    interco: dict = interco_bal(data=data, end_date=end_date, dCoAAdler=dCoAAdler)
    rpr_df: pd.DataFrame = interco.get('rpr_df')
    rpp_df: pd.DataFrame = interco.get('rpp_df')

    rpr_total_row: pd.DataFrame = pd.DataFrame(data={'amount': [rpr_df['amount'].sum()], 'Description': 'Total'},
                                               index=[
                                                   '9999'])
    rpr_df = pd.concat([rpr_df, rpr_total_row], ignore_index=False)

    rpp_total_row: pd.DataFrame = pd.DataFrame(data={'amount': [rpp_df['amount'].sum()], 'Description': 'Total'},
                                               index=[
                                                   '9999'])
    rpp_df = pd.concat([rpp_df, rpp_total_row], ignore_index=False)
    document.add_paragraph('Related Party Receivables')
    tbl_rpr = document.add_table(rows=1, cols=2)
    heading_cells = tbl_rpr.rows[0].cells
    heading_cells[0].text = 'Description'
    heading_cells[1].text = 'Amount'

    for _, row in rpr_df.iterrows():
        cells = tbl_rpr.add_row().cells
        cells[0].text = str(row['Description'])
        cells[1].text = number_format(-row.iloc[1])

    table_formatter(table_name=tbl_rpr, style_name='table_style_1', special=['Total'])

    document.add_paragraph('\n\nRelated Party Payables')

    tbl_rpp = document.add_table(rows=1, cols=2)
    heading_cells = tbl_rpp.rows[0].cells
    heading_cells[0].text = 'Description'
    heading_cells[1].text = 'Amount'

    for _, row in rpp_df.iterrows():
        cells = tbl_rpp.add_row().cells
        cells[0].text = str(row['Description'])
        cells[1].text = number_format(row.iloc[1])

    table_formatter(table_name=tbl_rpp, style_name='table_style_1', special=['Total'])


def apbalances(document, fAP: pd.DataFrame):
    tbl_ap = document.add_table(rows=1, cols=fAP.shape[1])
    tbl_ap.columns[0].width = Cm(11)
    heading_cells = tbl_ap.rows[0].cells

    for i in range(fAP.shape[1]):
        if i == 0:
            heading_cells[i].text = 'Supplier'
        else:
            heading_cells[i].text = list(fAP.columns)[i]

    for _, row in fAP.iterrows():
        cells = tbl_ap.add_row().cells
        for j in range(len(row)):
            if j == 0:
                cells[0].text = str(row['ledger_name'])
            else:
                cells[j].text = number_format(row.iloc[j])
    table_formatter(table_name=tbl_ap, style_name='table_style_1', special=['Total'])


def bsratios(bsdata: pd.DataFrame, pldata: pd.DataFrame, periods: list, end_date: datetime) -> dict:

    values: list = [np.nan] * len(bsdata.columns)
    df_ratios_bs = pd.DataFrame(data={'period': bsdata.columns.tolist(), 'cr': values, 'ato': values, 'roe': values})
    for period in periods:
        if period != datetime(year=end_date.year - 1, month=end_date.month, day=end_date.day):
            current_period: str = period.strftime('%Y-%m-%d')
            prior_year: str = int(period.strftime('%Y')) - 1 if int(period.strftime('%Y')) != 2020 else int(
                period.strftime('%Y'))
            previous_period: str = f"{prior_year}-{period.strftime('%m')}-{period.strftime('%d')}"
            # current_ratio https://corporatefinanceinstitute.com/resources/accounting/current-ratio-formula/  Liquidity ratio
            current_ratio: float = bsdata.loc[bsdata['Description'] == 'Current Assets', current_period].iloc[0] / - \
                bsdata.loc[bsdata['Description'] == 'Current Liabilities', current_period].iloc[0]
            if period != end_date:
            # asset turnover ratio https://corporatefinanceinstitute.com/resources/accounting/asset-turnover-ratio/ efficiency
                asset_turnover: float = pldata.loc[pldata['Description'] == 'Total Revenue', current_period].iloc[0] / ((
                                                                                                                                bsdata.loc[
                                                                                                                                    bsdata[
                                                                                                                                        'Description'] == 'Assets', current_period].iloc[
                                                                                                                                    0] +
                                                                                                                                bsdata.loc[
                                                                                                                                    bsdata[
                                                                                                                                        'Description'] == 'Assets', previous_period].iloc[
                                                                                                                                    0]) / 2)
                # roe https://corporatefinanceinstitute.com/resources/accounting/what-is-return-on-equity-roe/ profitability
                #TODO ROE Calculation for the period 30-09-2024 should use Net Profit for the period 01-10-2023 Till 30-09-2024. Currently It considers YTD Net Profit
                roe: float = pldata.loc[pldata['Description'] == 'Net Profit', current_period].iloc[0] / ((-bsdata.loc[
                    bsdata['Description'] == 'Equity', current_period].iloc[0] + -bsdata.loc[
                    bsdata['Description'] == 'Equity', previous_period].iloc[0]) / 2) * 100
            else:

                asset_turnover: float = pldata.loc[pldata['Description'] == 'Total Revenue', 'last12'].iloc[0] / ((
                                                                                                                                bsdata.loc[
                                                                                                                                    bsdata[
                                                                                                                                        'Description'] == 'Assets', current_period].iloc[
                                                                                                                                    0] +
                                                                                                                                bsdata.loc[
                                                                                                                                    bsdata[
                                                                                                                                        'Description'] == 'Assets', previous_period].iloc[
                                                                                                                                    0]) / 2)

                roe: float = pldata.loc[pldata['Description'] == 'Net Profit', 'last12'].iloc[0] / ((-bsdata.loc[
                    bsdata['Description'] == 'Equity', current_period].iloc[0] + -bsdata.loc[
                    bsdata['Description'] == 'Equity', previous_period].iloc[0]) / 2) * 100
            df_ratios_bs.loc[df_ratios_bs['period'] == current_period, 'cr'] = current_ratio
            df_ratios_bs.loc[df_ratios_bs['period'] == current_period, 'ato'] = asset_turnover
            df_ratios_bs.loc[df_ratios_bs['period'] == current_period, 'roe'] = roe
    df_ratios_bs.dropna(inplace=True)
    return df_ratios_bs


def main_bs_ratios(document, end_date: datetime, bsdata: pd.DataFrame, pldata: pd.DataFrame, periods: list):
    plt.style.use('ggplot')
    fig_bs, (ax1, ax2, ax3) = plt.subplots(nrows=3, ncols=1, sharex=True, figsize=(7.27, 10))
    bs_ratios_df: pd.DataFrame = bsratios(bsdata=bsdata, pldata=pldata, periods=periods, end_date=end_date)
    ax1.set_title('Current Ratio')
    ax1.plot([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in bs_ratios_df['period']], bs_ratios_df['cr'])
    for xy in zip([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in bs_ratios_df['period']],
                  bs_ratios_df['cr'].tolist()):
        ax1.annotate('{:,.1f}'.format(xy[1]), xy=xy)
    tick_locations = ax1.get_yticks()
    ax1.yaxis.set_major_locator(FixedLocator(tick_locations))
    ax1.yaxis.set_major_formatter(FixedFormatter(['{:,.1f}'.format(int(i)) for i in tick_locations]))

    ax2.set_title('Assets Turnover Ratio')
    ax2.plot([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in bs_ratios_df['period']], bs_ratios_df['ato'])
    for xy in zip([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in bs_ratios_df['period']],
                  bs_ratios_df['ato'].tolist()):
        ax2.annotate('{:,.1f}'.format(xy[1]), xy=xy)
    tick_locations = ax2.get_yticks()
    ax2.yaxis.set_major_locator(FixedLocator(tick_locations))
    ax2.yaxis.set_major_formatter(FixedFormatter(['{:,.1f}'.format(int(i)) for i in tick_locations]))

    ax3.set_title('Return on Equity')
    ax3.plot([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in bs_ratios_df['period']], bs_ratios_df['roe'])
    for xy in zip([datetime.strptime(i, '%Y-%m-%d').strftime('%Y') for i in bs_ratios_df['period']],
                  bs_ratios_df['roe'].tolist()):
        ax3.annotate('{:,.0f}%'.format(xy[1]), xy=xy)
    tick_locations = ax3.get_yticks()
    ax3.yaxis.set_major_locator(FixedLocator(tick_locations))
    ax3.yaxis.set_major_formatter(FixedFormatter(['{:,.0f}%'.format(int(i)) for i in tick_locations]))
    ax3.invert_xaxis()

    bs_graph_buf = BytesIO()
    plt.tight_layout(h_pad=3)
    plt.savefig(bs_graph_buf, format='png')
    plt.close(fig_bs)
    bs_graph_buf.seek(0)
    document.add_picture(bs_graph_buf)


def plotting_period(end_date: datetime, months: int) -> datetime:
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    delta = relativedelta(dt1=end_date, dt2=start_date)
    period: int = delta.months + (delta.years * 12) + 1
    if period < 6:
        start_date = end_date - relativedelta(months=months - 1)
        start_date = datetime(year=start_date.year, month=start_date.month, day=1)
    else:
        start_date
    return start_date


def closing_date(row, dCustomer: pd.DataFrame, cust_list: list) -> datetime:
    """Add credit period (in days) to the voucher date and convert that date to end of the month

    Args:
        row (_type_): a row in dataframe

    Returns:
        datetime: last date of the month to which voucher becomes due
    """
    ledger_code: str = row['ledger_code']

    if ledger_code in cust_list:
        credit_days: int = int(dCustomer.loc[dCustomer['ledger_code'] == ledger_code, 'credit_days'].squeeze())
        due_date = row['voucher_date'] + timedelta(days=credit_days)
        return due_date + relativedelta(day=31)
    else:
        pass


def already_collected(row, fGL: pd.DataFrame, fCollection: pd.DataFrame, dCustomer: pd.DataFrame) -> float:
    """Target collection for a given period is calculated by adding the credit period given to each customer.
    Invoices to which Target collection for a given period comprises may contain invoices which has been
    already collected prior they become due or before the beginning of target collection period. i.e. Invoice raised
    in 31/05/2024 which has 60 days credit period will become target collection in the period of 31/07/2024. But if
    such invoice has been collected on 15/06/2024, it should no longer be considered as Target collection for the period
    31/07/2024.

    Args:
        row (_type_): A row in the dataframe

    Returns:
        float: amount already collected out of target collection
    """

    fGL = fGL.loc[(fGL['transaction_type'].isin(VOUCHER_TYPES)) & (fGL['ledger_code'] >= 1000000000) & (
            fGL['ledger_code'] <= 1999999999)]
    cust_list: list = dCustomer['ledger_code'].tolist()
    fGL.loc[:, 'due_date'] = fGL.apply(closing_date, axis=1, args=[dCustomer, cust_list])
    start_date: datetime = row['due_date'].replace(day=1)
    due_inv_list: list = fGL.loc[
        (fGL['due_date'] >= start_date) & (fGL['due_date'] <= row['due_date']), 'voucher_number'].unique()
    collected_filt = (fCollection['invoice_number'].isin(due_inv_list)) & (fCollection['voucher_date'] < start_date)
    return fCollection.loc[collected_filt, 'voucher_amount'].sum()


def collection(fCollection: pd.DataFrame, end_date: datetime, fGL: pd.DataFrame,
               dCustomer: pd.DataFrame) -> pd.DataFrame:
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    # filters the collection date based on the selection
    fCollection1 = fCollection.loc[
        (fCollection['voucher_date'] >= start_date) & (fCollection['voucher_date'] <= end_date)]
    # convert collection date to last date of the month, so it can be grouped to know total collected per period.
    fCollection1 = fCollection1.groupby(pd.Grouper(key='voucher_date', freq='ME'))[
        'voucher_amount'].sum().reset_index().rename(columns={'voucher_date': 'due_date', 'voucher_amount': 'actual'})
    fCollection1 = fCollection1.loc[(fCollection1['due_date'] >= start_date) & (fCollection1['due_date'] <= end_date)]
    # Reasons for Finance / Receipt total for a period not match with 'Actual' in this report
    # 1. Credit notes are part of 'Actual' in this report
    # 2. Receipts other than from customers i.e. Employee Receivable is not part of this report
    # 3. Receipts that were not allocated to invoices are not part of this report.
    # for 3 above check fCollection/Invoice Number Contains RV/CN and Payment Date ->Blank
    fGL1 = fGL.copy()
    fGL1 = fGL1.loc[(fGL1['transaction_type'].isin(VOUCHER_TYPES)) & (fGL1['ledger_code'] >= 1000000000) & (
            fGL1['ledger_code'] <= 1999999999)]
    fGL1.loc[:, 'amount'] = fGL1['amount'] * -1
    cust_list: list = dCustomer['ledger_code'].tolist()
    fGL1.loc[:, 'due_date'] = fGL1.apply(closing_date, axis=1, args=[dCustomer, cust_list])
    fGL1 = fGL1.loc[(fGL1['due_date'] >= start_date) & (fGL1['due_date'] <= end_date)]
    fGL1 = fGL1.groupby(by=['due_date'], as_index=False)['amount'].sum()
    fGL1.loc[:, 'already_collected'] = fGL1.apply(already_collected, axis=1, args=[fGL, fCollection, dCustomer])
    fGL1.loc[:, 'amount'] = fGL1['amount'] - fGL1['already_collected']
    fGL1.drop(columns=['already_collected'], inplace=True)
    fGL1.rename(columns={'amount': 'target'}, inplace=True)

    combined: pd.DataFrame = pd.concat([fGL1.set_index('due_date'), fCollection1.set_index('due_date')], axis=1,
                                       join='outer').reset_index()
    return combined


def revenue(end_date: datetime, data: pd.DataFrame, fInvoices: pd.DataFrame, database: str, fData: pd.DataFrame,
            dJobs: pd.DataFrame, dCustomer: pd.DataFrame) -> dict:
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)

    first_invoice_dates: pd.Series = fInvoices.groupby('customer_code')['invoice_date'].min()
    fInvoices: pd.DataFrame = fInvoices.loc[
        (fInvoices['invoice_date'] >= start_date) & (fInvoices['invoice_date'] <= end_date), ['invoice_number',
                                                                                              'customer_code',
                                                                                              'amount',
                                                                                              'emp_id',
                                                                                              'invoice_date',
                                                                                              'type']]
    fInvoices.loc[:, 'new_or_old'] = fInvoices.apply(
        lambda row: 'Existing' if row['invoice_date'] > first_invoice_dates[row['customer_code']] else 'New',
        axis=1)
    new_or_old: pd.DataFrame = fInvoices.groupby(by=['invoice_date', 'new_or_old'], as_index=False)[
        'amount'].sum()

    rev_filt = (data['third_level'] == 'Direct Income') & (
            data['voucher_date'] <= end_date)
    rev_division: pd.DataFrame = data.loc[rev_filt, ['voucher_date', 'amount', 'second_level']].groupby(
        by=['voucher_date', 'second_level'], as_index=False).sum()
    if database == 'elite_security':

        sales_invoices: np.ndarray = data.loc[rev_filt, 'voucher_number'].unique()
        total_invoices: np.ndarray = fInvoices['invoice_number'].unique()
        worked_invoices: list = [
            inv for inv in sales_invoices if inv in total_invoices]
        rev_category: pd.DataFrame = data.loc[
            (data['voucher_number'].isin(worked_invoices)) & (data['third_level'] == 'Direct Income'), [
                'voucher_number', 'amount', 'voucher_date']].rename(
            columns={'voucher_number': 'invoice_number'})
        rev_category: pd.DataFrame = pd.merge(left=rev_category, right=fInvoices[['invoice_number', 'type']],
                                              on='invoice_number', how='left').drop(columns=['invoice_number']).groupby(
            by=['voucher_date', 'type'], as_index=False).sum()
        inv_emp: pd.DataFrame = fInvoices.groupby(by=['invoice_date', 'emp_id'], as_index=False)[
            'amount'].sum()

    elif database == 'nbn_logistics':
        fData.loc[:, 'amount'] = fData['credit'] - fData['debit']
        fData = fData.loc[(fData['voucher_date'] >= start_date) & (fData['voucher_date'] <= end_date) & (
            fData['ledger_code'].isin([4010201001, 4010201002, 4010201003, 4010201004])), ['voucher_date', 'amount',
                                                                                           'order_id']]
        fData = pd.merge(left=fData, right=dJobs[['order_id', 'customer_code', 'emp_id']], on='order_id', how='left')
        fdata: pd.DataFrame = pd.merge(left=fData, right=dCustomer[['customer_code', 'type']], on='customer_code',
                                       how='left').drop(columns=['order_id', 'customer_code', 'emp_id'])
        rev_category = fdata.groupby(by=[pd.Grouper(key='voucher_date', freq='ME'), 'type'], as_index=False)[
            'amount'].sum()
        inv_emp1: pd.DataFrame = \
            fData.groupby(by=[pd.Grouper(key='voucher_date', freq='ME'), 'emp_id'], as_index=False)['amount'].sum()
        inv_emp: pd.DataFrame = pd.DataFrame()
        for month in inv_emp1['voucher_date'].unique():
            others = pd.DataFrame(data={'voucher_date': [], 'emp_id': [], 'amount': []})
            df1 = inv_emp1.loc[inv_emp1['voucher_date'] == month, ['voucher_date', 'emp_id', 'amount']].nlargest(2,
                                                                                                                 'amount')
            others_sum = inv_emp1.loc[
                (~inv_emp1['emp_id'].isin(df1['emp_id'])) & (inv_emp1['voucher_date'] == month), 'amount'].sum()
            others = pd.DataFrame(data={'voucher_date': [month], 'emp_id': ['other'], 'amount': [others_sum]})
            inv_emp: pd.DataFrame = pd.concat([inv_emp, df1, others])
        inv_emp.rename(columns={'voucher_date': 'invoice_date'}, inplace=True)


    elif database == 'nbn_realestate':
        pass
    else:
        pass

    return {'rev_division': rev_division, 'rev_category': rev_category, 'new_or_old': new_or_old, 'inv_emp': inv_emp}


def revenue_dashboard(document, end_date: datetime, months: int, database: str, df_rev: dict):
    rev_division: pd.DataFrame = df_rev['rev_division']

    rev_division_plot: pd.DataFrame = rev_division.copy()
    rev_division = rev_division.loc[(rev_division['voucher_date'] <= end_date) & (
            rev_division['voucher_date'] >= plotting_period(end_date=end_date, months=months))].pivot_table(
        index='second_level', columns='voucher_date', values='amount',
        aggfunc='sum', fill_value=0).reset_index().rename(columns={'second_level': 'Description'})

    rev_division_line: pd.DataFrame = rev_division_plot.loc[(rev_division_plot['voucher_date'] <= end_date) & (
            rev_division_plot['voucher_date'] >= plotting_period(end_date=end_date, months=months))].pivot_table(
        index='voucher_date', columns='second_level', values='amount',
        aggfunc='sum', fill_value=0).reset_index().rename(columns={'voucher_date': 'Period'}).set_index(
        keys='Period')

    rev_division_pie_ytd: pd.DataFrame = rev_division_plot.loc[(rev_division_plot['voucher_date'] <= end_date) & (
            rev_division_plot['voucher_date'] >= datetime(year=end_date.year, month=1, day=1)), [
        'second_level', 'amount']].groupby(by='second_level').sum().reset_index().rename(
        columns={'second_level': 'Category'}).set_index(keys='Category')

    rev_division_pie_month: pd.DataFrame = rev_division_plot.loc[(rev_division_plot['voucher_date'] <= end_date) & (
            rev_division_plot['voucher_date'] >= datetime(year=end_date.year, month=end_date.month, day=1)), [
        'second_level', 'amount']].rename(columns={'second_level': 'Category'}).set_index(
        keys='Category')

    rev_category: pd.DataFrame = df_rev['rev_category']
    rev_category_plot: pd.DataFrame = rev_category.copy()

    rev_category = rev_category.loc[(rev_category['voucher_date'] <= end_date) & (
            rev_category['voucher_date'] >= plotting_period(end_date=end_date, months=months))].pivot_table(
        index='type',
        columns='voucher_date',
        values='amount',
        aggfunc='sum',
        fill_value=0).reset_index().rename(
        columns={'type': 'Description'})

    rev_category_line: pd.DataFrame = rev_category_plot.loc[(rev_category_plot['voucher_date'] <= end_date) & (
            rev_category_plot['voucher_date'] >= plotting_period(end_date=end_date, months=months))].pivot_table(
        index='voucher_date', columns='type', values='amount',
        aggfunc='sum', fill_value=0).reset_index().rename(columns={'voucher_date': 'Period'}).set_index(
        keys='Period')

    rev_category_pie: pd.DataFrame = df_rev['rev_category']
    rev_category_pie_ytd: pd.DataFrame = rev_category_pie.loc[(rev_category_pie['voucher_date'] <= end_date) & (
            rev_category_pie['voucher_date'] >= datetime(year=end_date.year, month=1, day=1)), ['type',
                                                                                                'amount']].groupby(
        by='type').sum()

    rev_category_pie_month: pd.DataFrame = rev_category_pie.loc[(rev_category_pie['voucher_date'] <= end_date) & (
            rev_category_pie['voucher_date'] >= datetime(year=end_date.year, month=end_date.month, day=1)), ['type',
                                                                                                             'amount']].groupby(
        by='type').sum()

    if database == 'elite_security':
        rev_summary = plt.figure()
        rev_summary.set_figheight(7)
        rev_summary.set_figwidth(10.5)
        ini_shape = (4, 5)
        ax1 = plt.subplot2grid(shape=ini_shape, loc=(0, 0), colspan=4)
        ax2 = plt.subplot2grid(shape=ini_shape, loc=(1, 0), colspan=2)
        ax3 = plt.subplot2grid(shape=ini_shape, loc=(1, 2), colspan=2)
        ax4 = plt.subplot2grid(shape=ini_shape, loc=(2, 0), colspan=4)
        ax5 = plt.subplot2grid(shape=ini_shape, loc=(3, 0), colspan=2)
        ax6 = plt.subplot2grid(shape=ini_shape, loc=(3, 2), colspan=2)
        ax7 = plt.subplot2grid(shape=ini_shape, loc=(0, 4), colspan=1)
        ax8 = plt.subplot2grid(shape=ini_shape, loc=(1, 4), colspan=1)
        ax9 = plt.subplot2grid(shape=ini_shape, loc=(2, 4), colspan=1)
        ax10 = plt.subplot2grid(shape=ini_shape, loc=(3, 4), colspan=1)

        ax1.table(cellText=[[j[0]] + [f'{i:,.0f}' for i in j if isinstance(i, float)] for j in rev_category.values],
                  colLabels=['Description'] + [i.strftime('%b') for i in rev_category.columns if i != 'Description'],
                  cellLoc='left', loc='best')
        ax1.set_title('Market/Related-party sales')
        ax1.axis('off')
        ax2.plot([i.strftime('%b') for i in rev_category_line.index], rev_category_line['Market'])
        tick_locations = ax2.get_yticks()
        ax2.yaxis.set_major_locator(FixedLocator(tick_locations))
        ax2.yaxis.set_major_formatter(FixedFormatter(['{:.0f}K'.format(i / 1_000) for i in tick_locations]))
        ax2.set_title('Market Sales')
        ax3.plot([i.strftime('%b') for i in rev_category_line.index], rev_category_line['Related'])
        tick_locations = ax3.get_yticks()
        ax3.yaxis.set_major_locator(FixedLocator(tick_locations))
        ax3.yaxis.set_major_formatter(FixedFormatter(['{:.0f}K'.format(i / 1_000) for i in tick_locations]))
        ax3.set_title('Related Sales')

        ax4.table(cellText=[[j[0]] + [f'{i:,.0f}' for i in j if isinstance(i, float)] for j in rev_division.values],
                  colLabels=['Description'] + [i.strftime('%b') for i in rev_division.columns if i != 'Description'],
                  cellLoc='left', loc='best')

        ax4.set_title('Division wise monthly sales')
        ax4.axis('off')

        ax5.plot([i.strftime('%b') for i in rev_division_line.index], rev_division_line['Manpower'])
        tick_locations = ax5.get_yticks()
        ax5.yaxis.set_major_locator(FixedLocator(tick_locations))
        ax5.yaxis.set_major_formatter(FixedFormatter(['{:.0f}K'.format(i / 1_000) for i in tick_locations]))
        ax5.set_title('Manpower Sales')
        ax6.plot([i.strftime('%b') for i in rev_division_line.index], rev_division_line['Projects'])
        tick_locations = ax6.get_yticks()
        ax6.yaxis.set_major_locator(FixedLocator(tick_locations))
        ax6.yaxis.set_major_formatter(FixedFormatter(['{:.0f}K'.format(i / 1_000) for i in tick_locations]))
        ax6.set_title('Projects Sales')
        ax7.pie(x=rev_category_pie_month['amount'], labels=rev_category_pie_month.index, autopct='%.0f%%',
                labeldistance=1,
                pctdistance=0.3)
        ax7.set_title('Month')

        ax8.pie(x=rev_category_pie_ytd['amount'], labels=rev_category_pie_ytd.index, autopct='%.0f%%', labeldistance=1,
                pctdistance=0.3)
        ax8.set_title('Year')

        ax9.pie(x=rev_division_pie_month['amount'], labels=rev_division_pie_month.index, autopct='%.1f%%',
                labeldistance=1,
                pctdistance=0.5)
        ax9.set_title('Month')

        ax10.pie(x=rev_division_pie_ytd['amount'], labels=rev_division_pie_ytd.index, autopct='%.1f%%', labeldistance=1,
                 pctdistance=0.5)
        ax10.set_title('Year')
        plt.tight_layout()
    elif database == 'nbn_logistics':
        rev_summary = plt.figure()
        rev_summary.set_figheight(7)
        rev_summary.set_figwidth(10.5)
        gs1 = GridSpec(nrows=1, ncols=1, figure=rev_summary, left=0.05, right=0.79, top=0.95, bottom=0.89)
        ax1 = rev_summary.add_subplot(gs1[:, :])
        gs2 = GridSpec(nrows=1, ncols=2, figure=rev_summary, left=0.05, right=0.79, top=0.85, wspace=0.15, bottom=0.66)
        ax2 = rev_summary.add_subplot(gs2[:, :-1])
        ax3 = rev_summary.add_subplot(gs2[:, -1])
        gs3 = GridSpec(nrows=1, ncols=1, figure=rev_summary, left=0.05, right=0.79, top=0.6, bottom=0.46)
        ax4 = rev_summary.add_subplot(gs3[:, :])
        gs4 = GridSpec(nrows=2, ncols=2, figure=rev_summary, left=0.05, right=0.79, top=0.44, wspace=0.15, hspace=0.2,
                       bottom=0.05)
        ax5 = rev_summary.add_subplot(gs4[:-1, :-1])
        ax6 = rev_summary.add_subplot(gs4[:-1, -1])
        ax7 = rev_summary.add_subplot(gs4[-1, :-1])
        ax8 = rev_summary.add_subplot(gs4[-1, -1])
        gs5 = GridSpec(nrows=4, ncols=1, figure=rev_summary, left=0.8, right=1, top=0.95, bottom=0.00)
        ax9 = rev_summary.add_subplot(gs5[0, 0])
        ax10 = rev_summary.add_subplot(gs5[1, 0])
        ax11 = rev_summary.add_subplot(gs5[2, 0])
        ax12 = rev_summary.add_subplot(gs5[3, 0])

        ax1.table(cellText=[[j[0]] + [f'{i:,.0f}' for i in j if isinstance(i, float)] for j in rev_category.values],
                  colLabels=['Description'] + [i.strftime('%b') for i in rev_category.columns if i != 'Description'],
                  cellLoc='left', loc='best')
        ax1.set_title('Market/Related-party sales', fontsize=9)
        ax1.axis('off')
        ax2.plot([i.strftime('%b') for i in rev_category_line.index], rev_category_line['Market'])
        tick_locations = ax2.get_yticks()
        ax2.yaxis.set_major_locator(FixedLocator(tick_locations))
        ax2.yaxis.set_major_formatter(FixedFormatter(['{:.0f}K'.format(i / 1_000) for i in tick_locations]))
        ax2.set_title('Market Sales', fontsize=9)
        ax3.plot([i.strftime('%b') for i in rev_category_line.index], rev_category_line['Related'])
        tick_locations = ax3.get_yticks()
        ax3.yaxis.set_major_locator(FixedLocator(tick_locations))
        ax3.yaxis.set_major_formatter(FixedFormatter(['{:.0f}K'.format(i / 1_000) for i in tick_locations]))
        ax3.set_title('Related Sales', fontsize=9)

        ax4.table(cellText=[[j[0]] + [f'{i:,.0f}' for i in j if isinstance(i, float)] for j in rev_division.values],
                  colLabels=['Description'] + [i.strftime('%b') for i in rev_division.columns if i != 'Description'],
                  cellLoc='left', loc='best')
        ax4.set_title('Division wise monthly sales', fontsize=9)
        ax4.axis('off')

        ax5.set_title('Clearance', fontsize=9)
        ax5.plot([i.strftime('%b') for i in rev_division_line.index], rev_division_line['Clearance'])
        tick_locations = ax5.get_yticks()
        ax5.yaxis.set_major_locator(FixedLocator(tick_locations))
        ax5.yaxis.set_major_formatter(FixedFormatter(['{:.0f}K'.format(i / 1_000) for i in tick_locations]))

        ax6.set_title('Transport', fontsize=9)
        ax6.plot([i.strftime('%b') for i in rev_division_line.index], rev_division_line['Transport'])
        tick_locations = ax6.get_yticks()
        ax6.yaxis.set_major_locator(FixedLocator(tick_locations))
        ax6.yaxis.set_major_formatter(FixedFormatter(['{:.0f}K'.format(i / 1_000) for i in tick_locations]))

        ax7.set_title('Freight', fontsize=9)
        ax7.plot([i.strftime('%b') for i in rev_division_line.index], rev_division_line['Freight'])
        tick_locations = ax7.get_yticks()
        ax7.yaxis.set_major_locator(FixedLocator(tick_locations))
        ax7.yaxis.set_major_formatter(FixedFormatter(['{:.0f}K'.format(i / 1_000) for i in tick_locations]))

        ax8.set_title('Other', fontsize=9)
        ax8.plot([i.strftime('%b') for i in rev_division_line.index], rev_division_line['Other'])
        tick_locations = ax8.get_yticks()
        ax8.yaxis.set_major_locator(FixedLocator(tick_locations))
        ax8.yaxis.set_major_formatter(FixedFormatter(['{:.0f}'.format(i) for i in tick_locations]))

        ax9.set_title('Month', fontsize=9)
        ax9.pie(x=rev_category_pie_month['amount'], labels=rev_category_pie_month.index, labeldistance=1,
                pctdistance=0.5, autopct='%.0f%%')
        ax10.set_title('Year', fontsize=9)
        ax10.pie(x=rev_category_pie_ytd['amount'], labels=rev_category_pie_ytd.index, labeldistance=1, pctdistance=0.5,
                 autopct='%.0f%%')
        ax11.set_title('Month', fontsize=9)
        ax11.pie(x=rev_division_pie_month['amount'], labels=rev_division_pie_month.index, labeldistance=1,
                 pctdistance=0.5, autopct='%.0f%%')
        ax12.set_title('Year', fontsize=9)
        ax12.pie(x=rev_division_pie_ytd['amount'], labels=rev_division_pie_ytd.index, labeldistance=1, pctdistance=0.5,
                 autopct='%.0f%%')
    elif database == 'premium':
        pass
    else:
        pass

    buf_revenue = BytesIO()
    plt.savefig(buf_revenue, format='png', dpi=2400)
    plt.close(rev_summary)
    buf_revenue.seek(0)
    document.add_picture(buf_revenue)


def data_output(refined: dict, welcome_info: dict) -> dict:
    merged: pd.DataFrame = refined['fGL']
    financial_periods_pl: list = sorted(list(
        set([welcome_info['end_date']] + pd.date_range(start=refined['fGL']['voucher_date'].min(),
                                                       end=welcome_info['end_date'],
                                                       freq='YE').to_pydatetime().tolist())),
        reverse=True)
    plcombined: pd.DataFrame = pd.DataFrame()
    bu_plcombined = refined['fGL']['business_unit_name'].unique()
    for f_year in financial_periods_pl:
        pl: dict = profitandloss(data=merged, end_date=f_year,
                                 basic_pl=True, bu=bu_plcombined, fBudget=refined['fBudget'],
                                 start_date=max(SYSTEM_CUT_OFF, datetime(year=f_year.year, month=1, day=1)))
        pl_period: pd.DataFrame = pl['df_basic']['cy_ytd_basic'].rename(
            columns={'amount': f'{f_year.date()}'}).set_index(
            keys='Description')
        plcombined = pd.concat([plcombined, pl_period], axis=1)
    last12:dict = profitandloss(basic_pl=True,bu=bu_plcombined,data=merged,end_date=welcome_info['end_date'],fBudget=refined['fBudget'],full_pl=False,mid_pl=False,start_date=(welcome_info['end_date']-relativedelta(years=1)+timedelta(days=1)))
    last12:pd.DataFrame = last12['df_basic']['last12'].rename(columns={'amount': 'last12'}).set_index(keys='Description')
    plcombined = pd.concat([plcombined, last12], axis=1)
    plcombined = plcombined.reset_index()
    df_pl: dict = profitandloss(basic_pl=True, data=merged, end_date=welcome_info['end_date'], full_pl=True,
                                bu=bu_plcombined, fBudget=refined['fBudget'],
                                start_date=datetime(year=welcome_info['end_date'].year - 1, month=1, day=1))
    cy_cp_basic: pd.DataFrame = df_pl['df_basic']['cy_cp_basic'].groupby('Description', as_index=False)['amount'].sum()
    cy_ytd_basic: pd.DataFrame = df_pl['df_basic']['cy_ytd_basic'].groupby('Description', as_index=False)[
        'amount'].sum()
    cy_pp_basic: pd.DataFrame = df_pl['df_basic']['cy_pp_basic'].groupby('Description', as_index=False)['amount'].sum()
    py_cp_basic: pd.DataFrame = df_pl['df_basic']['py_cp_basic'].groupby('Description', as_index=False)['amount'].sum()
    py_ytd_basic: pd.DataFrame = df_pl['df_basic']['py_ytd_basic'].groupby('Description', as_index=False)[
        'amount'].sum()
    cy_cp_basic_bud: pd.DataFrame = df_pl['df_basic']['cy_cp_basic_bud'].groupby('Description', as_index=False)[
        'amount'].sum()
    cy_ytd_basic_bud: pd.DataFrame = df_pl['df_basic']['cy_ytd_basic_bud'].groupby('Description', as_index=False)[
        'amount'].sum()

    ratios_pandl: dict = plratios(df_pl=df_pl, plcombined=plcombined)

    sort_order: list = coa_ordering(dCoAAdler=refined['dCoAAdler'])
    cp_month: pd.DataFrame = pd.concat(
        [cy_cp_basic.set_index('Description'), cy_pp_basic.set_index('Description'),
         py_cp_basic.set_index('Description'),
         cy_cp_basic_bud.set_index('Description')],
        axis=1, join='outer').reset_index()
    cp_month.fillna(value=0, inplace=True)
    cp_month['Description'] = pd.Categorical(cp_month['Description'], categories=[k for k in sort_order.keys()],
                                             ordered=True)
    cp_month.sort_values(by='Description', inplace=True)

    cy_cp_full: pd.DataFrame = df_pl['df_full']['cy_cp_full']
    cy_pp_full: pd.DataFrame = df_pl['df_full']['cy_pp_full']
    py_cp_full: pd.DataFrame = df_pl['df_full']['py_cp_full']
    cy_cp_full_bud: pd.DataFrame = df_pl['df_full']['cy_cp_full_bud']

    cp_month_full: pd.DataFrame = pd.concat(
        [cy_cp_full.set_index('Description'), cy_pp_full.set_index('Description'), py_cp_full.set_index('Description'),
         cy_cp_full_bud.set_index('Description')],
        axis=1, join='outer').reset_index()
    cp_month_full.fillna(value=0, inplace=True)
    cp_month_full['Description'] = pd.Categorical(cp_month_full['Description'],
                                                  categories=[k for k in sort_order.keys()],
                                                  ordered=True)
    cp_month_full.sort_values(by='Description', inplace=True)

    cat_profit = {}
    if welcome_info['database'] == 'nbn_logistics':
        cat_profit: dict = logistic_div(data=merged, ctgr=welcome_info['rev_cats'], end_date=welcome_info['end_date'])

    cy_ytd_basic: pd.DataFrame = df_pl['df_basic']['cy_ytd_basic']
    py_ytd_basic: pd.DataFrame = df_pl['df_basic']['py_ytd_basic']
    cy_ytd_basic_bud: pd.DataFrame = df_pl['df_basic']['cy_ytd_basic_bud']

    cp_ytd: pd.DataFrame = pd.concat(
        [cy_ytd_basic.set_index('Description'), py_ytd_basic.set_index('Description'),
         cy_ytd_basic_bud.set_index('Description')], axis=1, join='outer').reset_index()
    cp_ytd.fillna(value=0, inplace=True)
    cp_ytd['Description'] = pd.Categorical(cp_ytd['Description'], categories=[k for k in sort_order.keys()],
                                           ordered=True)
    cp_ytd.sort_values(by='Description', inplace=True)

    cy_ytd_basic_monthwise: pd.DataFrame = df_pl['df_basic']['cy_ytd_basic_monthwise']
    cy_ytd_basic_monthwise.fillna(value=0, inplace=True)
    cy_ytd_basic_monthwise['Description'] = pd.Categorical(cy_ytd_basic_monthwise['Description'],
                                                           categories=[k for k in sort_order.keys()],
                                                           ordered=True)
    cy_ytd_basic_monthwise.sort_values(by='Description', inplace=True)

    financial_periods_bs: list = sorted(list(
        set([welcome_info['end_date'],
             datetime(year=welcome_info['end_date'].year - 1, month=welcome_info['end_date'].month,
                      day=welcome_info['end_date'].day)] + list(
            pd.date_range(start=refined['fGL']['voucher_date'].min(), end=welcome_info['end_date'], freq='YE')))),
        reverse=True)
    bscombined: pd.DataFrame = pd.DataFrame()
    for f_year in financial_periods_bs:
        bs: pd.DataFrame = balancesheet(company_info=company_info, data=merged, end_date=f_year,
                                        dCoAAdler=refined['dCoAAdler'],
                                        database=welcome_info['database']).rename(
            columns={'amount': f'{f_year.date()}'})
        bscombined = pd.concat([bscombined, bs], axis=1)
    bscombined = bscombined.reset_index().rename(columns={'index': 'Description'})
    bscombined.fillna(value=0, inplace=True)

    return {'cp_month': cp_month, 'cp_month_full': cp_month_full, 'cp_ytd': cp_ytd,
            'cy_ytd_basic_monthwise': cy_ytd_basic_monthwise,
            'merged': merged, 'cat_profit': cat_profit, 'plcombined': plcombined, 'bscombined': bscombined,
            'financial_periods_bs': financial_periods_bs,
            'ratios_pandl': ratios_pandl, 'sort_order': sort_order}


def revenue_dashboard_two(document, df_rev: dict, welcome_info: dict, refined_data: dict):
    plt.style.use('ggplot')
    fig_sales, (new_existing, salesman_wise, col_graph) = plt.subplots(nrows=3, ncols=1, sharex=True)
    fig_sales.set_figheight(7)
    fig_sales.set_figwidth(10.5)

    # if welcome_info['database'] == 'elite_security':
    new_or_old: pd.DataFrame = df_rev['new_or_old'].groupby(['invoice_date', 'new_or_old'])['amount'].sum().unstack(
        fill_value=0)

    inv_emp: pd.DataFrame = df_rev['inv_emp']
    demp: pd.DataFrame = refined_data['dEmployee'].copy().reset_index()
    inv_emp = pd.merge(left=inv_emp, right=demp[['emp_name', 'emp_id']], on='emp_id', how='left')
    inv_emp['emp_name'] = inv_emp.apply(
        lambda x: ' '.join(x['emp_name'].split(sep=' ')[:2]).title() if pd.notna(x['emp_name']) else 'Others', axis=1)
    inv_emp.drop(columns=['emp_id'], inplace=True)
    inv_emp: pd.DataFrame = inv_emp.groupby(['invoice_date', 'emp_name'])['amount'].sum().unstack(fill_value=0)
    inv_emp.index = [i.strftime('%b') for i in inv_emp.index]

    monthly_collection: pd.DataFrame = collection(dCustomer=refined_data['dCustomer'],
                                                  end_date=welcome_info['end_date'], fGL=refined_data['fGL'],
                                                  fCollection=refined_data['fCollection'])

    new_or_old.index = [i.strftime('%b') for i in new_or_old.index]
    new_or_old.plot(kind='bar', stacked=True, ax=new_existing)
    new_existing.set_title('Revenue by Existing / New Customers')
    new_existing.legend()
    tick_locations = new_existing.get_yticks()
    new_existing.yaxis.set_major_locator(FixedLocator(tick_locations))
    new_existing.yaxis.set_major_formatter(FixedFormatter(['{:.0f}K'.format(i / 1_000) for i in tick_locations]))

    inv_emp.plot(kind='bar', stacked=True, ax=salesman_wise)
    salesman_wise.set_title('Revenue by Sales Person')
    salesman_wise.legend(loc='best')
    tick_locations = salesman_wise.get_yticks()
    salesman_wise.yaxis.set_major_locator(FixedLocator(tick_locations))
    salesman_wise.yaxis.set_major_formatter(FixedFormatter(['{:.0f}K'.format(i / 1_000) for i in tick_locations]))

    col_graph.set_title('Target Collection Vs Actual Collection')
    col_graph.plot([i.strftime('%b') for i in monthly_collection['due_date']], monthly_collection['target'],
                   label='Target')
    col_graph.plot([i.strftime('%b') for i in monthly_collection['due_date']], monthly_collection['actual'],
                   label='Actual')
    tick_locations = col_graph.get_yticks()
    col_graph.yaxis.set_major_locator(FixedLocator(tick_locations))
    col_graph.yaxis.set_major_formatter(FixedFormatter(['{:.0f}K'.format(i / 1_000) for i in tick_locations]))
    col_graph.legend()

    buf_sales = BytesIO()
    plt.tight_layout()
    plt.savefig(buf_sales, format='png', dpi=2400)
    plt.close(fig_sales)
    buf_sales.seek(0)
    document.add_picture(buf_sales)


def organic_sales(emp_id: str, mode: str, end_date: datetime, fInvoices: pd.DataFrame) -> float:
    if mode.lower() == 'month':
        start_date: datetime = datetime(year=end_date.year, month=end_date.month, day=1)
    elif mode.lower() == 'ytd':
        start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    else:
        raise ValueError(f'Invalid mode{mode}')

    customers: list = list(fInvoices.loc[
                               (fInvoices['invoice_date'] <= end_date) & (fInvoices['invoice_date'] >= start_date) & (
                                       fInvoices['emp_id'] == emp_id), 'customer_code'].unique())
    for customer in customers:
        first_sales_person: str = fInvoices.loc[(fInvoices['customer_code'] == customer), 'emp_id'].tolist()[0]
        if first_sales_person != emp_id:
            customers.remove(customer)
    self_sales: float = fInvoices.loc[
        (fInvoices['customer_code'].isin(customers)) & (fInvoices['invoice_date'] <= end_date) & (
                fInvoices['invoice_date'] >= start_date), 'amount'].sum()
    return self_sales


def sales_person(emp_ids: np.ndarray, dEmployee: pd.DataFrame, fInvoices: pd.DataFrame, end_date: datetime,
                 fGL: pd.DataFrame) -> dict:
    salesperson_stats: dict = {}
    for emp_id in emp_ids:
        doj: datetime = dEmployee.loc[dEmployee['emp_id'] == emp_id, 'doj'].iloc[0]
        cy_cp_rev: float = fInvoices.loc[(fInvoices['invoice_date'] <= end_date) & (
                fInvoices['invoice_date'] >= datetime(year=end_date.year, month=end_date.month, day=1)) & (
                                                 fInvoices['emp_id'] == emp_id), 'amount'].sum()
        cy_ytd_rev: float = fInvoices.loc[(fInvoices['invoice_date'] <= end_date) & (
                fInvoices['invoice_date'] >= datetime(year=end_date.year, month=1, day=1)) & (
                                                  fInvoices['emp_id'] == emp_id), 'amount'].sum()
        cy_cp_rev_org: float = organic_sales(emp_id=emp_id, mode='month', end_date=end_date, fInvoices=fInvoices)
        cy_ytd_rev_org: float = organic_sales(emp_id=emp_id, mode='ytd', end_date=end_date, fInvoices=fInvoices)
        cy_cp_customers: list = list(set(fInvoices.loc[(fInvoices['emp_id'] == emp_id) & (
                fInvoices['invoice_date'] >= datetime(year=end_date.year, month=end_date.month, day=1)) & (fInvoices[
                                                                                                               'invoice_date'] <= end_date), 'customer_code'].tolist()))
        customers_till: list = list(set(fInvoices.loc[(
                fInvoices['invoice_date'] <= end_date + relativedelta(day=31, months=-1)), 'customer_code'].tolist()))
        new_customers_added: int = len([customer for customer in cy_cp_customers if customer not in customers_till])

        ar_balance: pd.DataFrame = fGL.loc[
            (fGL['ledger_code'].isin(
                fInvoices.loc[fInvoices['customer_code'].isin(cy_cp_customers), 'ledger_code'].unique())) & (
                    fGL['voucher_date'] <= end_date), ['ledger_code', 'amount']].groupby(by='ledger_code').sum()
        monthly_rev: pd.DataFrame = fInvoices.loc[
            (fInvoices['invoice_date'] <= end_date) & (fInvoices['emp_id'] == emp_id) & (
                    fInvoices['invoice_date'] >= datetime(year=end_date.year, month=1, day=1)), ['invoice_date',
                                                                                                 'amount']].groupby(
            by='invoice_date').sum()
        cy_cp_rev_total: float = fInvoices.loc[(fInvoices['invoice_date'] <= end_date) & (
                fInvoices['invoice_date'] >= datetime(year=end_date.year, month=end_date.month,
                                                      day=1)), 'amount'].sum()
        cy_ytd_rev_total: float = fInvoices.loc[(fInvoices['invoice_date'] <= end_date) & (
                fInvoices['invoice_date'] >= datetime(year=end_date.year, month=1, day=1)), 'amount'].sum()
        cy_cp_rev_contrib_pct: float = cy_cp_rev / cy_cp_rev_total * 100
        cy_ytd_rev_contrib_pct: float = cy_ytd_rev / cy_ytd_rev_total * 100
        stats: dict = {'doj': doj.strftime('%d-%m-%Y'), 'cp_target': 0, 'cy_cp_rev': cy_cp_rev, 'ytd_target': 0,
                       'cy_ytd_rev': cy_ytd_rev, 'cy_cp_rev_org': cy_cp_rev_org,
                       'cy_ytd_rev_org': cy_ytd_rev_org,
                       'new_customers_added': new_customers_added, 'cy_cp_gp': 0, 'cy_ytd_gp': 0,
                       'ar_balance': ar_balance, 'monthly_rev': monthly_rev,
                       'cy_cp_rev_contrib_pct': f'{round(cy_cp_rev_contrib_pct, 1)}%',
                       'cy_ytd_rev_contrib_pct': f'{round(cy_ytd_rev_contrib_pct, 1)}%'}
        salesperson_stats[emp_id] = stats

    return salesperson_stats


def toc_customer(document, fInvoices: pd.DataFrame, end_date):
    customer_list: list = sorted(fInvoices.loc[(fInvoices['invoice_date'] >= datetime(year=end_date.year,
                                                                                      month=end_date.month, day=1)) & (
                                                       fInvoices['invoice_date'] <= end_date), 'cus_name'].unique())
    tbl_cust_toc = document.add_table(rows=1, cols=2)
    heading_cells = tbl_cust_toc.rows[0].cells
    heading_cells[0].text = 'Customer Name'
    heading_cells[1].text = 'Page #'

    for idx, row in enumerate(customer_list):
        cells = tbl_cust_toc.add_row().cells
        cells[0].text = str(row.upper())
        cells[1].text = str(idx + 1)

    table_formatter(table_name=tbl_cust_toc, style_name='table_style_1', special=[])


def toc_salseperson(document, dEmployee: pd.DataFrame, salesperson_list: np.ndarray):
    tbl_salesman_toc = document.add_table(rows=1, cols=2)
    heading_cells = tbl_salesman_toc.rows[0].cells
    heading_cells[0].text = 'Salesperson Name'
    heading_cells[1].text = 'Page #'

    for idx, row in enumerate(salesperson_list):
        cells = tbl_salesman_toc.add_row().cells
        cells[0].text = ' '.join(
            dEmployee.loc[dEmployee['emp_id'] == row, 'emp_name'].iloc[0].split(sep=' ')[:2]).title()
        cells[1].text = str(idx + 1)

    table_formatter(table_name=tbl_salesman_toc, style_name='table_style_1', special=[])
    document.add_page_break()


def settlement_days(invoices: list, fCollection: pd.DataFrame, end_date: datetime) -> list:
    col_days: list = []
    invoices = [inv for inv in invoices if not "CN" in inv]

    for invoice in invoices:
        inv_value: float = fCollection.loc[(fCollection['invoice_number'] == invoice), 'invoice_amount'].iloc[0]
        total_collection: float = fCollection.loc[(fCollection['invoice_number'] == invoice) & (
                fCollection['voucher_date'] <= end_date), 'voucher_amount'].sum()

        if (inv_value - total_collection) == 0:
            last_date: datetime = fCollection.loc[(fCollection['invoice_number'] == invoice) & (
                    fCollection['voucher_date'] <= end_date), 'voucher_date'].max()
            inv_date: datetime = fCollection.loc[(fCollection['invoice_number'] == invoice), 'invoice_date'].iloc[0]

            col_days.append(last_date - inv_date)

    return statistics.median(col_days) if col_days else timedelta(days=0)


def cust_ageing(customer: str, dCustomer: pd.DataFrame, fCollection: pd.DataFrame, end_date: datetime) -> pd.DataFrame:
    ledgers: list = dCustomer.loc[(dCustomer['cus_name'] == customer), 'ledger_code'].tolist()
    credit_days: int = int(dCustomer.loc[dCustomer['cus_name'].isin([customer]), 'credit_days'].iloc[0])
    invoices: np.ndarray = fCollection.loc[
        fCollection['ledger_code'].isin(ledgers) & (fCollection['invoice_date'] <= end_date), 'invoice_number'].unique()
    cust_soa: pd.DataFrame = fCollection.loc[(fCollection['invoice_number'].isin(invoices)), ['invoice_date',
                                                                                              'invoice_amount',
                                                                                              'voucher_amount',
                                                                                              'invoice_number',
                                                                                              'voucher_date']]
    inv_value_list: list = []
    age_bracket_list: list = []
    ranges = [(0, 'Not Due'), (30, '1-30'), (60, '31-60'),
              (90, '61-90'), (120, '91-120'), (121, '121-150'),
              (151, '151-180'), (181, '181-210'), (211, '211-240'),
              (241, '241-270'), (271, '271-300'), (300, '301-330'),
              (331, '331-360'), (float('inf'), 'More than 361')]
    for invoice in invoices:
        total_collection: float = cust_soa.loc[
            (cust_soa['invoice_number'] == invoice) & (cust_soa['voucher_date'] <= end_date), 'voucher_amount'].sum()
        inv_value: float = cust_soa.loc[(cust_soa['invoice_number'] == invoice), 'invoice_amount'].iloc[0]
        if (inv_value - total_collection) != 0:
            inv_value_list.append(inv_value - total_collection)
            days_passed: int = (end_date - cust_soa.loc[(cust_soa['invoice_number'] == invoice), 'invoice_date'].iloc[
                0] - timedelta(days=credit_days)).days
            for threshold, label in ranges:
                if days_passed <= threshold:
                    age_bracket_list.append(label)
                    break

    outstanding_df: pd.DataFrame = pd.DataFrame(
        data={'amount': inv_value_list, 'Age Bracket': age_bracket_list}).groupby(by='Age Bracket').sum()
    if not outstanding_df.empty:
        outstanding_df.reset_index(inplace=True)
        outstanding_df['Age Bracket'] = pd.Categorical(outstanding_df['Age Bracket'], categories=[i[1] for i in ranges],
                                                       ordered=True)
        outstanding_df.sort_values(by='Age Bracket', inplace=True)
        outstanding_df.set_index(keys='Age Bracket', drop=True, inplace=True)
    else:
        outstanding_df
    return outstanding_df


def customer_ratios(customers: list, fInvoices: pd.DataFrame, end_date: datetime, fCollection: pd.DataFrame,
                    dCustomer: pd.DataFrame, dEmployee: pd.DataFrame, fGL: pd.DataFrame, database: str,
                    fLogInv: pd.DataFrame) -> dict:
    customer_info: dict = {}
    for customer in customers:
        customer_since: datetime = fInvoices.loc[
            (fInvoices['cus_name'] == customer), 'invoice_date'].min() if not pd.isna(
            fInvoices.loc[(fInvoices['cus_name'] == customer), 'invoice_date'].min()) else "Not Applicable"
        total_revenue: float = fInvoices.loc[
            (fInvoices['cus_name'] == customer) & (fInvoices['invoice_date'] <= end_date), 'amount'].sum()
        cust_invoices: list = fInvoices.loc[(fInvoices['cus_name'] == customer), 'invoice_number'].to_list()

        last_receipt_dt: datetime = fCollection.loc[
            fCollection['invoice_number'].isin(cust_invoices), 'voucher_date'].max() if not pd.isna(fCollection.loc[
                                                                                                        fCollection[
                                                                                                            'invoice_number'].isin(
                                                                                                            cust_invoices), 'voucher_date'].max()) else "Not Collected"
        last_receipt_number: str = "Not Collected" if last_receipt_dt == "Not Collected" else \
            fCollection.loc[(fCollection['invoice_number'].isin(cust_invoices)) & (
                    fCollection['voucher_date'] == last_receipt_dt), 'voucher_number'].tail(1).iloc[0]
        last_receipt_amt: float = "Not Collected" if last_receipt_dt == "Not Collected" else fCollection.loc[
            (fCollection['voucher_number'] == last_receipt_number), 'voucher_amount'].sum()
        cy_cp_rev: float = fInvoices.loc[
            (fInvoices['cus_name'] == customer) & (fInvoices['invoice_date'] <= end_date) & (
                    fInvoices['invoice_date'] >= datetime(year=end_date.year, month=end_date.month,
                                                          day=1)), 'amount'].sum()
        cy_pp_rev: float = fInvoices.loc[(fInvoices['cus_name'] == customer) & (
                fInvoices['invoice_date'] <= end_date.replace(day=1) - timedelta(days=1)) & (
                                                 fInvoices['invoice_date'] >= end_date + relativedelta(day=1,
                                                                                                       months=-1)), 'amount'].sum()
        cy_ytd_rev: float = fInvoices.loc[
            (fInvoices['cus_name'] == customer) & (fInvoices['invoice_date'] <= end_date) & (
                    fInvoices['invoice_date'] >= datetime(year=end_date.year, month=1, day=1)), 'amount'].sum()
        py_ytd_rev: float = fInvoices.loc[(fInvoices['cus_name'] == customer) & (
                fInvoices['invoice_date'] <= datetime(year=end_date.year - 1, month=end_date.month,
                                                      day=end_date.day)) & (
                                                  fInvoices['invoice_date'] >= datetime(year=end_date.year - 1, month=1,
                                                                                        day=1)), 'amount'].sum()
        py_cp_rev: float = fInvoices.loc[(fInvoices['cus_name'] == customer) & (
                fInvoices['invoice_date'] <= datetime(year=end_date.year - 1, month=end_date.month,
                                                      day=end_date.day)) & (
                                                 fInvoices['invoice_date'] >= datetime(year=end_date.year - 1,
                                                                                       month=end_date.month,
                                                                                       day=1)), 'amount'].sum()
        collection_median: float = "Not Collected" if last_receipt_dt == "Not Collected" else settlement_days(
            invoices=cust_invoices, end_date=end_date, fCollection=fCollection)

        credit_days: int = dCustomer.loc[dCustomer['cus_name'].isin([customer]), 'credit_days'].iloc[0]
        date_established: datetime = dCustomer.loc[dCustomer['cus_name'].isin([customer]), 'date_established'].iloc[0]
        outstanding_bal: float = fGL.loc[
            (fGL['ledger_code'].isin(dCustomer.loc[dCustomer['cus_name'].isin([customer]), 'ledger_code'].tolist())) & (
                    fGL['voucher_date'] <= end_date), 'amount'].sum()
        cy_cp_rev_contrib_pct: float = cy_cp_rev / fInvoices.loc[(fInvoices['invoice_date'] <= end_date) & (
                fInvoices['invoice_date'] >= datetime(year=end_date.year, month=end_date.month,
                                                      day=1)), 'amount'].sum() * 100
        cy_ytd_rev_contrib_pct: float = cy_ytd_rev / fInvoices.loc[(fInvoices['invoice_date'] <= end_date) & (
                fInvoices['invoice_date'] >= datetime(year=end_date.year, month=1,
                                                      day=1)), 'amount'].sum() * 100
        monthly_rev: pd.DataFrame = fInvoices.loc[
            (fInvoices['cus_name'] == customer) & (fInvoices['invoice_date'] <= end_date) & (
                    fInvoices['invoice_date'] >= datetime(year=end_date.year, month=1, day=1)), ['invoice_date',
                                                                                                 'amount']].groupby(
            by=['invoice_date']).sum()
        monthly_rev.reset_index(inplace=True)
        monthly_rev.rename(columns={'invoice_date': 'Month', 'amount': 'Net Rev'}, inplace=True)
        monthly_rev.set_index(keys='Month', drop=True, inplace=True)
        if database == 'nbn_logistics':
            customer_code: str = dCustomer.loc[dCustomer['cus_name'] == customer, 'customer_code'].iloc[0]
            gross_rev: pd.DataFrame = fLogInv.loc[
                (fLogInv['invoice_date'] >= datetime(year=end_date.year, month=1, day=1)) & (
                        fLogInv['customer_code'] == customer_code), ['invoice_date', 'amount']]
            gross_rev: pd.DataFrame = gross_rev.groupby(by='invoice_date', as_index=False)['amount'].sum()
            gross_rev.rename(columns={'amount': 'Gross Rev', 'invoice_date': 'Month'}, inplace=True)
            gross_rev.set_index(keys='Month', inplace=True)
            monthly_rev = pd.concat([gross_rev, monthly_rev], axis=1)

        ageing: pd.DataFrame = cust_ageing(customer=customer, dCustomer=dCustomer, fCollection=fCollection,
                                           end_date=end_date)
        last_sales_person: str = fInvoices.loc[
            (fInvoices['invoice_date'] <= end_date) & (fInvoices['cus_name'] == customer), 'emp_id'].tail(
            1).iloc[0]
        last_sales_person = dEmployee.loc[dEmployee['emp_id'] == last_sales_person, 'emp_name'].iloc[0]
        stats: dict = {
            'customer_since': "Not Applicable" if customer_since == "Not Applicable" else customer_since.strftime(
                '%d-%m-%Y'),
            'total_revenue': total_revenue,
            'credit_score': 0,
            'last_receipt_amt': 0 if last_receipt_dt == "Not Collected" else last_receipt_amt,
            'cy_cp_rev': cy_cp_rev,
            'cy_pp_rev': cy_pp_rev,
            'last_receipt_dt': "Not Collected" if last_receipt_dt == "Not Collected" else last_receipt_dt.strftime(
                '%d-%m-%Y'),
            'cy_ytd_rev': cy_ytd_rev, 'py_cp_rev': py_cp_rev, 'py_ytd_rev': py_ytd_rev,
            'collection_median': "Not Collected" if last_receipt_dt == "Not Collected" else collection_median.days,
            'credit_days': credit_days, 'last_sales_person': last_sales_person,
            'customer_gp_cp': 0, 'outstanding_bal': -outstanding_bal, 'ageing': ageing,
            'date_established': date_established.strftime('%d-%m-%Y'),
            'cy_cp_rev_contrib_pct': f'{round(cy_cp_rev_contrib_pct, 1)}%',
            'cy_ytd_rev_contrib_pct': f'{round(cy_ytd_rev_contrib_pct, 1)}%',
            'cy_cp_roi': 0, 'customer_gp_ytd': 0,
            'cy_ytd_roi': 0, 'monthly_rev': monthly_rev, 'remarks': 0}
        customer_info[customer] = stats
    return customer_info


def cell_background(table, row: int, column: list, original: float, compare: float, good: str, bad: str):
    result = good if original >= compare else bad
    for idx, cell in enumerate(table.rows[row].cells):
        if idx in column:
            cell_xml_element = cell._tc
            table_cell_properties = cell_xml_element.get_or_add_tcPr()
            shade_obj = OxmlElement('w:shd')
            shade_obj.set(qn('w:fill'), result)
            table_cell_properties.append(shade_obj)


def expenses_allocation(orders: list, fData: pd.DataFrame, end_date: datetime, start_date: datetime) -> pd.DataFrame:
    """Calculate gross profit (excluding direct salaries) generated by all four types of revenue streams for a given job

    Args:
        job_id (str): a job number of an invoice which is partially or fully outstanding on a given date. 

    Returns:
        float: gross profit without salaries for a given job
    """
    # revenue ledgers and direct logistics expenses ledgers including "Logistics - Others"
    filt = (fData['order_id'].isin(orders)) & (fData['ledger_code'].isin(
        [5010202001, 5010202002, 5010202003, 5010202004, 4010201001, 4010201002, 4010201003, 4010201004])) & (
                   fData['voucher_date'] >= start_date) & (fData['voucher_date'] <= end_date)
    exp_allo: pd.DataFrame = fData.loc[filt, ['order_id', 'credit', 'debit']]
    exp_allo.loc[:, 'amount'] = fData['credit'] - fData['debit']
    exp_allo = exp_allo.groupby(by='order_id', as_index=False)['amount'].sum()
    return exp_allo


def job_revenue(orders: list, fData: pd.DataFrame, end_date: datetime, start_date: datetime) -> float:
    """ "Net" revenue to which each job has made for a give period

    Args:
        row (_type_): each row of the report

    Returns:
        float: Net revenue of each job
    """
    filt = (fData['order_id'].isin(orders)) & (fData['ledger_code'].isin(
        [4010201001, 4010201002, 4010201003, 4010201004])) & (fData['voucher_date'] >= start_date) & (
                   fData['voucher_date'] <= end_date)
    job_rev: pd.DataFrame = fData.loc[filt, ['order_id', 'credit', 'debit']]
    job_rev[:, 'amount'] = fData['credit'] - fData['debit']
    job_rev = job_rev.groupby(by='order_id', as_index=False)['amount'].sum()
    return job_rev


def gross_revenue(dCoAAdler: pd.DataFrame, fData: pd.DataFrame,
                  end_date: datetime) -> float:  # this function was never used. can be deleted.
    """Returns gross revenue reported for the whole period for a given job

    Args:
        row (_type_): a given row

    Returns:
        float: gross revenue
    """

    # amount net of amount debited or credited for a given customer is considered as gross revenue. 
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    jobs_list_filt = (fData['voucher_date'] >= start_date) & (fData['voucher_date'] <= end_date) & (
            fData['type'] == 'SI')
    jobs_list: np.ndarray = fData.loc[jobs_list_filt, 'order_id'].unique()
    customers: list = dCoAAdler.loc[dCoAAdler['second_level'].isin(
        ['Trade Receivables', 'Due from Related Parties', 'Deleted']), 'ledger_code'].to_list()
    filt = (fData['ledger_code'].isin(customers)) & (fData['order_id'].isin(jobs_list))
    gross_rev: pd.DataFrame = fData.loc[filt, ['order_id', 'credit', 'debit']]
    gross_rev.loc[:, 'amount'] = fData['credit'] - fData['debit']
    gross_rev = gross_rev.groupby(by='order_id', as_index=False)['amount'].sum()
    return gross_rev


def invoices_raised(fData: pd.DataFrame, start_date: datetime, end_date: datetime) -> list:
    """Certain jobs contains more than one invoices raised in same or different periods. 

    Args:
        job_id (str): a given job number

    Returns:
        list: list of invoices raised for a given job. 
    """
    invoices = fData.loc[
        (fData['type'] == 'SI') & (fData['voucher_date'] >= start_date) & (fData['voucher_date'] <= end_date), [
            'order_id', 'voucher_number']].drop_duplicates().groupby(by='order_id')['voucher_number'].apply(
        list).to_dict()
    return invoices


def assign_category(row, exclusion: dict) -> str:
    if row in exclusion['ruwais']['customers']:
        return 'ruwais'
    elif row in exclusion['qafco']['customers']:
        return 'qafco'
    else:
        return 'general'


def assign_employee(row, exclusion: dict) -> str:
    if row in exclusion['ruwais']['staff']:
        return 'ruwais'
    elif row in exclusion['qafco']['staff']:
        return 'qafco'
    else:
        return 'general'


def total_revenue(fData: pd.DataFrame, dLogContract: pd.DataFrame, exclusion: dict) -> dict:
    fData: pd.DataFrame = fData.loc[(fData['voucher_date'] >= datetime(year=2022, month=1, day=1))]
    fData.loc[:, 'amount'] = fData['credit'] - fData['debit']
    fData = pd.merge(left=fData, right=dLogContract, on='order_id', how='left')
    fData['voucher_date'] = fData['voucher_date'] + pd.offsets.MonthEnd(0)
    fData.loc[:, 'category'] = fData['customer_code'].apply(assign_category, args=[exclusion])
    df_for_rev: pd.DataFrame = fData.copy()
    df_for_rev = df_for_rev.loc[df_for_rev['ledger_code'].isin([4010201001, 4010201002, 4010201003, 4010201004])]
    overall_rev: pd.DataFrame = df_for_rev.pivot_table(columns='ledger_code', values='amount', aggfunc='sum',
                                                       fill_value=0,
                                                       index=['category',
                                                              pd.Grouper(key='voucher_date', freq='ME')]).reset_index()
    overall_rev.rename(
        columns={4010201001: 'custom', 4010201002: 'transport', 4010201003: 'freight', 4010201004: 'other'},
        inplace=True)
    overall_rev.loc[:, 'total_rev'] = overall_rev['custom'] + overall_rev['transport'] + overall_rev['freight'] + \
                                      overall_rev['other']
    invoice_rev: pd.DataFrame = df_for_rev.pivot_table(columns='ledger_code', values='amount', aggfunc='sum',
                                                       index=['voucher_date', 'voucher_number', 'order_id', 'category'],
                                                       fill_value=0).reset_index()
    invoice_rev.rename(
        columns={4010201001: 'custom', 4010201002: 'transport', 4010201003: 'freight', 4010201004: 'other'},
        inplace=True)
    filt = fData['ledger_code'].isin(
        [5010202001, 5010202002, 5010202003, 5010202004, 4010201001, 4010201002, 4010201003, 4010201004]) & (
                   fData['voucher_date'] >= datetime(year=2022, month=1, day=1))
    profit_wo_salary: pd.DataFrame = fData.loc[filt, ['order_id', 'credit', 'debit', 'voucher_date', 'category']]
    profit_wo_salary.loc[:, 'prwosal'] = profit_wo_salary['credit'] - profit_wo_salary['debit']
    profit_wo_salary = \
        profit_wo_salary.groupby(by=[pd.Grouper(key='voucher_date', freq='ME'), 'order_id', 'category'],
                                 as_index=False)[
            'prwosal'].sum()

    return {'overall_rev': overall_rev, 'invoice_rev': invoice_rev, 'profit_wo_salary': profit_wo_salary}


def salary_cost(fGL: pd.DataFrame, dEmployee: pd.DataFrame, exclusion: dict) -> float:
    sal_cost: pd.DataFrame = fGL.loc[
        fGL['cost_center'].isin(dEmployee['emp_id'].unique()) & (fGL['first_level'] == 'Staff Cost - Logistics') & (
                fGL['voucher_date'] >= datetime(year=2022, month=1, day=1))]
    sal_cost.loc[:, 'category'] = sal_cost['cost_center'].apply(assign_employee, args=[exclusion])
    sal_cost = sal_cost.pivot_table(index=['voucher_date', 'category'], columns='ledger_code', aggfunc='sum',
                                    fill_value=0,
                                    values='amount').reset_index()
    sal_cost.loc[:, 'custom'] = sal_cost[5010201001] + sal_cost[5010201004]
    sal_cost.loc[:, 'transport'] = sal_cost[5010201002] + sal_cost[5010201005]
    sal_cost.loc[:, 'freight'] = sal_cost[5010201003] + sal_cost[5010201006]
    sal_cost.drop(columns=[5010201001, 5010201002, 5010201003, 5010201004, 5010201005, 5010201006], inplace=True)
    return sal_cost


def salary_allocation(row, sal_allocation: pd.DataFrame, overall_rev: pd.DataFrame) -> pd.Series:
    voucher_date = row['voucher_date']
    custom, transport, freight = 0, 0, 0
    if row['category'] == 'ruwais':
        custom_overall: float = overall_rev.loc[
            (overall_rev['voucher_date'] == voucher_date) & (overall_rev['category'] == 'ruwais'), 'custom'].iloc[0]
        try:
            custom_allocation: float = sal_allocation.loc[(sal_allocation['voucher_date'] == voucher_date) & (
                    sal_allocation['category'] == 'ruwais'), 'custom'].iloc[0]
        except IndexError:
            custom_allocation = 0
        custom: float = row['custom'] / custom_overall * custom_allocation

        transport_overall: float = overall_rev.loc[
            (overall_rev['voucher_date'] == voucher_date) & (overall_rev['category'] == 'ruwais'), 'transport'].iloc[0]
        try:
            transport_allocation: float = sal_allocation.loc[(sal_allocation['voucher_date'] == voucher_date) & (
                    sal_allocation['category'] == 'ruwais'), 'transport'].iloc[0]
        except IndexError:
            transport_allocation = 0
        transport: float = row['transport'] / transport_overall * transport_allocation

        freight_overall: float = overall_rev.loc[
            (overall_rev['voucher_date'] == voucher_date) & (overall_rev['category'] == 'ruwais'), 'freight'].iloc[0]
        try:
            freight_allocation: float = sal_allocation.loc[(sal_allocation['voucher_date'] == voucher_date) & (
                    sal_allocation['category'] == 'ruwais'), 'freight'].iloc[0]
        except IndexError:
            freight_allocation = 0
        freight: float = row['freight'] / freight_overall * freight_allocation
    elif row['category'] == 'qafco':

        custom_overall: float = overall_rev.loc[(overall_rev['voucher_date'] == voucher_date) & (
            overall_rev['category'].isin(['qafco', 'general'])), 'custom'].sum()
        try:
            custom_allocation: float = sal_allocation.loc[(sal_allocation['voucher_date'] == voucher_date) & (
                sal_allocation['category'].isin(['qafco', 'general'])), 'custom'].sum()
        except IndexError:
            custom_allocation = 0
        custom: float = row['custom'] / custom_overall * custom_allocation

        transport_overall: float = overall_rev.loc[(overall_rev['voucher_date'] == voucher_date) & (
            overall_rev['category'].isin(['qafco', 'general'])), 'transport'].sum()
        try:
            transport_allocation: float = sal_allocation.loc[(sal_allocation['voucher_date'] == voucher_date) & (
                sal_allocation['category'].isin(['qafco', 'general'])), 'transport'].sum()
        except IndexError:
            transport_allocation = 0
        transport: float = row['transport'] / transport_overall * transport_allocation

        freight_overall: float = overall_rev.loc[
            (overall_rev['voucher_date'] == voucher_date) & (overall_rev['category'].isin(['qafco'])), 'freight'].iloc[
            0]
        try:
            freight_allocation: float = sal_allocation.loc[(sal_allocation['voucher_date'] == voucher_date) & (
                sal_allocation['category'].isin(['qafco'])), 'freight'].iloc[0]
        except IndexError:
            freight_allocation = 0
        freight: float = row['freight'] / freight_overall * freight_allocation
    else:
        custom_overall: float = overall_rev.loc[(overall_rev['voucher_date'] == voucher_date) & (
            overall_rev['category'].isin(['general', 'qafco'])), 'custom'].sum()
        try:
            custom_allocation: float = sal_allocation.loc[(sal_allocation['voucher_date'] == voucher_date) & (
                sal_allocation['category'].isin(['general', 'qafco'])), 'custom'].sum()
        except IndexError:
            custom_allocation = 0
        custom: float = row['custom'] / custom_overall * custom_allocation

        transport_overall: float = overall_rev.loc[(overall_rev['voucher_date'] == voucher_date) & (
            overall_rev['category'].isin(['general', 'qafco'])), 'transport'].sum()
        try:
            transport_allocation: float = sal_allocation.loc[(sal_allocation['voucher_date'] == voucher_date) & (
                sal_allocation['category'].isin(['general', 'qafco'])), 'transport'].sum()
        except IndexError:
            transport_allocation = 0
        transport: float = row['transport'] / transport_overall * transport_allocation

        freight_overall: float = overall_rev.loc[
            (overall_rev['voucher_date'] == voucher_date) & (overall_rev['category'] == 'general'), 'freight'].iloc[0]
        try:
            freight_allocation: float = sal_allocation.loc[(sal_allocation['voucher_date'] == voucher_date) & (
                    sal_allocation['category'] == 'general'), 'freight'].iloc[0]
        except IndexError:
            freight_allocation = 0
        freight: float = row['freight'] / freight_overall * freight_allocation

    return pd.Series({'custom_sal': custom, 'transport_sal': transport, 'freight_sal': freight})


def initial_profit(fGL: pd.DataFrame, invoice_values: dict, dEmployee: pd.DataFrame, exclusion: dict) -> float:
    sal_allocation: pd.DataFrame = salary_cost(fGL=fGL, dEmployee=dEmployee, exclusion=exclusion)

    overall_rev: pd.DataFrame = invoice_values['overall_rev']

    invoice_values['invoice_rev'][['custom_sal', 'transport_sal', 'freight_sal']] = invoice_values['invoice_rev'].apply(
        salary_allocation, args=[sal_allocation, overall_rev], axis=1)
    invoice_rev: pd.DataFrame = invoice_values['invoice_rev'].groupby(by=['voucher_date', 'order_id', 'category'],
                                                                      as_index=False).sum(
        ['custom', 'transport', 'freight', 'other', 'custom_sal', 'transport_sal', 'freight_sal'])
    prwosal: pd.DataFrame = invoice_values['profit_wo_salary']
    prwosal = prwosal.set_index(keys=['voucher_date', 'order_id'])
    invoice_rev = invoice_rev.set_index(keys=['voucher_date', 'order_id'])
    invoice_rev = pd.concat([invoice_rev, prwosal], axis=1)
    invoice_rev.fillna(value=0, inplace=True)
    invoice_rev.reset_index(inplace=True)
    return invoice_rev


def overhead_allocation_nbnl(row, overhead_monthly: pd.DataFrame, overall_rev: pd.DataFrame):
    revenue: float = row['total_rev']
    voucher_date: datetime = row['voucher_date']
    overhead: float = overhead_monthly.loc[overhead_monthly['voucher_date'] == voucher_date, 'amount'].iloc[0]
    monthly_rev: float = overall_rev.loc[overall_rev['voucher_date'] == voucher_date, 'total_rev'].sum()
    return revenue / monthly_rev * overhead


def job_profitability(fTimesheet: pd.DataFrame, fGL: pd.DataFrame, end_date: datetime, dEmployee: pd.DataFrame,
                      dExclude: pd.DataFrame, fOT: pd.DataFrame, fInvoices: pd.DataFrame, cogs_map: dict,
                      dJobs: pd.DataFrame, database, fData: pd.DataFrame) -> pd.DataFrame:
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    periods: list = pd.date_range(start=start_date, end=end_date, freq='ME').to_pydatetime().tolist()
    fGL = fGL.loc[:,
          ['cost_center', 'voucher_date', 'ledger_code', 'amount', 'third_level', 'second_level', 'first_level']]
    cy_cp_cus_np = None
    cy_cp_emp_np = None
    cy_ytd_emp_np = None
    cy_ytd_cus_np = None

    if database == 'elite_security':
        fGL = fGL.loc[~fGL['ledger_code'].isin([5010101002, 5010101003])]
        emp_list_full: list = dEmployee.index.tolist()
        driversandcleaners: list = dEmployee.loc[
            dEmployee['designation'].isin(['HEAVY DRIVER', 'DRIVER', 'CAMP SUPERVISOR'])].index.tolist()
        emp_list: list = [i for i in emp_list_full if i not in driversandcleaners]
        timesheet_sum: dict = {'dc_emp_beni': None, 'dc_trpt': None, 'dc_out': None, 'dc_sal': None}
        timesheet_jobs: dict = {'dc_emp_beni': None, 'dc_trpt': None, 'dc_out': None, 'dc_sal': None}
        timesheet_grand_sum: dict = {'dc_emp_beni': None, 'dc_trpt': None, 'dc_out': None, 'dc_sal': None}
        periodic_allocation: dict = {}

        for period in periods:
            st_date: datetime = period + relativedelta(day=1)
            fGL_fitlered: pd.DataFrame = fGL.loc[(fGL['voucher_date'] >= st_date) & (fGL['voucher_date'] <= period) &
                                                 (fGL['second_level'] == 'Manpower Cost'), ['cost_center',
                                                                                            'voucher_date',
                                                                                            'ledger_code',
                                                                                            'amount']]
            fGL_emp: pd.DataFrame = fGL_fitlered.loc[fGL_fitlered['cost_center'].isin(emp_list)]
            fGL_other: pd.DataFrame = \
                fGL_fitlered.loc[~fGL_fitlered['cost_center'].isin(emp_list), ['amount', 'ledger_code']].groupby(
                    'ledger_code',
                    as_index=False)[
                    'amount'].sum()
            fGL_emp = fGL_emp.groupby(by=['cost_center', 'voucher_date', 'ledger_code'], as_index=False)['amount'].sum()
            fGL_emp = fGL_emp.loc[fGL_emp['amount'] != 0]
            # TODO You may group this to cogs map using the ledger code. to be fixed. it will reduce the no of iteretion by approx 12.5%
            fTimesheet_filtered: pd.DataFrame = fTimesheet.loc[
                (fTimesheet['v_date'] >= st_date) & (fTimesheet['v_date'] <= period)]
            fTimesheet_filtered = fTimesheet_filtered.groupby(['cost_center', 'order_id', 'v_date']).size().reset_index(
                name='count')
            billable_jobs: list = fTimesheet_filtered.loc[
                fTimesheet_filtered['order_id'].str.contains('ESS/CTR'), 'order_id'].unique().tolist()

            for c in dExclude.columns:
                if c not in ['job_type', 'group']:
                    valid_jobs: list = dExclude.loc[dExclude[c] == False]['job_type'].tolist() + billable_jobs
                    timesheet_sum[c] = \
                        fTimesheet_filtered.loc[fTimesheet_filtered['order_id'].isin(valid_jobs)].groupby(
                            ['cost_center', 'v_date'], as_index=False)['count'].sum()
                    timesheet_jobs[c] = fTimesheet_filtered.loc[fTimesheet_filtered['order_id'].isin(valid_jobs)]
                    timesheet_grand_sum[c] = timesheet_sum[c]['count'].sum()
            allocation_dict: dict = {}
            unallocated_amount: float = 0
            for _, i in fGL_emp.iterrows():
                df_type: str = [(k, v) for k, v in cogs_map.items() if i['ledger_code'] in v][0][0]
                # TODO (a) YOU MAY FILTER df_sum/timesheet_sum and timesheet_detailed/timesheet_jobs only for those cost_centers apperiring in fGL_Emp. which will reduce the number of iterations.
                # Also filter by the ledger as well 
                df_sum: pd.DataFrame = timesheet_sum[df_type]
                timesheet_detailed: pd.DataFrame = timesheet_jobs[df_type]
                try:
                    total_days: int = df_sum.loc[(df_sum['v_date'] == i['voucher_date']) & (
                            df_sum['cost_center'] == i['cost_center']), 'count'].iloc[0]
                    timesheet_detailed = timesheet_detailed.loc[(timesheet_detailed['v_date'] == i['voucher_date']) & (
                            timesheet_detailed['cost_center'] == i['cost_center']), ['order_id', 'count']]
                    allocation_dict_init = {}
                    for _, j in timesheet_detailed.iterrows():
                        # TODO (a) only those cost centers having a value will return a value from below. 
                        allocated: float = i['amount'] / total_days * j['count']
                        allocation_dict_init[j['order_id']] = allocated
                    allocation_dict = {k: allocation_dict_init.get(k, 0) + allocation_dict.get(k, 0) for k in
                                       set(allocation_dict) | set(allocation_dict_init)}
                except IndexError:
                    unallocated_amount += i['amount']
                    allocation_dict['Un-Allocated'] = unallocated_amount
            fOT_filtered: pd.DataFrame = fOT.loc[(fOT['voucher_date'] >= st_date) & (fOT['voucher_date'] <= period)]
            fOT_filtered: dict = fOT_filtered.groupby(by='order_id')['amount'].sum().to_dict()
            allocation_dict = {k: allocation_dict.get(k, 0) + fOT_filtered.get(k, 0) for k in
                               set(allocation_dict) | set(fOT_filtered)}
            inv_filtered_cust: dict = fInvoices.loc[
                (fInvoices['invoice_date'] >= st_date) & (fInvoices['invoice_date'] <= period), ['order_id',
                                                                                                 'amount']].groupby(
                'order_id')['amount'].sum().to_dict()
            allocation_dict = {k: allocation_dict.get(k, 0) + inv_filtered_cust.get(k, 0) for k in
                               set(allocation_dict) | set(inv_filtered_cust)}
            for i in cogs_map:
                z: float = fGL_other.loc[fGL_other['ledger_code'].isin(cogs_map[i])]['amount'].sum()
                if z != 0:
                    for _, row in timesheet_jobs[i].groupby(by='order_id', as_index=False)['count'].sum().iterrows():
                        overhead_allocation: dict = {}
                        value: float = z / timesheet_grand_sum[i] * row['count']
                        overhead_allocation[row['order_id']] = value
                        allocation_dict = {k: allocation_dict.get(k, 0) + overhead_allocation.get(k, 0) for k in
                                           set(allocation_dict) | set(overhead_allocation)}
            acc_types: list = dExclude.loc[dExclude['group'].isin(['Accommodation']), 'job_type'].tolist()
            accommodation_cost: float = sum([v for k, v in allocation_dict.items() if k in acc_types])
            non_accomo_sum: int = fTimesheet_filtered.loc[~fTimesheet_filtered['order_id'].isin(acc_types)][
                'count'].sum()
            non_accomo: pd.DataFrame = fTimesheet_filtered.loc[~fTimesheet_filtered['order_id'].isin(acc_types)]
            for _, row in non_accomo.iterrows():
                accommodation_allocation: dict = {}
                value: float = accommodation_cost / non_accomo_sum * row['count']
                accommodation_allocation[row['order_id']] = value
                allocation_dict = {k: allocation_dict.get(k, 0) + accommodation_allocation.get(k, 0) for k in
                                   set(allocation_dict) | set(accommodation_allocation)}

            if 'AC-ACCOMODATION' in allocation_dict:
                del allocation_dict['AC-ACCOMODATION']

            if 'AC' in allocation_dict:
                del allocation_dict['AC']
            periodic_allocation[period] = allocation_dict

        cy_cp: pd.DataFrame = pd.DataFrame(list(periodic_allocation[end_date].items()), columns=['order_id', 'amount'])
        cy_cp = pd.merge(left=cy_cp, right=dJobs[['order_id', 'customer_code', 'emp_id']], on='order_id', how='left')
        cy_cp_cus: pd.DataFrame = cy_cp.groupby(by='customer_code', as_index=False)['amount'].sum()
        cy_cp_emp: pd.DataFrame = cy_cp.groupby(by='emp_id', as_index=False)['amount'].sum()
        cy_ytd: pd.DataFrame = pd.DataFrame()
        for period in periods:
            month_df: pd.DataFrame = pd.DataFrame(list(periodic_allocation[period].items()),
                                                  columns=['order_id', 'amount'])
            month_df['voucher_date'] = period
            cy_ytd = pd.concat([month_df, cy_ytd])
        cy_ytd = pd.merge(left=cy_ytd, right=dJobs[['order_id', 'customer_code', 'emp_id']], on='order_id',
                          how='left')
        cy_ytd_cus: pd.DataFrame = \
            cy_ytd.groupby(by=[pd.Grouper(key='voucher_date', freq='ME'), 'customer_code'], as_index=False)[
                'amount'].sum()
        cy_ytd_emp: pd.DataFrame = cy_ytd.groupby(by='emp_id', as_index=False)['amount'].sum()
    elif database == 'premium':
        pass
    elif database == 'nbn_logistics':
        exclusion: dict = {
            'ruwais': {'customers': ['CUS0794', 'CUS0781', 'CUS0787', 'CUS0613', 'CUS0756', 'CUS0813', 'CUS0810'],
                       'staff': ['NBNL0095', 'NBNL0096', 'NBNL0106']},
            'qafco': {'customers': ['CUS0806'],
                      'staff': ['NBNL0108','NBNL0112']}}
        invoice_values: dict = total_revenue(dLogContract=dJobs, fData=fData, exclusion=exclusion)
        nbnl_profitability: pd.DataFrame = initial_profit(dEmployee=dEmployee, fGL=fGL, invoice_values=invoice_values,
                                                          exclusion=exclusion)
        nbnl_profitability.loc[:, 'total_rev'] = nbnl_profitability['custom'] + nbnl_profitability['transport'] + \
                                                 nbnl_profitability['freight'] + nbnl_profitability['other']
        overhead_monthly: pd.DataFrame = \
            fGL.loc[fGL['third_level'].isin(['Overhead', 'Finance Cost']), ['amount', 'voucher_date']].groupby(
                by='voucher_date', as_index=False)['amount'].sum()
        nbnl_profitability.loc[:, 'overhead'] = nbnl_profitability.apply(overhead_allocation_nbnl,
                                                                         args=[overhead_monthly,
                                                                               invoice_values['overall_rev']], axis=1)
        nbnl_profitability = pd.merge(left=nbnl_profitability, right=dJobs, on='order_id', how='left')
        nbnl_profitability.loc[:, 'amount'] = nbnl_profitability['prwosal'] + nbnl_profitability['custom_sal'] + \
                                              nbnl_profitability['transport_sal'] + nbnl_profitability['freight_sal']
        nbnl_profitability.loc[:, 'net_profit'] = nbnl_profitability['amount'] + nbnl_profitability['overhead']
        nbnl_profitability.to_csv('nbnl_profitability_aug.csv')
        filt_for_month = (nbnl_profitability['voucher_date'] <= end_date) & (
                nbnl_profitability['voucher_date'] >= datetime(year=end_date.year, month=end_date.month, day=1))

        cy_cp_cus: pd.DataFrame = nbnl_profitability.loc[filt_for_month, ['customer_code', 'amount']]
        cy_cp_cus = cy_cp_cus.groupby(by='customer_code', as_index=False)['amount'].sum()

        cy_cp_emp: pd.DataFrame = nbnl_profitability.loc[filt_for_month, ['emp_id', 'amount']]
        cy_cp_emp = cy_cp_emp.groupby(by='emp_id', as_index=False)['amount'].sum()

        cy_cp_cus_np: pd.DataFrame = nbnl_profitability.loc[filt_for_month, ['customer_code', 'net_profit']]
        cy_cp_cus_np = cy_cp_cus_np.groupby(by='customer_code', as_index=False)['net_profit'].sum()

        cy_cp_emp_np: pd.DataFrame = nbnl_profitability.loc[filt_for_month, ['emp_id', 'net_profit']]
        cy_cp_emp_np = cy_cp_emp_np.groupby(by='emp_id', as_index=False)['net_profit'].sum()

        filt_for_ytd = (nbnl_profitability['voucher_date'] >= start_date) & (
                nbnl_profitability['voucher_date'] <= end_date)

        cy_ytd_cus: pd.DataFrame = nbnl_profitability.loc[filt_for_ytd, ['voucher_date', 'customer_code', 'amount']]
        cy_ytd_cus = \
            cy_ytd_cus.groupby(by=[pd.Grouper(key='voucher_date', freq='ME'), 'customer_code'], as_index=False)[
                'amount'].sum()

        cy_ytd_emp: pd.DataFrame = nbnl_profitability.loc[filt_for_ytd, ['voucher_date', 'emp_id', 'amount']]
        cy_ytd_emp = cy_ytd_emp.groupby(by=[pd.Grouper(key='voucher_date', freq='ME'), 'emp_id'], as_index=False)[
            'amount'].sum()

        cy_ytd_cus_np: pd.DataFrame = nbnl_profitability.loc[
            filt_for_ytd, ['voucher_date', 'customer_code', 'net_profit']]
        cy_ytd_cus_np = \
            cy_ytd_cus_np.groupby(by=[pd.Grouper(key='voucher_date', freq='ME'), 'customer_code'], as_index=False)[
                'net_profit'].sum()

        cy_ytd_emp_np: pd.DataFrame = nbnl_profitability.loc[filt_for_ytd, ['voucher_date', 'emp_id', 'net_profit']]
        cy_ytd_emp_np = cy_ytd_emp_np.groupby(by=[pd.Grouper(key='voucher_date', freq='ME'), 'emp_id'], as_index=False)[
            'net_profit'].sum()

        periodic_allocation: dict = nbnl_profitability.loc[filt_for_ytd, ['voucher_date', 'order_id', 'amount']]
        periodic_allocation = periodic_allocation.groupby('voucher_date').apply(
            lambda g: dict(zip(g['order_id'], g['amount']))).to_dict()

    else:
        pass
    return {'periodic_allocation': periodic_allocation, 'cy_cp_cus': cy_cp_cus, 'cy_ytd_cus': cy_ytd_cus,
            'cy_cp_emp': cy_cp_emp, 'cy_ytd_emp': cy_ytd_emp, 'cy_ytd_emp_np': cy_ytd_emp_np,
            'cy_ytd_cus_np': cy_ytd_cus_np, 'cy_cp_cus_np': cy_cp_cus_np, 'cy_cp_emp_np': cy_cp_emp_np}


def salespersonstats(document, salesperson_list: list, dEmployee: pd.DataFrame, fInvoices: pd.DataFrame,
                     profitability: dict, end_date: datetime, fGL: pd.DataFrame):
    salesperson_stats: dict = sales_person(emp_ids=salesperson_list, dEmployee=dEmployee, fInvoices=fInvoices,
                                           end_date=end_date, fGL=fGL)
    for idx, salesperson in enumerate(salesperson_list):
        if (idx + 1) % 2 == 0:
            document.add_paragraph('\n\n\n')
        salesperson_name: str = ' '.join(
            dEmployee.loc[dEmployee['emp_id'] == salesperson, 'emp_name'].iloc[0].split(sep=' ')[:2]).title()
        salutation: str = "Mr." if dEmployee.loc[dEmployee['emp_id'] == salesperson, 'sex'].iloc[0] == 'Male' else "Ms."
        full_name: str = f'{salutation}{salesperson_name}'
        status: bool = dEmployee.loc[dEmployee['emp_id'] == salesperson, 'termination_date'].iloc[0] < datetime(
            year=end_date.year, month=end_date.month, day=1)
        cy_cp_pl_company_title = document.add_paragraph().add_run(full_name)
        apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
        tbl_salesman_main_1 = document.add_table(rows=2, cols=2)
        tbl_salesman_main_th_1 = tbl_salesman_main_1.rows[0]
        tbl_salesman_main_th_1.cells[0].text = 'Date of Join'
        tbl_salesman_main_th_1.cells[1].text = 'New Customers Added'

        tbl_salesman_main_td_1 = tbl_salesman_main_1.rows[1]
        tbl_salesman_main_td_1.cells[0].text = str(salesperson_stats[salesperson]['doj'])
        tbl_salesman_main_td_1.cells[1].text = 'Resigned Employee' if status else number_format(
            num=salesperson_stats[salesperson]['new_customers_added'])
        table_formatter(table_name=tbl_salesman_main_1, style_name='table_style_1', special=[])

        tbl_salesman_main_2 = document.add_table(rows=2, cols=2)
        tbl_salesman_main_th_2 = tbl_salesman_main_2.rows[0]
        tbl_salesman_main_th_2.cells[0].text = 'CP Target'
        tbl_salesman_main_th_2.cells[1].text = 'YTD Target'

        tbl_salesman_main_td_2 = tbl_salesman_main_2.rows[1]
        tbl_salesman_main_td_2.cells[0].text = 'Resigned Employee' if status else number_format(
            num=salesperson_stats[salesperson]['cp_target'])
        tbl_salesman_main_td_2.cells[1].text = 'Resigned Employee' if status else number_format(
            num=salesperson_stats[salesperson]['ytd_target'])
        table_formatter(table_name=tbl_salesman_main_2, style_name='table_style_1', special=[])

        tbl_salesman_rev_1 = document.add_table(rows=2, cols=2)
        tbl_salesman_rev_th_1 = tbl_salesman_rev_1.rows[0]
        tbl_salesman_rev_th_1.cells[0].text = 'CY CP Revenue'
        tbl_salesman_rev_th_1.cells[1].text = 'CY YTD Revenue'

        tbl_salesman_rev_td_1 = tbl_salesman_rev_1.rows[1]
        tbl_salesman_rev_td_1.cells[0].text = number_format(salesperson_stats[salesperson]['cy_cp_rev'])
        tbl_salesman_rev_td_1.cells[1].text = number_format(salesperson_stats[salesperson]['cy_ytd_rev'])
        table_formatter(table_name=tbl_salesman_rev_1, style_name='table_style_1', special=[])

        tbl_salesman_rev_2 = document.add_table(rows=2, cols=2)
        tbl_salesman_rev_th_2 = tbl_salesman_rev_2.rows[0]
        tbl_salesman_rev_th_2.cells[0].text = 'CY CP Own\nRevenue'
        tbl_salesman_rev_th_2.cells[1].text = 'CY YTD Own\nRevenue'

        tbl_salesman_rev_td_2 = tbl_salesman_rev_2.rows[1]
        tbl_salesman_rev_td_2.cells[0].text = number_format(num=salesperson_stats[salesperson]['cy_cp_rev_org'])
        tbl_salesman_rev_td_2.cells[1].text = number_format(salesperson_stats[salesperson]['cy_ytd_rev_org'])
        table_formatter(table_name=tbl_salesman_rev_2, style_name='table_style_1', special=[])

        tbl_salesman_gp_1 = document.add_table(rows=2, cols=2)
        tbl_salesman_gp_th_1 = tbl_salesman_gp_1.rows[0]
        tbl_salesman_gp_th_1.cells[0].text = 'CY CP GP'
        tbl_salesman_gp_th_1.cells[1].text = 'CY YTD GP'

        cy_cp_gp: float = profitability['cy_cp_emp'].loc[
            profitability['cy_cp_emp']['emp_id'] == salesperson, 'amount'].sum()
        cy_ytd_gp: float = profitability['cy_ytd_emp'].loc[
            profitability['cy_ytd_emp']['emp_id'] == salesperson, 'amount'].sum()

        tbl_salesman_gp_td_1 = tbl_salesman_gp_1.rows[1]
        tbl_salesman_gp_td_1.cells[0].text = number_format(num=cy_cp_gp)
        tbl_salesman_gp_td_1.cells[1].text = number_format(num=cy_ytd_gp)
        table_formatter(table_name=tbl_salesman_gp_1, style_name='table_style_1', special=[])

        tbl_salesman_gp_2 = document.add_table(rows=2, cols=2)
        tbl_salesman_gp_th_2 = tbl_salesman_gp_2.rows[0]
        tbl_salesman_gp_th_2.cells[0].text = 'CY CP Revenue\nContribution'
        tbl_salesman_gp_th_2.cells[1].text = 'CY YTD Revenue\nContribution'

        tbl_salesman_gp_td_2 = tbl_salesman_gp_2.rows[1]
        tbl_salesman_gp_td_2.cells[0].text = str(salesperson_stats[salesperson]['cy_cp_rev_contrib_pct'])
        tbl_salesman_gp_td_2.cells[1].text = str(salesperson_stats[salesperson]['cy_ytd_rev_contrib_pct'])
        table_formatter(table_name=tbl_salesman_gp_2, style_name='table_style_1', special=[])
        if (idx + 1) % 2 == 0:
            document.add_page_break()


def customer_specifics(document, fInvoices: pd.DataFrame, end_date: datetime, dCustomer: pd.DataFrame,
                       dJobs: pd.DataFrame, path, fCollection: pd.DataFrame, dEmployee: pd.DataFrame,
                       fTimesheet: pd.DataFrame, fOT: pd.DataFrame, dExclude: pd.DataFrame, fGL: pd.DataFrame,
                       database: str, fData: pd.DataFrame, fLogInv: pd.DataFrame):
    customer_list: list = sorted(fInvoices.loc[(fInvoices['invoice_date'] >= datetime(year=end_date.year,
                                                                                      month=end_date.month, day=1)) & (
                                                       fInvoices['invoice_date'] <= end_date), 'cus_name'].unique())
    customer_info: dict = customer_ratios(customers=customer_list, fInvoices=fInvoices, end_date=end_date,
                                          fCollection=fCollection, dCustomer=dCustomer, dEmployee=dEmployee, fGL=fGL,
                                          database=database, fLogInv=fLogInv)
    profitability: dict = job_profitability(fTimesheet=fTimesheet, fGL=fGL, end_date=end_date, dEmployee=dEmployee,
                                            dExclude=dExclude, fOT=fOT, fInvoices=fInvoices, cogs_map=cogs_ledger_map,
                                            dJobs=dJobs, database=database, fData=fData)
    heading_format = {'fontfamily': 'Georgia', 'color': 'k', 'fontweight': 'bold', 'fontsize': 10}
    cy_cp_profit_cus: pd.DataFrame = profitability['cy_cp_cus']
    cy_ytd_profit_cus: pd.DataFrame = profitability['cy_ytd_cus']
    cy_cp_net_profit_cus: pd.DataFrame = profitability['cy_cp_cus_np']
    cy_ytd_net_profit_cus: pd.DataFrame = profitability['cy_ytd_cus_np']

    for customer in customer_list:
        cus_code: list = dCustomer.loc[(dCustomer['cus_name'] == customer), 'customer_code'].tolist()
        table_title = document.add_table(rows=1, cols=2)
        table_title.columns[0].width = Cm(12)
        r0c0 = table_title.cell(0, 0)
        cy_cp_pl_company_title = r0c0.add_paragraph().add_run(customer.upper())
        apply_style_properties(cy_cp_pl_company_title, style_picker(name='company_title'))
        r0c1 = table_title.cell(0, 1)
        cust_logo = r0c1.add_paragraph().add_run()
        try:
            logo = cust_logo.add_picture(f'{path}\{cus_code[0]}.png', width=Inches(0.79), height=Inches(1))
            logo = document.paragraphs[-1]
            logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except FileNotFoundError:
            logo = cust_logo.add_picture(f'{path}\default.png')
            logo = document.paragraphs[-1]
            logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        tbl_cust_main = document.add_table(rows=2, cols=4)
        tbl_cust_main_th = tbl_cust_main.rows[0]
        tbl_cust_main_th.cells[0].text = 'Date of Establishment'
        tbl_cust_main_th.cells[1].text = 'Customer Since'
        tbl_cust_main_th.cells[2].text = 'Salesperson'
        tbl_cust_main_th.cells[3].text = 'Balance'
        tbl_cust_main_td = tbl_cust_main.rows[1]
        tbl_cust_main_td.cells[0].text = str(customer_info[customer]['date_established'])
        tbl_cust_main_td.cells[1].text = str(customer_info[customer]['customer_since'])
        tbl_cust_main_td.cells[2].text = ' '.join(
            str(customer_info[customer]['last_sales_person']).split(sep=' ')[:2]).title()
        tbl_cust_main_td.cells[3].text = number_format(num=customer_info[customer]['outstanding_bal'])
        table_formatter(table_name=tbl_cust_main, style_name='table_style_1', special=[])

        tbl_cust_rev_1 = document.add_table(rows=2, cols=4)
        tbl_cust_rev_th_1 = tbl_cust_rev_1.rows[0]
        tbl_cust_rev_th_1.cells[0].text = f'CY CP Revenue\n({end_date.strftime("%B")} Month)'
        tbl_cust_rev_th_1.cells[1].text = 'CY YTD Revenue'
        tbl_cust_rev_th_1.cells[2].text = 'CY CP Rev \nContribution'
        tbl_cust_rev_th_1.cells[
            3].text = f'CY PP Revenue\n({(end_date.replace(day=1) - timedelta(days=1)).strftime("%B")} Month)'
        tbl_cust_rev_td_1 = tbl_cust_rev_1.rows[1]
        arrow_type_cy: str = '⇑' if customer_info[customer]['cy_cp_rev'] >= customer_info[customer][
            'py_cp_rev'] else '⇓'
        arrow_type_ytd: str = '⇑' if customer_info[customer]['cy_ytd_rev'] >= customer_info[customer][
            'py_ytd_rev'] else '⇓'
        tbl_cust_rev_td_1.cells[0].text = f"{number_format(num=customer_info[customer]['cy_cp_rev'])} {arrow_type_cy}"
        tbl_cust_rev_td_1.cells[1].text = f"{number_format(num=customer_info[customer]['cy_ytd_rev'])} {arrow_type_ytd}"
        tbl_cust_rev_td_1.cells[2].text = str(customer_info[customer]['cy_cp_rev_contrib_pct'])
        tbl_cust_rev_td_1.cells[3].text = number_format(num=customer_info[customer]['cy_pp_rev'])
        table_formatter(table_name=tbl_cust_rev_1, style_name='table_style_1', special=[])

        cell_background(table=tbl_cust_rev_1, row=1, column=[0], original=customer_info[customer]['cy_cp_rev'],
                        compare=customer_info[customer]['py_cp_rev'], good='d2f3cf', bad='ffd8d5')
        cell_background(table=tbl_cust_rev_1, row=1, column=[1], original=customer_info[customer]['cy_ytd_rev'],
                        compare=customer_info[customer]['py_ytd_rev'], good='d2f3cf', bad='ffd8d5')

        tbl_cust_rev_2 = document.add_table(rows=2, cols=4)
        tbl_cust_rev_th_2 = tbl_cust_rev_2.rows[0]
        tbl_cust_rev_th_2.cells[0].text = 'PY CP Revenue'
        tbl_cust_rev_th_2.cells[1].text = 'PY YTD Revenue'
        tbl_cust_rev_th_2.cells[2].text = 'CY YTD Rev \nContribution'
        tbl_cust_rev_th_2.cells[3].text = 'Total Revenue Made'
        tbl_cust_rev_td_2 = tbl_cust_rev_2.rows[1]

        tbl_cust_rev_td_2.cells[0].text = number_format(num=customer_info[customer]['py_cp_rev'])
        tbl_cust_rev_td_2.cells[1].text = number_format(num=customer_info[customer]['py_ytd_rev'])
        tbl_cust_rev_td_2.cells[2].text = str(customer_info[customer]['cy_ytd_rev_contrib_pct'])
        tbl_cust_rev_td_2.cells[3].text = number_format(num=customer_info[customer]['total_revenue'])
        table_formatter(table_name=tbl_cust_rev_2, style_name='table_style_1', special=[])

        tbl_cust_col = document.add_table(rows=2, cols=4)
        tbl_cust_col_th = tbl_cust_col.rows[0]
        tbl_cust_col_th.cells[0].text = 'Credit Score'
        tbl_cust_col_th.cells[1].text = 'Credit Days'
        tbl_cust_col_th.cells[2].text = 'Median Collection\nDays'
        tbl_cust_col_th.cells[3].text = 'Last Collection Date\n and Amount'
        tbl_cust_col_td = tbl_cust_col.rows[1]
        tbl_cust_col_td.cells[0].text = number_format(customer_info[customer]['credit_score'])
        tbl_cust_col_td.cells[1].text = number_format(num=customer_info[customer]['credit_days'])
        tbl_cust_col_td.cells[2].text = str(customer_info[customer]['collection_median'])

        tbl_cust_col_td.cells[
            3].text = f"{number_format(num=customer_info[customer]['last_receipt_amt'])}\n{str(customer_info[customer]['last_receipt_dt'])}"
        table_formatter(table_name=tbl_cust_col, style_name='table_style_1', special=[])

        tbl_cust_gp = document.add_table(rows=2, cols=4)
        tbl_cust_gp_th = tbl_cust_gp.rows[0]
        tbl_cust_gp_th.cells[0].text = 'Profit Month'
        tbl_cust_gp_th.cells[1].text = 'Profit YTD'
        tbl_cust_gp_th.cells[2].text = 'ROI Month'
        tbl_cust_gp_th.cells[3].text = 'ROI YTD'
        tbl_cust_gp_td = tbl_cust_gp.rows[1]
        cp_gp_pct = round(cy_cp_profit_cus.loc[cy_cp_profit_cus['customer_code'].isin(cus_code), 'amount'].sum() /
                          customer_info[customer]['cy_cp_rev'] * 100, 2)
        ytd_gp_pct = round(cy_ytd_profit_cus.loc[cy_ytd_profit_cus['customer_code'].isin(cus_code), 'amount'].sum() /
                           customer_info[customer]['cy_ytd_rev'] * 100, 2)

        monthly_rev: pd.DataFrame = customer_info[customer]['monthly_rev']
        monthly_rev.reset_index(inplace=True)

        if database == 'nbn_logistics':
            cp_np_pct = round(
                cy_cp_net_profit_cus.loc[cy_cp_net_profit_cus['customer_code'].isin(cus_code), 'net_profit'].sum() /
                customer_info[customer]['cy_cp_rev'] * 100, 2)
            ytd_np_pct = round(
                cy_ytd_net_profit_cus.loc[cy_ytd_net_profit_cus['customer_code'].isin(cus_code), 'net_profit'].sum() /
                customer_info[customer]['cy_ytd_rev'] * 100, 2)
            tbl_cust_gp_td.cells[
                0].text = f"GP-{number_format(num=cy_cp_profit_cus.loc[cy_cp_profit_cus['customer_code'].isin(cus_code), 'amount'].sum())} | {cp_gp_pct}% \nNP-{number_format(num=cy_cp_net_profit_cus.loc[cy_cp_net_profit_cus['customer_code'].isin(cus_code), 'net_profit'].sum())} | {cp_np_pct}%"
            tbl_cust_gp_td.cells[
                1].text = f"GP-{number_format(num=cy_ytd_profit_cus.loc[cy_ytd_profit_cus['customer_code'].isin(cus_code), 'amount'].sum())} | {ytd_gp_pct}% \nNP-{number_format(num=cy_ytd_net_profit_cus.loc[cy_ytd_net_profit_cus['customer_code'].isin(cus_code), 'net_profit'].sum())} | {ytd_np_pct}%"
            cy_cp_roi = cy_cp_profit_cus.loc[cy_cp_profit_cus['customer_code'].isin(cus_code), 'amount'].sum() / \
                        monthly_rev.loc[(monthly_rev['Month'] == end_date), 'Gross Rev'].iloc[0] * 100
            tbl_cust_gp_td.cells[2].text = str(f"{round(number=cy_cp_roi, ndigits=1)}%")
            cy_ytd_roi = cy_cp_profit_cus.loc[cy_ytd_profit_cus['customer_code'].isin(cus_code), 'amount'].sum() / \
                         monthly_rev['Gross Rev'].sum() * 100
            tbl_cust_gp_td.cells[3].text = str(f"{round(number=cy_ytd_roi, ndigits=1)}%")
        elif database == 'elite_security':
            tbl_cust_gp_td.cells[
                0].text = f"GP-{number_format(num=cy_cp_profit_cus.loc[cy_cp_profit_cus['customer_code'].isin(cus_code), 'amount'].sum())} | {cp_gp_pct}%"
            tbl_cust_gp_td.cells[
                1].text = f"GP-{number_format(num=cy_ytd_profit_cus.loc[cy_ytd_profit_cus['customer_code'].isin(cus_code), 'amount'].sum())} | {ytd_gp_pct}%"
            tbl_cust_gp_td.cells[2].text = str(customer_info[customer]['cy_cp_roi'])
            tbl_cust_gp_td.cells[3].text = str(customer_info[customer]['cy_ytd_roi'])
        elif database == 'nbn_realestate':
            pass
        else:
            pass
        table_formatter(table_name=tbl_cust_gp, style_name='table_style_1', special=[])

        fig, ((age_tbl, age_pie), (rev_tbl, rev_bar)) = plt.subplots(nrows=2, ncols=2, figsize=(8, 5))

        ageing: pd.DataFrame = customer_info[customer]['ageing']

        ageing.reset_index(inplace=True)

        if not ageing.empty:
            age_tbl.set_title('Receivable Ageing', loc='left', **heading_format)
            age_tbl.table(cellText=[[i[0], f'{i[1]:,.0f}'] for i in ageing.values], colLabels=ageing.columns,
                          cellLoc='center', loc='center')
            age_tbl.axis('off')

            age_pie.pie(x=ageing.loc[ageing['amount'] >= 0, 'amount'],
                        labels=ageing.loc[ageing['amount'] >= 0, 'Age Bracket'], autopct='%.1f%%')
            age_pie.axis('off')
        else:
            age_tbl.text(s='Zero Balance', x=0.5, y=0.5, ha='center', va='center', fontsize=28)
            age_tbl.axis('off')
            age_pie.text(s='Zero Balance', x=0.5, y=0.5, ha='center', va='center', fontsize=28)
            age_pie.axis('off')
        rev_tbl.set_title('Monthly Sales', loc='left', **heading_format)
        if database == 'nbn_logistics':
            rev_tbl.table(cellText=[[i[0].strftime('%B'), f'{i[1]:,.0f}', f'{i[2]:,.0f}'] for i in monthly_rev.values],
                        colLabels=monthly_rev.columns, cellLoc='center', loc='center')
        elif database == 'elite_security':
            rev_tbl.table(cellText=[[i[0].strftime('%B'), f'{i[1]:,.0f}'] for i in monthly_rev.values],
                        colLabels=monthly_rev.columns, cellLoc='center', loc='center')
        elif database == 'nbn_realestate':
            pass
        else:
            pass
        rev_tbl.axis('off')
        cust_profit_ytd: pd.DataFrame = cy_ytd_profit_cus.loc[
            cy_ytd_profit_cus['customer_code'].isin(cus_code), ['voucher_date', 'amount']]
        monthly_rev.set_index(keys='Month', inplace=True)
        monthly_rev.rename(columns={'Net Rev': 'Revenue'}, inplace=True)
        cust_profit_ytd = cust_profit_ytd.set_index(keys='voucher_date').rename(columns={'amount': 'GP'})
        monthly_rev_ytd: pd.DataFrame = pd.concat([monthly_rev, cust_profit_ytd], axis=1)

        monthly_rev_ytd = monthly_rev_ytd[['GP', 'Revenue']]

        monthly_rev_ytd.plot(kind='bar', stacked=True, ax=rev_bar)
        rev_bar.set_xticklabels([date.strftime('%b') for date in monthly_rev_ytd.index], rotation=0)
        gp_line = rev_bar.twinx()
        monthly_rev_ytd.loc[:, 'GP %'] = monthly_rev_ytd['GP'] / monthly_rev_ytd['Revenue']
        gp_line.plot([date.strftime('%b') for date in monthly_rev_ytd.index], monthly_rev_ytd['GP %'], color='black',
                     label='Customer GP')
        gp_line.legend()
        tick_locations = rev_bar.get_yticks()
        rev_bar.yaxis.set_major_locator(FixedLocator(tick_locations))
        rev_bar.yaxis.set_major_formatter(FixedFormatter(['{:,}'.format(int(i)) for i in tick_locations]))

        buf = BytesIO()
        plt.tight_layout()
        plt.savefig(buf, format='png')
        plt.close(fig)
        buf.seek(0)
        document.add_picture(buf)

        document.add_page_break()

    salesperson_list: np.ndarray = fInvoices.loc[(fInvoices['invoice_date'] <= end_date) & (
            fInvoices['invoice_date'] >= datetime(year=end_date.year, month=1, day=1)), 'emp_id'].unique()

    toc_salseperson(dEmployee=dEmployee, document=document, salesperson_list=salesperson_list)
    salespersonstats(salesperson_list=salesperson_list, document=document, dEmployee=dEmployee,
                     profitability=profitability, fInvoices=fInvoices, fGL=fGL, end_date=end_date)
    return profitability


def topcustomers(fInvoices: pd.DataFrame, end_date: datetime, mode: str, div: str, type: str, cnt: int) -> pd.DataFrame:
    if mode.lower() == 'month':
        start_date: datetime = datetime(
            year=end_date.year, month=end_date.month, day=1)
    elif mode.lower() == 'ytd':
        start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    else:
        raise ValueError(f'Invalid mode{mode}')

    if div.lower() == 'guarding':
        pattern = 'CTR'
    elif div.lower() == 'elv':
        pattern = 'ORD|CRD'
    else:
        raise ValueError(f'Invalid div{div}')
    topfivecustomers: pd.DataFrame = fInvoices.loc[
        (fInvoices['invoice_date'] >= start_date) & (fInvoices['invoice_date'] <= end_date) & (
                fInvoices['type'] == type) & (
            fInvoices['order_id'].str.contains(pat=pattern)), [
            'amount', 'cus_name']].groupby('cus_name').sum().sort_values(by='amount', ascending=False).head(
        cnt).reset_index().rename(columns={'cus_name': 'Customer', 'amount': 'Amount'})
    total_row: pd.DataFrame = pd.DataFrame(data={'Customer': ['Total'], 'Amount': [topfivecustomers['Amount'].sum()]})
    if not topfivecustomers.empty:
        topfivecustomers = pd.concat([topfivecustomers, total_row], ignore_index=True)
    return topfivecustomers


def revenue_change(fInvoices: pd.DataFrame, end_date: datetime, mode: str, order: bool) -> pd.DataFrame:
    start_date: datetime = datetime(year=end_date.year, month=end_date.month, day=1)
    pp_end: datetime = start_date - timedelta(days=1)
    pp_start: datetime = datetime(year=pp_end.year, month=pp_end.month, day=1)
    py_cp_start: datetime = datetime(year=end_date.year - 1, month=end_date.month, day=1)
    py_cp_end: datetime = py_cp_start + relativedelta(day=31)

    cy_cp: pd.Dataframe = fInvoices.loc[
        (fInvoices['invoice_date'] >= start_date) & (fInvoices['invoice_date'] <= end_date), ['amount',
                                                                                              'cus_name']].groupby(
        by='cus_name').sum().rename(columns={'amount': 'cycp'})
    cy_pp: pd.DataFrame = fInvoices.loc[
        (fInvoices['invoice_date'] >= pp_start) & (fInvoices['invoice_date'] <= pp_end), ['amount',
                                                                                          'cus_name']].groupby(
        by='cus_name').sum().rename(columns={'amount': 'cypp'})
    py_cp: pd.DataFrame = fInvoices.loc[
        (fInvoices['invoice_date'] >= py_cp_start) & (fInvoices['invoice_date'] <= py_cp_end), ['amount',
                                                                                                'cus_name']].groupby(
        by='cus_name').sum().rename(columns={'amount': 'pycp'})

    revenue_period: pd.DataFrame = pd.concat([cy_cp, cy_pp, py_cp], axis=1).fillna(0).reset_index().rename(
        columns={'cus_name': 'Customer'})

    revenue_period['Variance'] = revenue_period.apply(lambda x: x['cycp'] - x[f'{mode}'], axis=1)
    revenue_period.sort_values(by='Variance', ascending=order, inplace=True)
    revenue_period.drop(columns=['cycp', 'cypp', 'pycp'], inplace=True)
    total_row: pd.DataFrame = pd.DataFrame(
        data={'Customer': ['Total'], 'Variance': [revenue_period.head(5)['Variance'].sum()]})
    revenue_period = pd.concat([revenue_period.head(5), total_row], ignore_index=True)
    if order:
        revenue_period['Variance'] = revenue_period.apply(lambda x: x['Variance'] * -1, axis=1)
    return revenue_period


def revenue_movement(fInvoices: pd.DataFrame, end_date: datetime, document, database: str):
    rows_report_1: int = 4
    cols_report_1: int = 2
    if database == 'elite_security':
        cp_in_guard_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='month',
                                                    div='guarding',
                                                    type='Related', cnt=5)
        cp_in_elv_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='month', div='elv',
                                                  type='Related', cnt=5)
        cp_ex_guard_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='month',
                                                    div='guarding',
                                                    type='Market', cnt=5)
        cp_ex_elv_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='month', div='elv',
                                                  type='Market', cnt=5)
        ytd_in_guard_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='ytd', div='guarding',
                                                     type='Related', cnt=5)
        ytd_in_elv_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='ytd', div='elv',
                                                   type='Related', cnt=5)
        ytd_ex_guard_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='ytd', div='guarding',
                                                     type='Market', cnt=5)
        ytd_ex_elv_df: pd.DataFrame = topcustomers(fInvoices=fInvoices, end_date=end_date, mode='ytd', div='elv',
                                                   type='Market',
                                                   cnt=5)
        cus_info_1 = {0: {0: {'name': 'Current Month Internal Guarding (Top Five)', 'df': cp_in_guard_df},
                          1: {'name': 'Current Month Internal ELV (Top Five)', 'df': cp_in_elv_df}},
                      1: {0: {'name': '\nCurrent Month External Guarding (Top Five)', 'df': cp_ex_guard_df},
                          1: {'name': '\nCurrent Month External ELV (Top Five)', 'df': cp_ex_elv_df}},
                      2: {0: {'name': '\nYear to Date Internal Guarding (Top Five)', 'df': ytd_in_guard_df},
                          1: {'name': '\nYear to Date Internal ELV (Top Five)', 'df': ytd_in_elv_df}},
                      3: {0: {'name': '\nYear to Date External Guarding (Top Five)', 'df': ytd_ex_guard_df},
                          1: {'name': '\nYear to Date External ELV (Top Five)', 'df': ytd_ex_elv_df}}}

        keydatacus1 = document.add_table(rows=rows_report_1, cols=cols_report_1)

        for row in range(rows_report_1):
            for col in range(cols_report_1):
                row_0 = keydatacus1.rows[row].cells
                row_0[col].text = cus_info_1[row][col]['name']
                df: pd.DataFrame = cus_info_1[row][col]['df']
                inner_tbl = row_0[col].add_table(rows=1, cols=2)
                inner_tbl_hdr = inner_tbl.rows[0].cells
                inner_tbl_hdr[0].text = 'Customer'
                inner_tbl_hdr[1].text = 'Amount'
                for _, j in df.iterrows():
                    cells = inner_tbl.add_row().cells
                    cells[0].text = str(j['Customer'])
                    cells[1].text = number_format(j.iloc[1])
                table_formatter(table_name=inner_tbl, style_name='table_style_1', special=['Total'])
        document.add_page_break()

    inc_pp: pd.DataFrame = revenue_change(fInvoices=fInvoices, end_date=end_date, mode='cypp', order=True)
    dec_pp: pd.DataFrame = revenue_change(fInvoices=fInvoices, end_date=end_date, mode='cypp', order=False)
    inc_py: pd.DataFrame = revenue_change(fInvoices=fInvoices, end_date=end_date, mode='pycp', order=True)
    dec_py: pd.DataFrame = revenue_change(fInvoices=fInvoices, end_date=end_date, mode='pycp', order=False)

    cus_info_2 = {0: {0: {'name': 'Top 5 Customers with Increased\nRevenue compared to previous month', 'df': inc_pp},
                      1: {'name': 'Top 5 Customers with Decreased\nRevenue compared to previous month', 'df': dec_pp}},
                  1: {0: {'name': '\nTop 5 Customers with Increased\nRevenue compared to previous year', 'df': inc_py},
                      1: {'name': '\nTop 5 Customers with Decreased\nRevenue compared to previous year', 'df': dec_py}}}
    rows_report_2: int = 2
    cols_report_2: int = 2
    keydatacus2 = document.add_table(rows=rows_report_1, cols=cols_report_1)  # r4,c2

    for row in range(rows_report_2):
        for col in range(cols_report_2):
            row_0 = keydatacus2.rows[row].cells
            row_0[col].text = cus_info_2[row][col]['name']
            df: pd.DataFrame = cus_info_2[row][col]['df']
            inner_tbl = row_0[col].add_table(rows=1, cols=2)
            inner_tbl_hdr = inner_tbl.rows[0].cells
            inner_tbl_hdr[0].text = 'Customer'
            inner_tbl_hdr[1].text = 'Amount'
            for _, j in df.iterrows():
                cells = inner_tbl.add_row().cells
                cells[0].text = str(j['Customer'])
                cells[1].text = number_format(j.iloc[1])
            table_formatter(table_name=inner_tbl, style_name='table_style_1', special=['Total'])
    document.add_page_break()


def service_period(doj: datetime, end_date: datetime) -> str:
    days_passed: int = (end_date - doj).days
    service_ranges: list = [(365, '< One Year'), (730, '1-2 Years'), (1095, '2-3 Years'),
                            (1460, '3-4 Years'), (float('inf'), '4 Years +')]
    for threshold, label in service_ranges:
        if days_passed <= threshold:
            return label
            break


def emp_age(dob: datetime, end_date: datetime) -> str:
    age: int = end_date.year - dob.year - ((end_date.month, end_date.day) < (dob.month, dob.day))
    service_ranges: list = [(25, '< 25'), (35, '26-35 Years'), (45, '36-45 Years'), (float('inf'), '46 +')]
    for threshold, label in service_ranges:
        if age <= threshold:
            return label
            break


def employee_related(dEmployee: pd.DataFrame, end_date: datetime, database: str) -> dict:
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)
    total_pie_slices: int = 5
    dEmployee['termination_date'] = pd.to_datetime(dEmployee['termination_date'])
    dEmployee['dob'] = pd.to_datetime(dEmployee['dob'])
    dEmployee['doj'] = pd.to_datetime(dEmployee['doj'])
    dEmployee = dEmployee.loc[(~dEmployee['emp_id'].isin(['ESS0015-OLD', 'ESS0016'])) & (
        dEmployee['emp_id'].str.contains('ESS|NBNL')) & (dEmployee['doj'] <= end_date) & (
                                      (dEmployee['termination_date'] >= start_date) | (
                                  dEmployee['termination_date'].isna()))]
    emp_types: dict = {'MGMT': 'Staff', 'STAFF': 'Staff', 'ELV STAFF': 'Staff', 'LABOUR': 'Labour',
                       'LABOUR A': 'Labour',
                       'LABOUR A 2': 'Labour', 'LABOUR A 3': 'Labour', 'LABOUR A 4': 'Labour', 'ELV LABOUR': 'Labour'}
    current_emp: pd.DataFrame = dEmployee.loc[
        (dEmployee['termination_date'] > end_date) | (dEmployee['termination_date'].isna())]
    gender: dict = current_emp.value_counts(subset='sex').to_dict()

    type: list = [emp_types[i] for i in current_emp['emp_type'].tolist()]
    type: dict = {item: type.count(item) for item in set(type)}

    if database == 'elite_security':
        dept: list = [i if i == 'ELV' else 'Guarding' for i in current_emp['dept']]
    elif database == 'nbn_logistics':
        dept: list = current_emp['dept'].tolist()
    elif database == 'nbn_realestate':
        pass
    else:
        pass
    dept: dict = {item: dept.count(item) for item in set(dept)}

    designation: list = current_emp['designation'].tolist()
    designation: dict = {item.title(): designation.count(item) for item in set(designation)}
    designation = dict(sorted(designation.items(), key=lambda item: item[1], reverse=True))
    d_sliced: dict = dict(islice(designation.items(), total_pie_slices - 1))
    d_sliced['Others'] = sum(
        [i[1] for i in dict(islice(designation.items(), total_pie_slices - 1, len(designation))).items()])

    nationality: list = [i for i in current_emp['nationality'].tolist()]
    nationality: dict = {item.title(): nationality.count(item) for item in set(nationality)}
    nationality = dict(sorted(nationality.items(), key=lambda item: item[1], reverse=True))
    n_sliced: dict = dict(islice(nationality.items(), total_pie_slices - 1))
    n_sliced['Others'] = sum(
        [i[1] for i in dict(islice(nationality.items(), total_pie_slices - 1, len(nationality))).items()])

    opening_emp = len(dEmployee.loc[(dEmployee['doj'] <= start_date - timedelta(days=1))])
    joined_emp: int = len(dEmployee.loc[(dEmployee['doj'] >= start_date)])
    resigned_emp: int = -len(dEmployee.loc[(dEmployee['termination_date'] <= end_date)])
    closing_emp: int = opening_emp + joined_emp + resigned_emp
    emp_movement: dict = {f'Employees at {start_date.date()}': opening_emp,
                          'New Joiners': joined_emp, 'Resigned Employees': resigned_emp,
                          f'Employees at {end_date.date()}': closing_emp}
    current_emp.loc[:, 'Service'] = current_emp.loc[:, 'doj'].apply(func=service_period, args=[end_date])
    current_emp.loc[:, 'Age'] = current_emp.loc[:, 'dob'].apply(func=emp_age, args=[end_date])

    service: list = current_emp['Service'].tolist()
    service: dict = {item: service.count(item) for item in set(service)}

    age: list = current_emp['Age'].tolist()
    age: dict = {item: age.count(item) for item in set(age)}

    df_new_joiner: pd.DataFrame = dEmployee.loc[dEmployee['doj'] >= start_date, ['doj', 'emp_id']].rename(
        columns={'doj': 'Period', 'emp_id': 'Joined'})
    new_joiners = df_new_joiner.groupby(pd.Grouper(key='Period', freq='ME')).count()

    df_resigned: pd.DataFrame = dEmployee.loc[
        dEmployee['termination_date'] <= end_date, ['emp_id', 'termination_date']].rename(
        columns={'termination_date': 'Period', 'emp_id': 'Resigned'})
    emp_resigned = df_resigned.groupby(pd.Grouper(key='Period', freq='ME')).count()

    total_employees: pd.DataFrame = pd.concat([new_joiners, emp_resigned], axis=1)
    total_employees['Total Employees'] = (total_employees['Joined'] - total_employees['Resigned'])
    total_employees.drop(columns=['Joined', 'Resigned'], inplace=True)
    total_employees['Total Employees'] = total_employees['Total Employees'].cumsum() + opening_emp

    employee_data: dict = {'Gender': gender, 'Type': type, 'Department': dept, 'Nationality': n_sliced,
                           'Employee Age': age, 'Service Period': service, 'Designation': d_sliced,
                           'Employee Movement': emp_movement,
                           'new_joiner': new_joiners, 'resigned_emp': emp_resigned, 'total_employees': total_employees}
    return employee_data


def hrrelated(document, dEmployee: pd.DataFrame, database: str, end_date: datetime):
    change_orientation(document=document, method='l')
    emp_data: dict = employee_related(dEmployee=dEmployee, database=database, end_date=end_date)
    plt.style.use('ggplot')
    hr_fig_1, ((gender, type), (dept, nationality)) = plt.subplots(nrows=2, ncols=2, figsize=(10.5, 7))

    gender.set_title('Gender')
    gender.pie(x=list(emp_data['Gender'].values()), labels=list(emp_data['Gender'].keys()), autopct='%.0f%%',
               labeldistance=1, pctdistance=0.3)
    gender.axis('off')

    type.set_title('Category')
    type.pie(x=list(emp_data['Type'].values()), labels=list(emp_data['Type'].keys()), autopct='%.0f%%', labeldistance=1,
             pctdistance=0.3)
    type.axis('off')

    dept.set_title('Department')
    dept.pie(x=list(emp_data['Department'].values()), labels=list(emp_data['Department'].keys()), autopct='%.0f%%',
             labeldistance=1, pctdistance=0.3)
    dept.axis('off')

    nationality.set_title('Nationality')
    nationality.pie(x=list(emp_data['Nationality'].values()), labels=list(emp_data['Nationality'].keys()),
                    autopct='%.0f%%',
                    labeldistance=1, pctdistance=0.3)
    nationality.axis('off')

    hr_graph_1_buf = BytesIO()
    plt.tight_layout()
    plt.savefig(hr_graph_1_buf, format='png')
    plt.close(hr_fig_1)
    hr_graph_1_buf.seek(0)
    document.add_picture(hr_graph_1_buf)
    document.add_page_break()

    hr_fig_2, ((age, service), (designation, movement)) = plt.subplots(nrows=2, ncols=2, figsize=(10.5, 7))
    age.set_title('Age')
    age.pie(x=list(emp_data['Employee Age'].values()), labels=list(emp_data['Employee Age'].keys()), autopct='%.0f%%',
            labeldistance=1, pctdistance=0.3)
    age.axis('off')

    service.set_title('Service')
    service.pie(x=list(emp_data['Service Period'].values()), labels=list(emp_data['Service Period'].keys()),
                autopct='%.0f%%', labeldistance=1, pctdistance=0.3)
    service.axis('off')

    designation.set_title('Designation')
    designation.pie(x=list(emp_data['Designation'].values()), labels=list(emp_data['Designation'].keys()),
                    autopct='%.0f%%',
                    labeldistance=1, pctdistance=0.7)
    designation.axis('off')

    emp_move = pd.DataFrame(list(emp_data['Employee Movement'].items()), columns=['Description', '# of Emp'])

    movement.set_title('Movement')
    movement.table(cellText=[i for i in emp_move.values], colLabels=emp_move.columns, cellLoc='center', loc='center')
    movement.axis('off')

    hr_graph_2_buf = BytesIO()
    plt.tight_layout()
    plt.savefig(hr_graph_2_buf, format='png')
    plt.close(hr_fig_2)
    hr_graph_2_buf.seek(0)
    document.add_picture(hr_graph_2_buf)
    document.add_page_break()

    hr_fig_3, (jo_re, total_staff) = plt.subplots(nrows=2, ncols=1, sharex=True, figsize=(7.27, 10))

    jo_re.set_title('New Joiners and Leavers')
    jo_re.plot([i.strftime('%b') for i in emp_data['new_joiner'].index], emp_data['new_joiner']['Joined'],
               label='New Joiners')
    jo_re.plot([i.strftime('%b') for i in emp_data['new_joiner'].index], emp_data['resigned_emp']['Resigned'],
               label='Resigned')
    jo_re.legend()

    total_staff.set_title('Total Manpower')
    total_staff.plot([i.strftime('%b') for i in emp_data['new_joiner'].index],
                     emp_data['total_employees']['Total Employees'], label='Total Employees')

    change_orientation(document=document, method='p', )
    hr_graph_3_buf = BytesIO()
    plt.tight_layout()
    plt.savefig(hr_graph_3_buf, format='png')
    plt.close(hr_fig_3)
    hr_graph_3_buf.seek(0)
    document.add_picture(hr_graph_3_buf)
    document.add_page_break()


def operations(ftimesheet: pd.DataFrame, financial: pd.DataFrame, end_date: datetime, dExclude: pd.DataFrame) -> dict:
    start_date: datetime = datetime(year=end_date.year, month=1, day=1)

    ftimesheet: pd.DataFrame = ftimesheet.loc[
        (ftimesheet['v_date'] >= start_date) & (ftimesheet['v_date'] <= end_date) & (
            ~ftimesheet['order_id'].isin(['discharged', 'not_joined']))]
    ftimesheet.loc[:, 'v_date'] = ftimesheet.apply(lambda x: x['v_date'] + relativedelta(day=31), axis=1)
    df_transport: pd.DataFrame = ftimesheet.copy()
    df_accommodation: pd.DataFrame = ftimesheet.copy()
    df_unproductive: pd.DataFrame = ftimesheet.copy()
    df_transport = df_transport.loc[
        ~df_transport['order_id'].isin(dExclude.loc[dExclude['dc_trpt'] == True, 'job_type'].tolist())]
    df_unproductive = df_unproductive.loc[
        df_unproductive['order_id'].isin(dExclude.loc[dExclude['dc_emp_beni'] == False, 'job_type'].tolist())]

    df_transport = df_transport.groupby(by=['v_date'])['cost_center'].count().reset_index().rename(
        columns={'cost_center': 'trpt_md'})
    df_accommodation = df_accommodation.groupby(by=['v_date'])['cost_center'].count().reset_index().rename(
        columns={'cost_center': 'acco_md'})
    df_unproductive = df_unproductive.groupby(by=['v_date'])['cost_center'].count().reset_index().rename(
        columns={'cost_center': 'unproductive_md'})
    financial: pd.DataFrame = financial.loc[
        financial['Description'].isin(['Transportation - Manpower', 'Accommodation - Manpower'])]
    trpt = financial.index[financial['Description'] == 'Transportation - Manpower'][0]
    acc = financial.index[financial['Description'] == 'Accommodation - Manpower'][0]
    financial = financial.transpose().reset_index().rename(columns={trpt: 'Transport', acc: 'Accommodation'}).drop(0)
    financial.loc[:, 'Transport'] = financial['Transport'] * -1
    financial.loc[:, 'Accommodation'] = financial['Accommodation'] * -1
    financial['voucher_date'] = pd.to_datetime(financial['voucher_date'], format='%m/%d/%Y %H:%M')
    operations: pd.DataFrame = pd.concat(
        [financial.set_index('voucher_date'), df_transport.set_index('v_date'), df_accommodation.set_index('v_date'),
         df_unproductive.set_index('v_date')], axis=1)
    operations.loc[:, 'productive_md'] = operations['acco_md'] - operations['unproductive_md']
    return operations


def opsrelated(fTimesheet: pd.DataFrame, end_date: datetime, dExclude: pd.DataFrame, financial: pd.DataFrame, document,
               profitability: dict):
    ops_data: pd.DataFrame = operations(ftimesheet=fTimesheet, financial=financial, end_date=end_date,
                                        dExclude=dExclude)
    plt.style.use('ggplot')
    fig_ops_1, (cost_line, ph_line) = plt.subplots(nrows=2, ncols=1, sharex=True, figsize=(7.27, 10))

    cost_line.set_title('Transportation and Accommodation Expenses')
    cost_line.plot([i.strftime('%b') for i in ops_data.index], ops_data['Transport'], label='Transport')
    for xy in zip([i.strftime('%b') for i in ops_data.index], ops_data['Transport'].tolist()):
        cost_line.annotate('{:,}K'.format(int(xy[1] / 1_000)), xy=xy)
    cost_line.plot([i.strftime('%b') for i in ops_data.index], ops_data['Accommodation'], label='Accommodation')
    for xy in zip([i.strftime('%b') for i in ops_data.index], ops_data['Accommodation'].tolist()):
        cost_line.annotate('{:,}K'.format(int(xy[1] / 1_000)), xy=xy)
    tick_locations = cost_line.get_yticks()
    cost_line.yaxis.set_major_locator(FixedLocator(tick_locations))
    cost_line.yaxis.set_major_formatter(FixedFormatter(['{:.0f}K'.format(i / 1_000) for i in tick_locations]))
    cost_line.legend()

    ph_line.set_title('Transportation and Accommodation Per Head')
    ph_line.plot([i.strftime('%b') for i in ops_data.index], (ops_data['Transport'] / ops_data['trpt_md']) * 30,
                 label='Transport')
    ph_line.plot([i.strftime('%b') for i in ops_data.index], (ops_data['Accommodation'] / ops_data['acco_md']) * 30,
                 label='Accommodation')
    tick_locations = ph_line.get_yticks()
    ph_line.yaxis.set_major_locator(FixedLocator(tick_locations))
    ph_line.yaxis.set_major_formatter(FixedFormatter(['{:,}'.format(i) for i in tick_locations]))
    ph_line.legend()

    ops_graph_1_buf = BytesIO()
    plt.tight_layout()
    plt.savefig(ops_graph_1_buf, format='png')
    plt.close(fig_ops_1)
    ops_graph_1_buf.seek(0)
    document.add_picture(ops_graph_1_buf)
    document.add_page_break()

    fig_ops_2, (bill_nonbil, efficiency, non_billable) = plt.subplots(nrows=3, ncols=1, figsize=(7.73, 10.63),
                                                                      sharex=True,
                                                                      gridspec_kw={'height_ratios': [1, 1, 2]})

    bill_nonbil.set_title('Billable Vs Non-Billable Mandays')
    bill_nonbil.plot([i.strftime('%b') for i in ops_data.index], ops_data['productive_md'], label='Productive')
    bill_nonbil.plot([i.strftime('%b') for i in ops_data.index], ops_data['unproductive_md'], label='Un-productive')
    tick_locations = bill_nonbil.get_yticks()
    bill_nonbil.yaxis.set_major_locator(FixedLocator(tick_locations))
    bill_nonbil.yaxis.set_major_formatter(FixedFormatter(['{:.0f}K'.format(i / 1_000) for i in tick_locations]))
    bill_nonbil.legend()

    efficiency.set_title('Manpower Utilization Efficiency')

    efficiency.plot([i.strftime('%b') for i in ops_data.index], (ops_data['productive_md'] / ops_data['acco_md']) * 100,
                    label='Efficiency')
    tick_locations = efficiency.get_yticks()
    efficiency.yaxis.set_major_locator(FixedLocator(tick_locations))
    efficiency.yaxis.set_major_formatter(FixedFormatter(['{:,.0f}%'.format(int(i)) for i in tick_locations]))
    efficiency.legend()

    periods = pd.date_range(start=datetime(year=end_date.year, month=1, day=1), end=end_date,
                            freq='ME').to_pydatetime().tolist()

    c = {}
    exclude_dict = dExclude.groupby('group')['job_type'].apply(set).to_dict()
    for t in periods:
        period_allocation = profitability['periodic_allocation'].get(t, {})
        a = {}
        for group, job_types in exclude_dict.items():
            for job_type in job_types:
                if job_type in period_allocation:
                    a[group] = a.get(group, 0) + period_allocation[job_type]
        a = {k: v for k, v in a.items() if v != 0}
        c[t] = a
    results_df = pd.DataFrame.from_dict(c, orient='index').fillna(0) * -1

    non_billable.set_title('Non-Billable Cost')
    for p in results_df.columns:
        non_billable.plot([i.strftime('%b') for i in results_df.index], results_df[p], label=p)
    tick_locations = non_billable.get_yticks()
    non_billable.yaxis.set_major_locator(FixedLocator(tick_locations))
    non_billable.yaxis.set_major_formatter(FixedFormatter(['{:,}'.format(int(i)) for i in tick_locations]))
    non_billable.legend()

    ops_graph_2_buf = BytesIO()
    plt.tight_layout()
    plt.savefig(ops_graph_2_buf, format='png')
    plt.close(fig_ops_2)
    ops_graph_2_buf.seek(0)
    document.add_picture(ops_graph_2_buf)
    document.add_page_break()


def cohart(fSalesTill2020: pd.DataFrame, end_date: datetime, fInvoices: pd.DataFrame, document,
           dCustomer: pd.DataFrame) -> dict:
    fInvoices: pd.DataFrame = fInvoices.loc[
        (fInvoices['invoice_date'] <= end_date), ['invoice_date', 'amount', 'customer_code']]

    fInvoices['invoice_date'] = fInvoices['invoice_date'] + pd.offsets.MonthEnd(0)
    fSalesTill2020: pd.DataFrame = fSalesTill2020.loc[
        fSalesTill2020['ledger_code'].isin([4010201001, 4010201002, 4010201003, 4010201004]), ['invoice_date', 'amount',
                                                                                               'customer_code']]

    invoice_combined: pd.DataFrame = pd.concat([fInvoices, fSalesTill2020])
    invoice_combined.sort_values(by='invoice_date', inplace=True)

    min_date: datetime = invoice_combined['invoice_date'].min().date()
    max_date: datetime = invoice_combined['invoice_date'].max().date()
    cohort_qty: pd.DataFrame = pd.DataFrame()
    cohort_value: pd.DataFrame = pd.DataFrame()
    no_months: int = (relativedelta(max_date, min_date).months +
                      relativedelta(max_date, min_date).years * 12) + 1
    dates_list: list = list(invoice_combined['invoice_date'].unique())
    for i in range(no_months):

        for j in range(i, no_months):  # column like 0,1,2
            pre_period_filt = (invoice_combined['invoice_date'] <= dates_list[i if i - 1 <
                                                                                   0 else i - 1])
            pre_period_customers: list = list(
                set(invoice_combined.loc[pre_period_filt, 'customer_code'].tolist()))

            ini_period_filt = (invoice_combined['invoice_date'] == dates_list[i])
            ini_period_customers = list(
                set(invoice_combined.loc[ini_period_filt, 'customer_code'].tolist()))
            if i == 0:
                customers: list = [
                    value for value in ini_period_customers if value in pre_period_customers]
            else:
                customers = [
                    value for value in ini_period_customers if value not in pre_period_customers]

            current_period_filt = (invoice_combined['invoice_date'] == dates_list[j])
            current_period_customers: list = list(
                set(invoice_combined.loc[current_period_filt, 'customer_code'].tolist()))

            in_both: list = [
                value for value in current_period_customers if value in customers]
            cohort_qty.loc[i, j] = len(in_both)

            cohort_value_filt = (invoice_combined['customer_code'].isin(in_both)) & (
                    invoice_combined['invoice_date'] == dates_list[j])
            cohort_value.loc[i, j] = invoice_combined.loc[cohort_value_filt, 'amount'].sum()

    report_one: pd.DataFrame = cohort_qty.iloc[-12:, -12:]
    report_one.reset_index(inplace=True, drop=True)
    report_one.fillna(value=0, inplace=True)

    start_date: datetime = end_date - relativedelta(years=1) + timedelta(days=1)
    cols: list = list(report_one.columns)
    periods: list = [i.strftime('%b') for i in pd.date_range(start=start_date, end=end_date, freq='ME').to_pydatetime()]
    rename_dict: dict = {col: periods[i] for i, col in enumerate(cols)}
    report_one.rename(columns=rename_dict, inplace=True)
    index_list = pd.Series(periods)
    report_one = pd.concat([index_list, report_one], axis=1)
    report_one.rename(columns={0: 'Month'}, inplace=True)

    cols: int = cohort_qty.shape[0]
    rows: int = cohort_qty.shape[1]
    months = 12
    final_dict = {}
    for i in range(months):
        row_start = rows - months - i
        row_end = rows - i
        col_start = cols - months - i
        col_end = cols - i

        df1 = cohort_qty.iloc[row_start:row_end, col_start:col_end].reset_index(drop=True)
        total = 0
        for j in range(months):
            numerator: float = df1.iloc[j, 11]
            denominator: float = df1.iloc[j, j]

            if denominator != 0:
                amount: float = numerator / denominator * 100
            else:
                amount = 0
            total += amount

        avg = total / months
        final_dict[i] = round(number=avg)

    retention_dict: dict = {col: final_dict[i] for i, col in enumerate(periods)}

    year_before: datetime = end_date - relativedelta(years=1) + pd.offsets.MonthEnd(0)

    six_months_before: datetime = end_date - relativedelta(months=6) + pd.offsets.MonthEnd(0)

    beforeoneyear: np.ndarray = invoice_combined.loc[
        (invoice_combined['invoice_date'] <= year_before), 'customer_code'].unique()

    beforesixmonths: np.ndarray = invoice_combined.loc[((invoice_combined['invoice_date'] > year_before) & (
            invoice_combined['invoice_date'] <= six_months_before)), 'customer_code'].unique()

    withinsixmonths: np.ndarray = invoice_combined.loc[((invoice_combined['invoice_date'] > six_months_before) & (
            invoice_combined['invoice_date'] <= end_date)), 'customer_code'].unique()

    dormant: list = [i for i in [j for j in beforesixmonths if j not in beforeoneyear] if i not in withinsixmonths]

    dormant: pd.DataFrame = \
        invoice_combined.loc[invoice_combined['customer_code'].isin(dormant), ['customer_code', 'amount']].groupby(
            by='customer_code', as_index=False)['amount'].sum().sort_values('amount', ascending=False).head(10)

    dormant = pd.merge(left=dormant, right=dCustomer[['customer_code', 'cus_name']], on='customer_code',
                       how='left').drop(columns=['customer_code'])
    dormant.rename(columns={'cus_name': 'Customer', 'amount': 'Gross Rev'}, inplace=True)

    cohort, (cohort_grid, cohort_graph, dorment_cust) = plt.subplots(nrows=3, ncols=1, figsize=(7.73, 10.63))
    cohort_grid.table(cellText=[[j[0]] + [i for i in j if isinstance(i, float)] for j in report_one.values],
                      colLabels=[i for i in report_one.columns],
                      cellLoc='left', loc='best')
    cohort_grid.set_title('Customer retention last 12 months')
    cohort_grid.axis('off')
    cohort_graph.set_title('Average Retention over last 12 months (%)')
    cohort_graph.plot(list(retention_dict.keys()), list(retention_dict.values()))
    dorment_cust.set_title('Dormant Customers for last 6 months')
    dorment_cust.table(cellText=[[' '.join(i[1].title().split(sep=' ')[:2]), f'{i[0]:,.0f}'] for i in dormant.values],
                       colLabels=[i for i in dormant.columns],
                       cellLoc='left', loc='best')

    dorment_cust.axis('off')

    cohort_data = BytesIO()
    plt.tight_layout()
    plt.savefig(cohort_data, format='png')
    plt.close(cohort)
    cohort_data.seek(0)
    document.add_picture(cohort_data)
    document.add_page_break()


def empctc(row, dEmployee: pd.DataFrame) -> float:
    special_emp:dict = {'emp_id':'NBNL0088',
                        'benefits':{'ticket':{'self':0,
                                                'spouse':1,
                                                'dependent':2},
                                    'insurance':{'self':1,
                                                'spouse':1,
                                                'dependent':2}}}
    policy: str = dEmployee.loc[(dEmployee['emp_id'] == row['emp_id']), 'leave_policy']
    basic: float = dEmployee.loc[(dEmployee['emp_id'] == row['emp_id']), 'ba']
    gross: float = \
        dEmployee.loc[
            (dEmployee['emp_id'] == row['emp_id']), ['ba', 'hra', 'tra', 'ma', 'oa', 'pda']].sum(
            axis=1).values[0]
    ticket_amt: float = dEmployee.loc[(dEmployee['emp_id'] == row['emp_id']), 'travel_cost']
    try:
        leave_policy, ticket_policy = int(policy.split(sep='-')[1].strip().split(sep=' ')[0]), \
            policy.split(sep='-')[-1].strip().split(sep=' ')[0]
    except IndexError:
        leave_policy, ticket_policy = 1, 1
    eos: float = basic * 12 / 365 * 30 / 12
    leave: float = gross * 12 / 365 * leave_policy / 12
    ticket: float = ticket_amt / (12 if ticket_policy == 'Yearly' else 24)
    ctc: float = gross + eos + leave + ticket
    return ctc
